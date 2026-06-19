import base64
import html
import json
import os
import re
import unicodedata
from io import BytesIO
from datetime import datetime
from collections import OrderedDict

from fastapi import HTTPException, Request

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.config import INFORMES_DIR, UPLOAD_DIR
from app.database import get_connection
from app.services.valoracion_comparacion import (
    preparar_matriz_homogeneizacion,
    preparar_resumen_comparacion,
    preparar_testigo_comparacion,
)


def limpiar_nombre_archivo(texto: str) -> str:
    texto = (texto or "").strip().replace(" ", "_")
    texto = re.sub(r"[^A-Za-z0-9_\-]", "", texto)
    return texto or "expediente"


def valor_o_guion(valor):
    if valor is None:
        return "-"
    texto = str(valor).strip()
    return texto if texto else "-"


def limpiar_texto(valor) -> str:
    return str(valor or "").strip()


def parsear_float_valoracion(valor):
    if valor is None:
        return None
    if isinstance(valor, (int, float)):
        return float(valor)
    texto = limpiar_texto(valor)
    if not texto:
        return None
    texto = (
        texto.replace("€", "")
        .replace("EUR/m2", "")
        .replace("€/m²", "")
        .replace("m²", "")
        .replace("m2", "")
        .strip()
    )
    texto = texto.replace(" ", "")
    if "," in texto and "." in texto:
        texto = texto.replace(".", "").replace(",", ".")
    elif "," in texto:
        texto = texto.replace(",", ".")
    try:
        return float(texto)
    except ValueError:
        return None


def formatear_numero_es(valor, decimales: int = 0) -> str:
    numero = parsear_float_valoracion(valor)
    if numero is None:
        return valor_o_guion(valor)
    texto = f"{numero:,.{decimales}f}"
    return texto.replace(",", "X").replace(".", ",").replace("X", ".")


def formatear_moneda_es(valor) -> str:
    if parsear_float_valoracion(valor) is None:
        return valor_o_guion(valor)
    return f"{formatear_numero_es(valor, 0)} €"


def formatear_precio_unitario_es(valor) -> str:
    if parsear_float_valoracion(valor) is None:
        return valor_o_guion(valor)
    return f"{formatear_numero_es(valor, 0)} €/m²"


def formatear_porcentaje_es(valor) -> str:
    if parsear_float_valoracion(valor) is None:
        return valor_o_guion(valor)
    return f"{formatear_numero_es(valor, 2)}%"


def formatear_superficie_es(valor) -> str:
    if parsear_float_valoracion(valor) is None:
        return valor_o_guion(valor)
    return f"{formatear_numero_es(valor, 2)} m²"


def formatear_coeficiente_es(valor) -> str:
    if parsear_float_valoracion(valor) is None:
        return valor_o_guion(valor)
    return f"{formatear_numero_es(valor, 2)}x"


def formatear_booleano_es(valor) -> str:
    texto = limpiar_texto(valor).lower()
    if texto in {"", "-"}:
        return "-"
    return "Sí" if texto in {"1", "true", "si", "sí"} else "No"


def row_to_dict(row) -> dict:
    return dict(row) if row is not None else {}


def get_row_value(row, key: str, default=""):
    if row is None:
        return default
    try:
        if key in row.keys():
            return row[key]
    except AttributeError:
        return row.get(key, default)
    return default


def imagen_url_pdf(nombre_archivo: str | None, base_url: str = "") -> str:
    nombre = limpiar_texto(nombre_archivo)
    if not nombre:
        return ""
    prefijo = base_url.rstrip("/")
    return f"{prefijo}/uploads/{nombre}" if prefijo else f"/uploads/{nombre}"


def mapa_overlay_data_uri(mapa) -> str:
    nombre_foto = limpiar_texto(mapa.get("imagen_base"))
    if not nombre_foto:
        return ""
    ruta_foto = os.path.join(UPLOAD_DIR, nombre_foto)
    cuadrantes_con_incidencia = [
        cuadrante["codigo_cuadrante"] for cuadrante in mapa.get("cuadrantes", [])
    ]
    try:
        overlay = generar_overlay_cuadrantes_mapa(
            ruta_foto,
            mapa.get("filas"),
            mapa.get("columnas"),
            cuadrantes_con_incidencia,
        )
    except Exception:
        return ""
    return "data:image/png;base64," + base64.b64encode(overlay.getvalue()).decode("ascii")


def clave_orden_cuadrante(codigo):
    texto = limpiar_texto(codigo)
    match = re.fullmatch(r"([A-Za-z]+)(\d+)", texto)
    if not match:
        return (1, texto.upper(), 0)
    letras, numero = match.groups()
    return (0, letras.upper(), int(numero))


ESTADO_INSPECCION_LABELS = {
    "no_necesita_reparacion": "No necesita reparación",
    "necesita_reparacion": "Necesita reparación",
    "defecto_grave": "Defecto grave",
    "no_inspeccionado": "No inspeccionado",
}

INSPECCION_GENERAL_GROUPS = [
    (
        "Información general / conjunto del inmueble",
        [
            ("puerta_entrada", "Puerta de entrada"),
            ("vestibulo", "Vestíbulo"),
            ("ventilacion_cruzada", "Ventilación cruzada"),
            ("ventilacion_general_inmueble", "Ventilación general del inmueble"),
            ("iluminacion_natural_general", "Iluminación natural general"),
            ("orientacion_general", "Orientación general"),
            ("reformado_cambio_uso", "Reformado / cambio de uso"),
        ],
    ),
    (
        "Estructura",
        [
            ("estructura_vertical", "Estructura vertical"),
            ("estructura_horizontal", "Estructura horizontal"),
            ("forjados_voladizos", "Forjados / voladizos"),
            ("cubiertas", "Cubiertas"),
            ("soleras_losas", "Soleras / losas"),
        ],
    ),
    (
        "Instalaciones generales",
        [
            ("instalacion_electrica_general", "Instalación eléctrica general"),
            ("agua_acs", "Agua / ACS"),
            ("calefaccion", "Calefacción"),
            ("climatizacion", "Climatización"),
        ],
    ),
    (
        "Carpinterías generales",
        [
            ("carpinterias_generales", "Carpinterías generales"),
            ("persianas_generales", "Persianas generales"),
            ("barandillas_generales", "Barandillas generales"),
            ("vierteaguas_generales", "Vierteaguas generales"),
        ],
    ),
]

INSPECCION_ESTANCIA_BASE_ITEMS = [
    ("puerta", "Puerta"),
    ("revestimiento", "Revestimiento"),
    ("iluminacion", "Iluminación"),
    ("mobiliario", "Mobiliario"),
    ("mecanismos_electricos", "Mecanismos eléctricos"),
    ("humedades", "Humedades"),
    ("techo", "Techo"),
    ("pavimento", "Pavimento"),
]

INSPECCION_ESTANCIA_BANO_ITEMS = [
    ("banera_ducha", "Bañera / ducha"),
    ("mampara", "Mampara"),
    ("lavabo", "Lavabo"),
    ("inodoro", "Inodoro"),
    ("bide", "Bidé"),
    ("espejo", "Espejo"),
    ("ventilacion_forzada", "Ventilación forzada"),
    ("condensacion", "Condensación"),
    ("griferia", "Grifería"),
    ("sifones", "Sifones"),
    ("desagues", "Desagües"),
    ("llaves_paso", "Llaves de paso"),
]

INSPECCION_ESTANCIA_COCINA_ITEMS = [
    ("extractor", "Extractor"),
    ("encimera", "Encimera"),
    ("zona_coccion", "Zona de cocción"),
    ("frigorifico", "Frigorífico"),
    ("horno", "Horno"),
    ("fregadero", "Fregadero"),
    ("griferia", "Grifería"),
    ("sifones", "Sifones"),
    ("desagues", "Desagües"),
    ("llaves_paso", "Llaves de paso"),
    ("conexion_lavavajillas", "Conexión lavavajillas"),
]

INSPECCION_ESTANCIA_HABITABLE_ITEMS = [
    ("persiana", "Persiana"),
    ("cajon_persiana", "Cajón persiana"),
    ("carpinteria_estancia", "Carpintería estancia"),
    ("cierre_manivela", "Cierre / manivela"),
    ("tomas_corriente", "Tomas de corriente"),
]

INSPECCION_EXTERIOR_ITEMS = [
    ("fachadas", "Fachadas"),
    ("cubiertas_exteriores", "Cubiertas exteriores"),
    ("patios_exteriores", "Patios exteriores"),
    ("terrazas_balcones", "Terrazas / balcones"),
    ("jardines", "Jardines"),
    ("entorno_inmediato", "Entorno inmediato"),
    ("carpinterias_exteriores", "Carpinterías exteriores"),
    ("barandillas_exteriores", "Barandillas exteriores"),
    ("rejas_exteriores", "Rejas exteriores"),
    ("toldos", "Toldos"),
    ("tendederos", "Tendederos"),
]

INSPECCION_ELEMENTOS_COMUNES_ITEMS = [
    ("portal_acceso", "Portal / acceso"),
    ("vestibulo_comun", "Vestíbulo común"),
    ("pasillos_comunes", "Pasillos comunes"),
    ("escaleras", "Escaleras"),
    ("ascensor", "Ascensor"),
    ("patio_luces", "Patio de luces"),
    ("patio_ventilacion", "Patio de ventilación"),
    ("fachada_comun", "Fachada común"),
    ("cubierta_comun", "Cubierta común"),
    ("cuarto_instalaciones_comunes", "Cuarto de instalaciones comunes"),
]
ESTADO_HABITABILIDAD_LABELS = {
    "cumple": "Cumple",
    "no_cumple": "No cumple",
    "no_aplica": "No aplica",
    "no_inspeccionado": "No inspeccionado",
}
CONCLUSION_HABITABILIDAD_LABELS = {
    "apto": "Apto",
    "apto_con_deficiencias": "Apto con deficiencias",
    "no_apto": "No apto",
}
HABITABILIDAD_GENERAL_ITEMS = [
    ("ventilacion_general", "Ventilación general"),
    ("iluminacion_natural_general", "Iluminación natural general"),
    ("salubridad_general", "Salubridad general"),
    ("seguridad_uso", "Seguridad de uso"),
    ("instalaciones_basicas", "Instalaciones básicas"),
    ("accesibilidad_basica", "Accesibilidad básica"),
    ("adecuacion_uso_residencial", "Adecuación al uso residencial"),
]
HABITABILIDAD_ESTANCIA_ITEMS = [
    ("ventilacion", "Ventilación"),
    ("iluminacion", "Iluminación"),
    ("humedades_condensaciones", "Humedades / condensaciones"),
    ("salubridad", "Salubridad"),
    ("seguridad_uso_estancia", "Seguridad de uso"),
]
HABITABILIDAD_EXTERIOR_ITEMS = [
    ("patio_ventilacion", "Patio de ventilación"),
    ("fachada_humedades", "Fachada / humedades"),
    ("cubierta_filtraciones", "Cubierta / filtraciones"),
]
VALORACION_ENCARGO_ITEMS = [
    ("nombre_solicitante", "Solicitante"),
    ("nif_cif_solicitante", "NIF/CIF"),
    ("domicilio_solicitante", "Domicilio"),
    ("entidad_financiera", "Entidad financiera"),
    ("finalidad_valoracion", "Finalidad de la valoración"),
    ("finalidad_otro", "Finalidad: otro / matiz"),
    ("alcance_valoracion", "Alcance de la valoración"),
    ("fecha_valoracion", "Fecha de valoración"),
    ("finalidad_valoracion_detallada", "Finalidad detallada"),
]
VALORACION_BASE_VALOR_ITEMS = [
    ("base_valor", "Base de valor"),
    ("base_valor_otro", "Base de valor: otro"),
    ("definicion_base_valor", "Definición de la base de valor"),
]
VALORACION_DOCUMENTACION_ITEMS = [
    ("documentacion_utilizada", "Documentación utilizada"),
    ("datos_registrales", "Datos registrales"),
]
VALORACION_IDENTIFICACION_ITEMS = [
    ("identificacion_bien", "Identificación del bien"),
    ("superficie_valoracion", "Superficie de valoración"),
    ("superficie_util", "Superficie útil"),
    ("superficie_construida", "Superficie construida"),
    ("superficie_registral", "Superficie registral"),
    ("superficie_catastral", "Superficie catastral"),
    ("superficie_terraza", "Superficie de terraza"),
    ("superficie_zonas_comunes", "Superficie de zonas comunes"),
    ("superficie_total", "Superficie total"),
    ("superficie_comprobada", "Superficie comprobada"),
    ("superficie_computable", "Superficie computable"),
    ("superficie_adoptada_calculo", "Superficie adoptada para cálculo"),
    ("criterio_superficie_adoptada", "Criterio de superficie adoptada"),
]
VALORACION_SITUACION_LEGAL_ITEMS = [
    ("situacion_ocupacion", "Situación de ocupación"),
    ("situacion_urbanistica", "Situación urbanística"),
    ("servidumbres", "Servidumbres"),
    ("linderos", "Linderos"),
]
VALORACION_ENTORNO_ITEMS = [
    ("ubicacion_valoracion", "Ubicación"),
    ("descripcion_entorno", "Descripción del entorno"),
    ("grado_consolidacion", "Grado de consolidación"),
    ("antiguedad_entorno", "Antigüedad del entorno"),
    ("rasgos_urbanos", "Rasgos urbanos"),
    ("nivel_renta", "Nivel de renta"),
    ("uso_predominante", "Uso predominante"),
    ("equipamientos", "Equipamientos"),
    ("infraestructuras", "Infraestructuras"),
]
VALORACION_EDIFICIO_INMUEBLE_ITEMS = [
    ("tipo_edificio", "Tipo de edificio"),
    ("numero_portales", "Número de portales"),
    ("numero_escaleras", "Número de escaleras"),
    ("numero_ascensores", "Número de ascensores"),
    ("estado_conservacion", "Estado de conservación"),
    ("antiguedad", "Antigüedad"),
    ("calidades", "Calidades"),
    ("vistas", "Vistas"),
    ("uso_residencial", "Uso residencial"),
]
VALORACION_CONSTRUCTIVO_ITEMS = [
    ("estructura", "Estructura"),
    ("cubierta", "Cubierta"),
    ("cerramientos", "Cerramientos"),
    ("aislamiento", "Aislamiento"),
    ("carpinteria", "Carpintería"),
    ("acristalamiento", "Acristalamiento"),
    ("instalaciones", "Instalaciones"),
]
VALORACION_ESTADO_ITEMS = [
    ("estado_inmueble", "Estado actual del inmueble"),
    ("regimen_ocupacion", "Régimen de ocupación"),
    ("inmueble_arrendado", "Inmueble arrendado"),
    ("observaciones_portal", "Observaciones del portal"),
    (
        "observaciones_cuadro_contadores",
        "Observaciones del cuadro de contadores",
    ),
    ("fecha_visita", "Fecha de visita"),
    ("fecha_emision", "Fecha de emisión"),
    ("fecha_caducidad", "Fecha de caducidad"),
]
VALORACION_METODO_ITEMS = [
    ("criterios_metodo_valoracion", "Criterios / método de valoración"),
    ("testigos_comparables", "Testigos comparables"),
    ("observaciones_testigos", "Observaciones sobre testigos"),
    ("variables_mercado", "Variables de mercado"),
    ("metodo_homogeneizacion", "Método de homogeneización"),
]
VALORACION_METODOS_ECO_ITEMS = [
    ("metodo_comparacion_aplicado", "Comparación aplicada"),
    ("metodo_comparacion_descartado", "Comparación descartada"),
    ("metodo_comparacion_justificacion", "Justificación comparación"),
    ("metodo_comparacion_observaciones", "Observaciones comparación"),
    ("metodo_coste_aplicado", "Coste aplicado"),
    ("metodo_coste_descartado", "Coste descartado"),
    ("metodo_coste_justificacion", "Justificación coste"),
    ("metodo_coste_observaciones", "Observaciones coste"),
    ("metodo_actualizacion_rentas_aplicado", "Actualización de rentas aplicada"),
    ("metodo_actualizacion_rentas_descartado", "Actualización de rentas descartada"),
    ("metodo_actualizacion_rentas_justificacion", "Justificación actualización de rentas"),
    ("metodo_actualizacion_rentas_observaciones", "Observaciones actualización de rentas"),
    ("metodo_residual_aplicado", "Residual aplicado"),
    ("metodo_residual_descartado", "Residual descartado"),
    ("metodo_residual_justificacion", "Justificación residual"),
    ("metodo_residual_observaciones", "Observaciones residual"),
]
VALORACION_RESULTADO_ITEMS = [
    ("valor_unitario", "Valor unitario"),
    ("valor_resultante", "Valor resultante"),
    ("valor_tasacion_final", "Valor de tasación final"),
]
VALORACION_LIMITACIONES_ITEMS = [
    (
        "condicionantes_limitaciones_valoracion",
        "Condicionantes y limitaciones",
    ),
    ("observaciones_valoracion", "Observaciones"),
]
VALORACION_INCIDENCIAS_ITEMS = [
    ("incidencias_condicionantes_manuales", "Condicionantes manuales"),
    ("incidencias_advertencias_manuales", "Advertencias manuales"),
    ("incidencias_limitaciones_manuales", "Limitaciones manuales"),
]
VALORACION_NOTA_ECO = (
    "El presente informe no constituye tasación hipotecaria regulada salvo que "
    "expresamente se indique y se cumplan los requisitos legales aplicables. "
    "Su estructura técnica se inspira en criterios de trazabilidad, prudencia "
    "y justificación propios de estándares profesionales de valoración."
)
COMPARABLES_COLUMNAS = [
    ("direccion_testigo", "Dirección"),
    ("fuente_testigo", "Fuente"),
    ("fuente_tipo", "Tipo de fuente"),
    ("fuente_detalle", "Detalle de fuente"),
    ("fecha_testigo", "Fecha"),
    ("fecha_captura", "Fecha de captura"),
    ("precio_oferta", "Precio oferta"),
    ("precio_depurado", "Precio depurado"),
    ("superficie_tomada", "Superficie tomada"),
    ("tipo_superficie_tomada", "Tipo de superficie"),
    ("precio_unitario_inicial", "€/m² inicial"),
    ("unitario_homogeneizado", "€/m² homogeneizado"),
    ("unitario_para_resumen", "€/m² para resumen"),
    ("incluido_calculo", "Incluido en cálculo"),
    ("peso_porcentaje", "Peso"),
    ("representatividad", "Representatividad"),
    ("motivo_ponderacion", "Motivo ponderación"),
    ("motivo_exclusion", "Motivo exclusión"),
    ("valor_unitario", "Valor unitario"),
    ("superficie_construida", "Sup. constr."),
    ("superficie_util", "Sup. útil"),
    ("tipologia", "Tipología"),
    ("planta", "Planta"),
    ("dormitorios", "Dorm."),
    ("banos", "Baños"),
    ("estado_conservacion", "Estado"),
    ("antiguedad", "Antigüedad"),
    ("calidad_constructiva", "Calidad"),
    ("dato_verificado", "Dato verificado"),
    ("testigo_visitado", "Testigo visitado"),
    ("fiabilidad_dato", "Fiabilidad"),
    ("similitud_inmueble", "Similitud"),
    ("estado_mercado", "Estado mercado"),
    ("observaciones_economicas", "Observaciones económicas"),
    ("advertencias_calculo_texto", "Advertencias cálculo inicial"),
    ("visitado", "Visitado"),
    ("observaciones", "Observaciones"),
]
COMPARABLES_FORMATTERS = {
    "precio_oferta": formatear_moneda_es,
    "precio_depurado": formatear_moneda_es,
    "precio_unitario_inicial": formatear_precio_unitario_es,
    "unitario_homogeneizado": formatear_precio_unitario_es,
    "valor_unitario": formatear_precio_unitario_es,
    "valor_unitario_base": formatear_precio_unitario_es,
    "valor_unitario_ajustado": formatear_precio_unitario_es,
    "superficie_tomada": formatear_superficie_es,
    "superficie_construida": formatear_superficie_es,
    "superficie_util": formatear_superficie_es,
    "dato_verificado": formatear_booleano_es,
    "testigo_visitado": formatear_booleano_es,
    "visitado": formatear_booleano_es,
    "incluido_calculo": formatear_booleano_es,
    "peso_porcentaje": formatear_porcentaje_es,
}
VALORACION_CAMPOS_CONTEXTO = list(
    OrderedDict.fromkeys(
        [campo for campo, _ in VALORACION_ENCARGO_ITEMS]
        + [campo for campo, _ in VALORACION_BASE_VALOR_ITEMS]
        + [campo for campo, _ in VALORACION_DOCUMENTACION_ITEMS]
        + [campo for campo, _ in VALORACION_IDENTIFICACION_ITEMS]
        + [campo for campo, _ in VALORACION_SITUACION_LEGAL_ITEMS]
        + [campo for campo, _ in VALORACION_ENTORNO_ITEMS]
        + [campo for campo, _ in VALORACION_EDIFICIO_INMUEBLE_ITEMS]
        + [campo for campo, _ in VALORACION_CONSTRUCTIVO_ITEMS]
        + [campo for campo, _ in VALORACION_ESTADO_ITEMS]
        + [campo for campo, _ in VALORACION_METODO_ITEMS]
        + [campo for campo, _ in VALORACION_METODOS_ECO_ITEMS]
        + [campo for campo, _ in VALORACION_RESULTADO_ITEMS]
        + [campo for campo, _ in VALORACION_LIMITACIONES_ITEMS]
        + [campo for campo, _ in VALORACION_INCIDENCIAS_ITEMS]
        + ["incidencias_automaticas_visibles", "incidencias_manuales_visibles"]
    )
)
GRAVEDAD_CUADRANTE_LABELS = {
    "leve": "Leve",
    "media": "Media",
    "grave": "Grave",
}


def configurar_documento(doc: Document) -> None:
    seccion = doc.sections[-1]
    seccion.top_margin = Cm(2.5)
    seccion.bottom_margin = Cm(2.5)
    seccion.left_margin = Cm(2.5)
    seccion.right_margin = Cm(2.5)

    estilos = doc.styles

    estilos["Normal"].font.name = "Arial"
    estilos["Normal"].font.size = Pt(10)

    estilos["Title"].font.name = "Arial"
    estilos["Title"].font.size = Pt(22)
    estilos["Title"].font.bold = True

    estilos["Heading 1"].font.name = "Arial"
    estilos["Heading 1"].font.size = Pt(16)
    estilos["Heading 1"].font.bold = True

    estilos["Heading 2"].font.name = "Arial"
    estilos["Heading 2"].font.size = Pt(13)
    estilos["Heading 2"].font.bold = True

    estilos["Heading 3"].font.name = "Arial"
    estilos["Heading 3"].font.size = Pt(11)
    estilos["Heading 3"].font.bold = True


def add_parrafo(
    doc: Document,
    texto: str,
    bold: bool = False,
    centrado: bool = False,
    espacio_despues: int = 6,
):
    p = doc.add_paragraph()
    if centrado:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(espacio_despues)
    return p


def add_etiqueta_valor(doc: Document, etiqueta: str, valor) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)

    r1 = p.add_run(f"{etiqueta}: ")
    r1.bold = True
    r1.font.name = "Arial"
    r1.font.size = Pt(10)

    r2 = p.add_run(valor_o_guion(valor))
    r2.font.name = "Arial"
    r2.font.size = Pt(10)


def add_etiqueta_valor_si_hay(doc: Document, etiqueta: str, valor) -> None:
    if limpiar_texto(valor):
        add_etiqueta_valor(doc, etiqueta, valor)


def add_bloque_campos_si_hay(doc: Document, campos, datos, vacio: str | None = None) -> bool:
    hay_datos = False
    for campo, etiqueta in campos:
        if limpiar_texto(datos[campo]):
            add_etiqueta_valor(doc, etiqueta, datos[campo])
            hay_datos = True
    if not hay_datos and vacio:
        add_parrafo(doc, vacio)
    return hay_datos


def estado_inspeccion_legible(valor) -> str:
    estado = limpiar_texto(valor)
    if not estado:
        estado = "no_inspeccionado"
    return ESTADO_INSPECCION_LABELS.get(estado, valor_o_guion(estado))


def estado_habitabilidad_legible(valor) -> str:
    estado = limpiar_texto(valor)
    if not estado:
        estado = "no_inspeccionado"
    return ESTADO_HABITABILIDAD_LABELS.get(estado, valor_o_guion(estado))


def conclusion_habitabilidad_legible(valor) -> str:
    conclusion = limpiar_texto(valor)
    return CONCLUSION_HABITABILIDAD_LABELS.get(conclusion, valor_o_guion(conclusion))


def obtener_items_inspeccion_estancia(tipo_estancia: str):
    tipo = limpiar_texto(tipo_estancia).lower()
    items = list(INSPECCION_ESTANCIA_BASE_ITEMS)

    if tipo in {"baño", "aseo"}:
        items.extend(INSPECCION_ESTANCIA_BANO_ITEMS)
    elif tipo == "cocina":
        items.extend(INSPECCION_ESTANCIA_COCINA_ITEMS)
    elif tipo in {"salón", "salon", "dormitorio", "comedor", "recibidor", "pasillo"}:
        items.extend(INSPECCION_ESTANCIA_HABITABLE_ITEMS)

    return items


def add_estado_checklist(doc: Document, items, datos) -> None:
    for campo, etiqueta in items:
        add_etiqueta_valor(doc, etiqueta, estado_inspeccion_legible(datos[campo]))


def add_estado_checklist_habitabilidad(doc: Document, items, datos) -> None:
    for campo, etiqueta in items:
        add_etiqueta_valor(doc, etiqueta, estado_habitabilidad_legible(datos[campo]))


def recoger_incidencias_checklist(items, datos):
    incidencias = []
    for campo, etiqueta in items:
        estado = limpiar_texto(datos[campo])
        if estado in {"necesita_reparacion", "defecto_grave"}:
            incidencias.append((etiqueta, estado_inspeccion_legible(estado)))
    return incidencias


def fila_o_dict_vacio(fila, campos):
    if fila is None:
        return {campo: "" for campo in campos}
    return {campo: get_row_value(fila, campo, "") for campo in campos}


def add_titulo(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)

    run = p.add_run(texto)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(20)


def add_subtitulo(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)

    run = p.add_run(texto)
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(11)


def add_salto_pagina(doc: Document) -> None:
    doc.add_page_break()


def add_tabla_datos_expediente(doc: Document, expediente) -> None:
    tabla = doc.add_table(rows=0, cols=2)
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabla.style = "Table Grid"

    filas = [
        ("Número de expediente", expediente["numero_expediente"]),
        ("Cliente", expediente["cliente"]),
        ("Dirección", expediente["direccion"]),
        ("Código postal", expediente["codigo_postal"]),
        ("Ciudad", expediente["ciudad"]),
        ("Provincia", expediente["provincia"]),
        ("Tipo de inmueble", expediente["tipo_inmueble"]),
        ("Orientación", expediente["orientacion_inmueble"]),
        ("Año de construcción", expediente["anio_construccion"]),
        ("Uso del inmueble", expediente["uso_inmueble"]),
    ]

    for etiqueta, valor in filas:
        row = tabla.add_row().cells
        row[0].text = etiqueta
        row[1].text = valor_o_guion(valor)

    for fila in tabla.rows:
        for i, celda in enumerate(fila.cells):
            celda.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for parrafo in celda.paragraphs:
                for run in parrafo.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
                    if i == 0:
                        run.bold = True


def add_pie_figura(doc: Document, texto: str) -> None:
    """Añade un pie de figura discreto y homogéneo."""
    if not limpiar_texto(texto):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(texto)
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.italic = True


def siguiente_numero_figura(doc: Document) -> int:
    numero = getattr(doc, "_numero_figura", 1)
    setattr(doc, "_numero_figura", numero + 1)
    return numero


def add_imagen_si_existe(doc: Document, nombre_foto: str, pie: str | None = None) -> None:
    if not nombre_foto:
        return

    ruta_foto = os.path.join(UPLOAD_DIR, nombre_foto)
    if not os.path.exists(ruta_foto):
        add_parrafo(doc, f"Fotografía no localizada: {nombre_foto}")
        return

    try:
        doc.add_picture(ruta_foto, width=Cm(12.5))
        ultimo = doc.paragraphs[-1]
        ultimo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ultimo.paragraph_format.space_after = Pt(2)
        if pie:
            add_pie_figura(doc, pie)
    except Exception:
        add_parrafo(doc, f"No se ha podido insertar la fotografía: {nombre_foto}")


def add_separador_bloque(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("──────────────────────────────────────────────────────")
    run.font.name = "Arial"
    run.font.size = Pt(8)


def add_tabla_campos(doc: Document, filas) -> None:
    """Tabla compacta de campos técnicos para una patología."""
    tabla = doc.add_table(rows=0, cols=2)
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabla.style = "Table Grid"

    for etiqueta, valor in filas:
        row = tabla.add_row().cells
        row[0].text = valor_o_guion(etiqueta)
        row[1].text = valor_o_guion(valor)

    for fila in tabla.rows:
        for i, celda in enumerate(fila.cells):
            celda.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for parrafo in celda.paragraphs:
                parrafo.paragraph_format.space_after = Pt(0)
                for run in parrafo.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
                    if i == 0:
                        run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def indice_a_letras_cuadrante(indice: int) -> str:
    letras = ""
    valor = indice + 1
    while valor > 0:
        valor, resto = divmod(valor - 1, 26)
        letras = chr(65 + resto) + letras
    return letras


def generar_overlay_cuadrantes_mapa(
    ruta_imagen, filas, columnas, cuadrantes_con_incidencia
):
    from PIL import Image, ImageDraw, ImageFont

    if not ruta_imagen or not os.path.exists(ruta_imagen):
        raise FileNotFoundError("Imagen base no localizada")
    if not filas or not columnas or int(filas) <= 0 or int(columnas) <= 0:
        raise ValueError("Filas o columnas no válidas")

    imagen = Image.open(ruta_imagen).convert("RGBA")
    ancho, alto = imagen.size
    ancho_celda = ancho / int(columnas)
    alto_celda = alto / int(filas)
    draw = ImageDraw.Draw(imagen, "RGBA")
    font = ImageFont.load_default()
    incidencias = {limpiar_texto(codigo).upper() for codigo in cuadrantes_con_incidencia if limpiar_texto(codigo)}

    for fila_idx in range(int(filas)):
        for columna_idx in range(int(columnas)):
            x0 = int(round(columna_idx * ancho_celda))
            y0 = int(round(fila_idx * alto_celda))
            x1 = int(round((columna_idx + 1) * ancho_celda))
            y1 = int(round((fila_idx + 1) * alto_celda))
            codigo = f"{indice_a_letras_cuadrante(fila_idx)}{columna_idx + 1}"

            if codigo.upper() in incidencias:
                draw.rectangle(
                    [(x0, y0), (x1, y1)],
                    fill=(255, 0, 0, 28),
                    outline=(200, 0, 0, 255),
                    width=4,
                )

            draw.rectangle(
                [(x0, y0), (x1, y1)],
                outline=(0, 0, 0, 170),
                width=1,
            )

            bbox = draw.textbbox((0, 0), codigo, font=font)
            texto_ancho = bbox[2] - bbox[0]
            texto_alto = bbox[3] - bbox[1]
            texto_x = x0 + 6
            texto_y = y0 + 4
            draw.rectangle(
                [
                    (texto_x - 3, texto_y - 2),
                    (texto_x + texto_ancho + 3, texto_y + texto_alto + 2),
                ],
                fill=(255, 255, 255, 190),
            )
            draw.text((texto_x, texto_y), codigo, fill=(0, 0, 0, 255), font=font)

    buffer = BytesIO()
    imagen.save(buffer, format="PNG")
    buffer.seek(0)
    return buffer


def add_imagen_mapa_con_overlay(doc: Document, mapa) -> None:
    nombre_foto = limpiar_texto(mapa.get("imagen_base"))
    if not nombre_foto:
        return

    ruta_foto = os.path.join(UPLOAD_DIR, nombre_foto)
    cuadrantes_con_incidencia = [
        cuadrante["codigo_cuadrante"] for cuadrante in mapa.get("cuadrantes", [])
    ]

    try:
        overlay = generar_overlay_cuadrantes_mapa(
            ruta_foto,
            mapa.get("filas"),
            mapa.get("columnas"),
            cuadrantes_con_incidencia,
        )
        doc.add_picture(overlay, width=Cm(12.5))
        ultimo = doc.paragraphs[-1]
        ultimo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ultimo.paragraph_format.space_after = Pt(8)
    except Exception:
        add_imagen_si_existe(doc, nombre_foto)


def add_tabla_comparables(doc: Document, comparables) -> None:
    if not comparables:
        add_parrafo(doc, "No constan comparables registrados.")
        return

    tabla = doc.add_table(rows=1, cols=len(COMPARABLES_COLUMNAS))
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabla.style = "Table Grid"

    encabezado = tabla.rows[0].cells
    for indice, (_, titulo) in enumerate(COMPARABLES_COLUMNAS):
        encabezado[indice].text = titulo

    for comparable in comparables:
        fila = tabla.add_row().cells
        for indice, (campo, _) in enumerate(COMPARABLES_COLUMNAS):
            fila[indice].text = valor_o_guion(comparable[campo])

    for fila in tabla.rows:
        es_encabezado = fila is tabla.rows[0]
        for celda in fila.cells:
            celda.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for parrafo in celda.paragraphs:
                for run in parrafo.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8)
                    run.bold = es_encabezado


def add_valor_destacado(doc: Document, etiqueta: str, valor) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)

    r1 = p.add_run(f"{etiqueta}: ")
    r1.bold = True
    r1.font.name = "Arial"
    r1.font.size = Pt(13)

    r2 = p.add_run(valor_o_guion(valor))
    r2.bold = True
    r2.font.name = "Arial"
    r2.font.size = Pt(14)


def add_portada(doc: Document, expediente) -> None:
    add_titulo(doc, "INFORME PERICIAL")
    add_subtitulo(doc, "Inspección técnica y registro de patologías")

    for _ in range(4):
        doc.add_paragraph()

    add_parrafo(
        doc,
        f"Expediente: {valor_o_guion(expediente['numero_expediente'])}",
        bold=True,
        centrado=True,
        espacio_despues=10,
    )
    add_parrafo(
        doc,
        f"Cliente: {valor_o_guion(expediente['cliente'])}",
        centrado=True,
        espacio_despues=8,
    )
    add_parrafo(
        doc,
        f"Dirección: {valor_o_guion(expediente['direccion'])}",
        centrado=True,
        espacio_despues=8,
    )
    add_parrafo(
        doc,
        f"Fecha de emisión: {datetime.now().strftime('%d/%m/%Y')}",
        centrado=True,
        espacio_despues=8,
    )

    add_salto_pagina(doc)


def add_apartado_introduccion(doc: Document) -> None:
    doc.add_heading("1. Objeto del informe", level=1)
    add_parrafo(
        doc,
        "El presente informe tiene por objeto dejar constancia de los datos del expediente, "
        "las visitas realizadas al inmueble inspeccionado, las estancias revisadas y las "
        "patologías observadas durante la inspección técnica.",
    )


def add_apartado_datos_generales(doc: Document, expediente) -> None:
    doc.add_heading("2. Datos generales del expediente", level=1)
    add_tabla_datos_expediente(doc, expediente)
    doc.add_paragraph()

    add_etiqueta_valor(
        doc, "Observaciones generales", expediente["observaciones_generales"]
    )

    numero_unidad = "2.1"

    if expediente["observaciones_bloque"]:
        doc.add_heading("2.1 Características del bloque", level=2)
        add_etiqueta_valor(
            doc, "Observaciones del bloque", expediente["observaciones_bloque"]
        )
        numero_unidad = "2.2"

    doc.add_heading(f"{numero_unidad} Características de la unidad", level=2)
    add_etiqueta_valor(doc, "Planta de la unidad", expediente["planta_unidad"])
    add_etiqueta_valor(doc, "Puerta / unidad", expediente["puerta_unidad"])
    add_etiqueta_valor(
        doc, "Superficie construida", expediente["superficie_construida"]
    )
    add_etiqueta_valor(doc, "Superficie útil", expediente["superficie_util"])
    add_etiqueta_valor(doc, "Dormitorios", expediente["dormitorios_unidad"])
    add_etiqueta_valor(doc, "Baños", expediente["banos_unidad"])
    add_etiqueta_valor(
        doc, "Observaciones de la unidad", expediente["observaciones_unidad"]
    )


def add_apartado_bloque_judicial(doc: Document, numero_apartado: int, expediente) -> None:
    if limpiar_texto(expediente["destinatario"]) != "judicial":
        return

    doc.add_heading(f"{numero_apartado}. Encargo judicial", level=1)
    add_etiqueta_valor_si_hay(
        doc, "Procedimiento judicial", expediente["procedimiento_judicial"]
    )
    add_etiqueta_valor_si_hay(doc, "Juzgado", expediente["juzgado"])
    add_etiqueta_valor_si_hay(doc, "Auto judicial", expediente["auto_judicial"])
    add_etiqueta_valor_si_hay(
        doc, "Parte solicitante", expediente["parte_solicitante"]
    )
    add_etiqueta_valor_si_hay(doc, "Objeto de la pericia", expediente["objeto_pericia"])
    add_etiqueta_valor_si_hay(
        doc, "Alcance y limitaciones", expediente["alcance_limitaciones"]
    )
    add_etiqueta_valor_si_hay(
        doc, "Metodología pericial", expediente["metodologia_pericial"]
    )


def add_apartado_datos_patologias(doc: Document, numero_apartado: int, expediente) -> None:
    doc.add_heading(
        f"{numero_apartado}. Datos específicos del informe de patologías", level=1
    )

    ambitos = {
        "interior": "Interior",
        "exterior": "Exterior",
        "interior_exterior": "Interior y exterior",
    }
    add_etiqueta_valor(
        doc,
        "Ámbito de patologías",
        ambitos.get(limpiar_texto(expediente["ambito_patologias"]), expediente["ambito_patologias"]),
    )
    add_etiqueta_valor_si_hay(doc, "Descripción del daño", expediente["descripcion_dano"])
    add_etiqueta_valor_si_hay(doc, "Causa probable", expediente["causa_probable"])
    add_etiqueta_valor_si_hay(doc, "Pruebas e indicios", expediente["pruebas_indicios"])
    add_etiqueta_valor_si_hay(
        doc, "Evolución / preexistencia", expediente["evolucion_preexistencia"]
    )
    add_etiqueta_valor_si_hay(
        doc, "Propuesta de reparación", expediente["propuesta_reparacion"]
    )
    add_etiqueta_valor_si_hay(doc, "Urgencia / gravedad", expediente["urgencia_gravedad"])


def add_apartado_inspeccion_visita(doc: Document, numero_apartado: int, visitas, climatologias) -> None:
    doc.add_heading(f"{numero_apartado}. Inspección y visita", level=1)

    if not visitas:
        add_parrafo(doc, "No constan visitas registradas en el expediente.")
        return

    for indice, visita in enumerate(visitas, start=1):
        doc.add_heading(
            f"{numero_apartado}.{indice} Visita de fecha {valor_o_guion(visita['fecha'])}",
            level=2,
        )
        add_etiqueta_valor(doc, "Técnico", visita["tecnico"])
        add_etiqueta_valor(doc, "Observaciones de visita", visita["observaciones_visita"])
        climatologia = climatologias.get(visita["id"])
        if climatologia:
            add_etiqueta_valor(doc, "Climatología", climatologia["resumen"])
        else:
            add_etiqueta_valor(doc, "Climatología", "No consta climatología registrada")


def agrupar_patologias_interiores(patologias):
    grupos = OrderedDict()
    for patologia in patologias:
        nivel = limpiar_texto(patologia["nivel_nombre"])
        unidad = limpiar_texto(patologia["unidad_identificador"])
        estancia = limpiar_texto(patologia["estancia_nombre"]) or "Sin estancia"
        clave = (nivel, unidad, estancia)
        grupos.setdefault(
            clave,
            {
                "nivel": nivel,
                "unidad": unidad,
                "estancia": estancia,
                "items": [],
            },
        )
        grupos[clave]["items"].append(patologia)
    return grupos


def agrupar_patologias_exteriores(registros):
    grupos = OrderedDict()
    for registro in registros:
        zona = limpiar_texto(registro["zona_exterior"]) or "Sin zona exterior"
        grupos.setdefault(zona, []).append(registro)
    return grupos


def deduplicar_textos(valores) -> list[str]:
    resultado = []
    vistos = set()
    for valor in valores:
        texto = limpiar_texto(valor)
        if not texto:
            continue
        clave = texto.lower()
        if clave in vistos:
            continue
        vistos.add(clave)
        resultado.append(texto)
    return resultado


def unir_lista_natural(valores) -> str:
    items = deduplicar_textos(valores)
    if not items:
        return ""
    if len(items) == 1:
        return items[0]
    if len(items) == 2:
        return f"{items[0]} y {items[1]}"
    return f"{', '.join(items[:-1])} y {items[-1]}"


def obtener_rol_final_patologia(registro) -> str:
    rol_patologia_observado = (
        registro["rol_patologia_observado"]
        if "rol_patologia_observado" in registro.keys()
        else ""
    )
    rol_patologia_biblioteca = (
        registro["rol_patologia_biblioteca"]
        if "rol_patologia_biblioteca" in registro.keys()
        else ""
    )
    return limpiar_texto(
        rol_patologia_observado or rol_patologia_biblioteca
    ).lower()


def construir_resumen_patologias(prefijo: str, registros) -> str:
    causas = []
    efectos = []
    mixtas = []
    observaciones = []

    for registro in registros:
        rol = obtener_rol_final_patologia(registro)
        nombre_patologia = limpiar_texto(registro["patologia"])
        if rol == "causa":
            causas.append(nombre_patologia)
        elif rol == "efecto":
            efectos.append(nombre_patologia)
        elif rol == "mixta":
            mixtas.append(nombre_patologia)
        obs = limpiar_texto(registro["observaciones"])
        if obs:
            observaciones.append(obs)

    causas = deduplicar_textos(causas)
    efectos = deduplicar_textos(efectos)
    mixtas = deduplicar_textos(mixtas)
    observaciones = deduplicar_textos(observaciones)

    frases = []
    efectos_principales = list(efectos)
    if causas:
        efectos_principales = deduplicar_textos(efectos_principales + mixtas)
    descripcion_causas = (
        f"se observan procesos patológicos consistentes en {unir_lista_natural(causas)}"
        if len(causas) > 1
        else f"se observa {unir_lista_natural(causas)}"
    )

    if causas and efectos_principales:
        frases.append(
            f"{prefijo} {descripcion_causas}, "
            f"con efectos derivados consistentes en {unir_lista_natural(efectos_principales)}."
        )
    elif causas:
        frases.append(f"{prefijo} {descripcion_causas}.")
        if mixtas:
            frases.append(f"Asimismo, se aprecian {unir_lista_natural(mixtas)}.")
    elif efectos:
        frases.append(
            f"{prefijo} se observan daños consistentes en {unir_lista_natural(efectos)}, "
            "sin poder determinar con certeza su origen."
        )
        if mixtas:
            frases.append(f"Asimismo, se aprecian {unir_lista_natural(mixtas)}.")
    elif mixtas:
        frases.append(f"{prefijo} se observa {unir_lista_natural(mixtas)}.")

    if observaciones:
        frases.append(f"Asimismo, se aprecian {unir_lista_natural(observaciones)}.")

    return " ".join(frases)


def construir_titulo_grupo_interior(grupo) -> str:
    partes = []
    if grupo["nivel"]:
        partes.append(f"Nivel: {grupo['nivel']}")
    if grupo["unidad"]:
        partes.append(f"Unidad: {grupo['unidad']}")
    partes.append(f"Estancia: {grupo['estancia']}")
    return " · ".join(partes)


def detectar_incoherencias(registros) -> list[str]:
    causas = 0
    efectos = 0
    mixtas_sin_definir = 0

    for registro in registros:
        rol_observado = limpiar_texto(
            registro["rol_patologia_observado"]
            if "rol_patologia_observado" in registro.keys()
            else ""
        ).lower()
        rol_final = obtener_rol_final_patologia(registro)
        if rol_final == "causa":
            causas += 1
        elif rol_final == "efecto":
            efectos += 1
        elif rol_final == "mixta" and not rol_observado:
            mixtas_sin_definir += 1

    incoherencias = []
    if efectos and not causas:
        incoherencias.append(
            "Se han identificado efectos sin una causa claramente asociada. Se recomienda revisar el origen del daño."
        )
    if causas and not efectos:
        incoherencias.append(
            "Se han identificado causas sin manifestaciones visibles asociadas en el momento de la inspección."
        )
    if mixtas_sin_definir:
        incoherencias.append(
            "Existen patologías de carácter mixto cuya función no ha sido especificada (causa o efecto)."
        )
    return incoherencias


def construir_conclusion_tecnica_global(patologias_interiores, patologias_exteriores) -> str:
    registros = list(patologias_interiores) + list(patologias_exteriores)
    if not registros:
        return "No constan patologías registradas en el expediente."

    causas = []
    efectos = []
    mixtas = []
    incoherencias = []

    for registro in registros:
        rol = obtener_rol_final_patologia(registro)
        nombre_patologia = limpiar_texto(registro["patologia"])
        if rol == "causa":
            causas.append(nombre_patologia)
        elif rol == "efecto":
            efectos.append(nombre_patologia)
        elif rol == "mixta":
            mixtas.append(nombre_patologia)

    causas = deduplicar_textos(causas)
    efectos = deduplicar_textos(efectos)
    mixtas = deduplicar_textos(mixtas)
    incoherencias = deduplicar_textos(detectar_incoherencias(registros))

    afecta_interior = bool(patologias_interiores)
    afecta_exterior = bool(patologias_exteriores)
    if afecta_interior and afecta_exterior:
        alcance = "interior y exterior"
    elif afecta_interior:
        alcance = "interior"
    elif afecta_exterior:
        alcance = "exterior"
    else:
        alcance = ""

    frases = []
    causas_principales = deduplicar_textos(causas + mixtas)

    if causas_principales:
        frases.append(
            f"Del análisis realizado se concluye que existen procesos patológicos asociados a {unir_lista_natural(causas_principales)}."
        )
        if efectos:
            frases.append(
                f"Estos han generado efectos consistentes en {unir_lista_natural(efectos)}, afectando a {alcance}."
            )
        else:
            frases.append("Se han identificado causas sin manifestaciones visibles claras.")
    elif efectos:
        frases.append("Se han identificado daños sin una causa claramente definida.")

    if incoherencias:
        frases.append(
            "Asimismo, se recomienda revisar determinadas situaciones detectadas durante el análisis."
        )

    return " ".join(frases)


def construir_conclusion_pericial(patologias_interiores, patologias_exteriores) -> str:
    registros = list(patologias_interiores) + list(patologias_exteriores)
    if not registros:
        return "No se observan daños significativos que permitan establecer una conclusión pericial."

    incoherencias = detectar_incoherencias(registros)
    causas = []
    efectos = []

    for registro in registros:
        rol = obtener_rol_final_patologia(registro)
        nombre_patologia = limpiar_texto(registro["patologia"])
        if rol == "causa":
            causas.append(nombre_patologia)
        elif rol == "efecto":
            efectos.append(nombre_patologia)

    causas = deduplicar_textos(causas)
    causa_dominante = causas[0] if causas else ""
    efectos = deduplicar_textos(efectos)
    hay_causas = bool(causas)
    hay_efectos = bool(efectos)
    hay_incoherencias = bool(incoherencias)

    frases = []
    if hay_causas and hay_efectos and not hay_incoherencias:
        frases.append(
            f"Del análisis realizado se desprende que los daños observados tienen su origen en procesos asociados a {causa_dominante}."
        )
        frases.append(
            f"Las manifestaciones detectadas, consistentes en {unir_lista_natural(efectos)}, presentan coherencia con dicho origen."
        )
        frases.append(
            "La distribución de las lesiones en el inmueble resulta compatible con el mecanismo descrito."
        )
    elif hay_incoherencias:
        frases.append(
            "No puede determinarse con certeza el origen de los daños observados a partir de la información disponible."
        )
    elif hay_efectos:
        frases.append(
            f"Los daños observados resultan compatibles con procesos asociados a {causa_dominante or 'diversos mecanismos patológicos'}."
        )
    elif hay_causas:
        frases.append(
            "Se identifican procesos patológicos sin manifestaciones visibles suficientemente definidas en el momento de la inspección."
        )
    else:
        frases.append(
            "No se observan daños significativos que permitan establecer una conclusión pericial."
        )

    if patologias_interiores and patologias_exteriores:
        frases.append(
            "Los daños afectan tanto a zonas interiores como exteriores del inmueble."
        )

    return " ".join(frases)


def add_apartado_patologias_interiores(
    doc: Document, numero_apartado: int, patologias, conn
) -> None:
    doc.add_heading(f"{numero_apartado}. Patologías interiores", level=1)

    if not patologias:
        add_parrafo(doc, "No constan patologías interiores registradas.")
        return

    grupos = agrupar_patologias_interiores(patologias)
    indice_estancia = 1
    for grupo in grupos.values():
        if indice_estancia > 1:
            add_separador_bloque(doc)

        doc.add_heading(
            f"{numero_apartado}.{indice_estancia} {construir_titulo_grupo_interior(grupo)}",
            level=2,
        )
        resumen = construir_resumen_patologias(
            f"En la estancia {grupo['estancia']}",
            grupo["items"],
        )
        if resumen:
            add_parrafo(doc, resumen)
        incoherencias = detectar_incoherencias(grupo["items"])
        if incoherencias:
            add_parrafo(doc, "Observaciones técnicas:", bold=True)
            for item in incoherencias:
                add_parrafo(doc, f"- {item}")

        for indice_item, item in enumerate(grupo["items"], start=1):
            doc.add_heading(
                f"Patología {indice_item} · {valor_o_guion(item['patologia'])}",
                level=3,
            )
            referencias_cuadrantes = obtener_referencias_cuadrantes_patologia(
                conn, item["id"]
            )
            filas = []
            if referencias_cuadrantes:
                filas.append((
                    "Cuadrantes",
                    formatear_referencias_cuadrantes_patologia(referencias_cuadrantes),
                ))
            filas.extend([
                ("Localización del daño", item["localizacion_dano"]),
                ("Elemento", item["elemento"]),
                ("Patología", item["patologia"]),
                ("Observaciones", item["observaciones"]),
            ])
            add_tabla_campos(doc, filas)
            if item["foto"]:
                numero_figura = siguiente_numero_figura(doc)
                pie = (
                    f"Figura {numero_figura}. {valor_o_guion(item['patologia'])} "
                    f"en {valor_o_guion(item['localizacion_dano']).lower()} "
                    f"({valor_o_guion(grupo['estancia'])})"
                )
                add_imagen_si_existe(doc, item["foto"], pie=pie)
        indice_estancia += 1


def add_apartado_patologias_exteriores(
    doc: Document, numero_apartado: int, registros, conn
) -> None:
    doc.add_heading(f"{numero_apartado}. Patologías exteriores", level=1)

    if not registros:
        add_parrafo(doc, "No constan patologías exteriores registradas.")
        return

    zonas = {
        "fachada": "Fachada",
        "cubierta": "Cubierta",
        "medianera": "Medianera",
        "patio": "Patio",
        "terraza": "Terraza",
        "exterior_general": "Exterior general",
    }
    elementos = {
        "revestimiento": "Revestimiento",
        "cerramiento": "Cerramiento",
        "cornisa": "Cornisa",
        "alero": "Alero",
        "peto": "Peto",
        "impermeabilizacion": "Impermeabilizacion",
        "carpinteria_exterior": "Carpinteria exterior",
        "barandilla": "Barandilla",
        "bajante": "Bajante",
        "canalon": "Canalon",
        "forjado": "Forjado",
        "otro": "Otro",
    }
    localizaciones = {
        "horizontal": "Horizontal",
        "vertical": "Vertical",
        "encuentro": "Encuentro",
        "puntual": "Puntual",
    }

    grupos = agrupar_patologias_exteriores(registros)
    indice_zona = 1
    for zona, items in grupos.items():
        if indice_zona > 1:
            add_separador_bloque(doc)

        zona_legible = zonas.get(zona, zona)
        doc.add_heading(
            f"{numero_apartado}.{indice_zona} {zona_legible}",
            level=2,
        )
        resumen = construir_resumen_patologias(
            "En el exterior del inmueble",
            items,
        )
        if resumen:
            add_parrafo(doc, resumen)
        incoherencias = detectar_incoherencias(items)
        if incoherencias:
            add_parrafo(doc, "Observaciones técnicas:", bold=True)
            for item in incoherencias:
                add_parrafo(doc, f"- {item}")
        for indice_item, item in enumerate(items, start=1):
            doc.add_heading(
                f"Patología {indice_item} · {valor_o_guion(item['patologia'])}",
                level=3,
            )
            referencias_cuadrantes = obtener_referencias_cuadrantes_patologia(
                conn, item["id"]
            )
            elemento_exterior = elementos.get(
                limpiar_texto(item["elemento_exterior"]), item["elemento_exterior"]
            )
            localizacion_exterior = localizaciones.get(
                limpiar_texto(item["localizacion_dano_exterior"]),
                item["localizacion_dano_exterior"],
            )
            filas = []
            if referencias_cuadrantes:
                filas.append((
                    "Cuadrantes",
                    formatear_referencias_cuadrantes_patologia(referencias_cuadrantes),
                ))
            filas.extend([
                ("Elemento exterior", elemento_exterior),
                ("Localización del daño exterior", localizacion_exterior),
                ("Patología", item["patologia"]),
                ("Observaciones", item["observaciones"]),
            ])
            add_tabla_campos(doc, filas)
            if item["foto"]:
                numero_figura = siguiente_numero_figura(doc)
                pie = (
                    f"Figura {numero_figura}. {valor_o_guion(item['patologia'])} "
                    f"en {valor_o_guion(localizacion_exterior).lower()} "
                    f"({valor_o_guion(zona_legible)})"
                )
                add_imagen_si_existe(doc, item["foto"], pie=pie)
        indice_zona += 1


def describir_objeto_visita_informe(cur, visita) -> str:
    ambito = limpiar_texto(visita["ambito_visita"]) or "edificio_completo"
    if ambito == "edificio_completo":
        return "Edificio completo"
    if ambito == "nivel" and visita["nivel_id"]:
        nivel = cur.execute(
            "SELECT nombre_nivel FROM niveles_edificio WHERE id=?",
            (visita["nivel_id"],),
        ).fetchone()
        if nivel and limpiar_texto(nivel["nombre_nivel"]):
            return f"Nivel: {nivel['nombre_nivel']}"
    if visita["unidad_id"]:
        unidad = cur.execute(
            "SELECT identificador FROM unidades_expediente WHERE id=?",
            (visita["unidad_id"],),
        ).fetchone()
        identificador = limpiar_texto(unidad["identificador"]) if unidad else ""
        if identificador:
            if ambito == "zona_comun":
                return f"Zona común: {identificador}"
            if ambito == "exterior":
                return f"Exterior: {identificador}"
            return f"Unidad: {identificador}"
    etiquetas = {
        "nivel": "Nivel",
        "unidad": "Unidad",
        "zona_comun": "Zona común",
        "exterior": "Exterior",
    }
    return etiquetas.get(ambito, valor_o_guion(ambito))


def resumen_patologia_vinculada_mapa(cur, visita_id: int, patologia_id, patologia_detectada) -> str:
    if patologia_id:
        interior = cur.execute(
            """
            SELECT rp.patologia, e.nombre AS estancia_nombre
            FROM registros_patologias rp
            JOIN estancias e ON rp.estancia_id = e.id
            WHERE rp.id=? AND rp.visita_id=?
            """,
            (patologia_id, visita_id),
        ).fetchone()
        exterior = cur.execute(
            """
            SELECT patologia, zona_exterior
            FROM registros_patologias_exteriores
            WHERE id=? AND visita_id=?
            """,
            (patologia_id, visita_id),
        ).fetchone()

        texto_referencia = limpiar_texto(patologia_detectada).lower()
        if interior and exterior and texto_referencia:
            if limpiar_texto(interior["patologia"]).lower() == texto_referencia:
                exterior = None
            elif limpiar_texto(exterior["patologia"]).lower() == texto_referencia:
                interior = None

        if interior:
            estancia = limpiar_texto(interior["estancia_nombre"])
            if estancia:
                return f"{interior['patologia']} ({estancia})"
            return valor_o_guion(interior["patologia"])
        if exterior:
            zona = limpiar_texto(exterior["zona_exterior"])
            if zona:
                return f"{exterior['patologia']} ({zona})"
            return valor_o_guion(exterior["patologia"])

    return limpiar_texto(patologia_detectada)


def obtener_referencias_cuadrantes_patologia(conn, patologia_id):
    if not patologia_id:
        return []

    filas = conn.execute(
        """
        SELECT mp.titulo AS mapa_titulo, qmp.codigo_cuadrante
        FROM cuadrantes_mapa_patologia qmp
        JOIN mapas_patologia mp ON qmp.mapa_id = mp.id
        WHERE qmp.patologia_id = ?
        ORDER BY mp.id ASC, qmp.id ASC
        """,
        (patologia_id,),
    ).fetchall()

    if not filas:
        return []

    agrupado = OrderedDict()
    for fila in filas:
        mapa_titulo = limpiar_texto(fila["mapa_titulo"]) or "Mapa sin título"
        agrupado.setdefault(mapa_titulo, [])
        codigo = limpiar_texto(fila["codigo_cuadrante"])
        if codigo and codigo not in agrupado[mapa_titulo]:
            agrupado[mapa_titulo].append(codigo)

    return [
        {
            "mapa_titulo": "" if mapa_titulo == "Mapa sin título" else mapa_titulo,
            "cuadrantes": sorted(cuadrantes, key=clave_orden_cuadrante),
        }
        for mapa_titulo, cuadrantes in agrupado.items()
    ]


def formatear_referencias_cuadrantes_patologia(referencias) -> str:
    if not referencias:
        return ""

    if len(referencias) == 1:
        referencia = referencias[0]
        cuadrantes = ", ".join(referencia["cuadrantes"])
        mapa_titulo = limpiar_texto(referencia["mapa_titulo"])
        if mapa_titulo:
            return f"{cuadrantes} (Mapa: {mapa_titulo})"
        return cuadrantes

    partes = []
    for referencia in referencias:
        cuadrantes = ", ".join(referencia["cuadrantes"])
        mapa_titulo = limpiar_texto(referencia["mapa_titulo"])
        if mapa_titulo:
            partes.append(f"{cuadrantes} ({mapa_titulo})")
        else:
            partes.append(cuadrantes)
    return " | ".join(partes)


def cargar_mapas_patologia_utiles_visita(cur, visita):
    mapas = cur.execute(
        """
        SELECT *
        FROM mapas_patologia
        WHERE visita_id=?
        ORDER BY id ASC
        """,
        (visita["id"],),
    ).fetchall()

    resultado = []
    for mapa in mapas:
        cuadrantes = cur.execute(
            """
            SELECT *
            FROM cuadrantes_mapa_patologia
            WHERE mapa_id=?
            ORDER BY id ASC
            """,
            (mapa["id"],),
        ).fetchall()
        cuadrantes_utiles = []
        for cuadrante in cuadrantes:
            tiene_datos = any(
                [
                    limpiar_texto(cuadrante["descripcion"]),
                    limpiar_texto(cuadrante["gravedad"]),
                    cuadrante["patologia_id"],
                    limpiar_texto(cuadrante["foto_detalle"]),
                ]
            )
            if not tiene_datos:
                continue
            cuadrantes_utiles.append(
                {
                    "codigo_cuadrante": cuadrante["codigo_cuadrante"],
                    "descripcion": cuadrante["descripcion"],
                    "gravedad_label": GRAVEDAD_CUADRANTE_LABELS.get(
                        limpiar_texto(cuadrante["gravedad"]), valor_o_guion(cuadrante["gravedad"])
                    ),
                    "patologia_vinculada_resumen": resumen_patologia_vinculada_mapa(
                        cur,
                        visita["id"],
                        cuadrante["patologia_id"],
                        cuadrante["patologia_detectada"],
                    ),
                    "foto_detalle": cuadrante["foto_detalle"],
                }
            )

        if not cuadrantes_utiles:
            continue

        resultado.append(
            {
                "titulo": mapa["titulo"],
                "descripcion": mapa["descripcion"],
                "imagen_base": mapa["imagen_base"],
                "filas": mapa["filas"],
                "columnas": mapa["columnas"],
                "objeto_visita_label": describir_objeto_visita_informe(cur, visita),
                "cuadrantes": cuadrantes_utiles,
            }
        )

    return resultado


def add_apartado_mapas_patologia(
    doc: Document, numero_apartado: int, bloques_mapas
) -> bool:
    bloques_utiles = [bloque for bloque in bloques_mapas if bloque["mapas"]]
    if not bloques_utiles:
        return False

    doc.add_heading(f"{numero_apartado}. Mapas de localización de daños", level=1)

    multiple_visitas = len(bloques_utiles) > 1
    for indice_visita, bloque in enumerate(bloques_utiles, start=1):
        visita = bloque["visita"]
        mapas = bloque["mapas"]
        prefijo_visita = f"{numero_apartado}.{indice_visita}"

        if multiple_visitas:
            doc.add_heading(
                f"{prefijo_visita} Visita de fecha {valor_o_guion(visita['fecha'])}",
                level=2,
            )

        for indice_mapa, mapa in enumerate(mapas, start=1):
            prefijo_mapa = (
                f"{prefijo_visita}.{indice_mapa}" if multiple_visitas else f"{numero_apartado}.{indice_mapa}"
            )
            doc.add_heading(f"{prefijo_mapa} {valor_o_guion(mapa['titulo'])}", level=2 if not multiple_visitas else 3)
            add_etiqueta_valor_si_hay(doc, "Descripción", mapa["descripcion"])
            add_etiqueta_valor_si_hay(doc, "Contexto", mapa["objeto_visita_label"])
            if mapa["imagen_base"]:
                add_parrafo(doc, "Imagen base con cuadrícula:", bold=True)
                add_imagen_mapa_con_overlay(doc, mapa)

            for indice_cuadrante, cuadrante in enumerate(mapa["cuadrantes"], start=1):
                prefijo_cuadrante = f"{prefijo_mapa}.{indice_cuadrante}"
                doc.add_heading(
                    f"{prefijo_cuadrante} Cuadrante {valor_o_guion(cuadrante['codigo_cuadrante'])}",
                    level=3,
                )
                add_etiqueta_valor_si_hay(doc, "Descripción", cuadrante["descripcion"])
                if limpiar_texto(cuadrante["gravedad_label"]) and cuadrante["gravedad_label"] != "-":
                    add_etiqueta_valor(doc, "Gravedad", cuadrante["gravedad_label"])
                add_etiqueta_valor_si_hay(
                    doc,
                    "Patología vinculada",
                    cuadrante["patologia_vinculada_resumen"],
                )
                if cuadrante["foto_detalle"]:
                    add_parrafo(doc, "Foto detalle:", bold=True)
                    add_imagen_si_existe(doc, cuadrante["foto_detalle"])

    return True


def add_apartado_analisis_tecnico(doc: Document, numero_apartado: int, expediente) -> None:
    doc.add_heading(f"{numero_apartado}. Análisis técnico", level=1)
    campos = [
        ("Descripción del daño", expediente["descripcion_dano"]),
        ("Causa probable", expediente["causa_probable"]),
        ("Pruebas e indicios", expediente["pruebas_indicios"]),
        ("Evolución / preexistencia", expediente["evolucion_preexistencia"]),
        ("Urgencia / gravedad", expediente["urgencia_gravedad"]),
    ]
    hay_datos = False
    for etiqueta, valor in campos:
        if limpiar_texto(valor):
            add_etiqueta_valor(doc, etiqueta, valor)
            hay_datos = True
    if not hay_datos:
        add_parrafo(doc, "No constan observaciones globales adicionales de análisis técnico.")


def add_apartado_propuesta_reparacion(doc: Document, numero_apartado: int, expediente) -> None:
    doc.add_heading(f"{numero_apartado}. Propuesta de reparación", level=1)
    if limpiar_texto(expediente["propuesta_reparacion"]):
        add_parrafo(doc, expediente["propuesta_reparacion"])
    else:
        add_parrafo(doc, "No consta propuesta de reparación registrada.")


def add_apartado_conclusion_patologias(
    doc: Document, numero_apartado: int, patologias_interiores, patologias_exteriores
) -> None:
    doc.add_heading(f"{numero_apartado}. Conclusiones técnicas", level=1)
    add_parrafo(
        doc,
        construir_conclusion_tecnica_global(
            patologias_interiores,
            patologias_exteriores,
        ),
    )


def add_apartado_conclusion_pericial(
    doc: Document, numero_apartado: int, patologias_interiores, patologias_exteriores
) -> None:
    doc.add_heading(f"{numero_apartado}. Conclusiones periciales", level=1)
    add_parrafo(
        doc,
        construir_conclusion_pericial(
            patologias_interiores,
            patologias_exteriores,
        ),
    )


def cargar_habitabilidad_visita(cur, visita_id: int) -> dict:
    general_campos = [campo for campo, _ in HABITABILIDAD_GENERAL_ITEMS] + [
        "conclusion_habitabilidad",
        "observaciones_generales_habitabilidad",
    ]
    exterior_campos = [campo for campo, _ in HABITABILIDAD_EXTERIOR_ITEMS] + [
        "observaciones_exterior_habitabilidad"
    ]

    general = fila_o_dict_vacio(
        cur.execute(
            """
            SELECT *
            FROM habitabilidad_general_visita
            WHERE visita_id = ?
            """,
            (visita_id,),
        ).fetchone(),
        general_campos,
    )
    exterior = fila_o_dict_vacio(
        cur.execute(
            """
            SELECT *
            FROM habitabilidad_exterior
            WHERE visita_id = ?
            """,
            (visita_id,),
        ).fetchone(),
        exterior_campos,
    )
    estancias = cur.execute(
        """
        SELECT e.id AS estancia_id,
               e.nombre AS estancia_nombre,
               e.tipo_estancia,
               e.planta,
               e.acabado_pavimento,
               e.acabado_paramento,
               e.acabado_techo,
               he.*
        FROM estancias e
        LEFT JOIN habitabilidad_estancias he
               ON he.estancia_id = e.id
              AND he.visita_id = e.visita_id
        WHERE e.visita_id = ?
        ORDER BY e.id ASC
        """,
        (visita_id,),
    ).fetchall()

    return {
        "general": general,
        "exterior": exterior,
        "estancias": estancias,
    }


def add_apartado_condiciones_habitabilidad(
    doc: Document, numero_apartado: int, visitas_habitabilidad
) -> None:
    doc.add_heading(f"{numero_apartado}. Condiciones generales de habitabilidad", level=1)

    if not visitas_habitabilidad:
        add_parrafo(doc, "No constan comprobaciones generales de habitabilidad registradas.")
        return

    for indice, bloque in enumerate(visitas_habitabilidad, start=1):
        visita = bloque["visita"]
        general = bloque["habitabilidad"]["general"]
        climatologia = bloque["climatologia"]
        doc.add_heading(
            f"{numero_apartado}.{indice} Visita de fecha {valor_o_guion(visita['fecha'])}",
            level=2,
        )
        add_etiqueta_valor(doc, "Técnico", visita["tecnico"])
        if climatologia:
            add_etiqueta_valor(doc, "Climatología", climatologia["resumen"])
        for campo, etiqueta in HABITABILIDAD_GENERAL_ITEMS:
            add_etiqueta_valor(doc, etiqueta, estado_habitabilidad_legible(general[campo]))
        if limpiar_texto(general["observaciones_generales_habitabilidad"]):
            add_etiqueta_valor(
                doc,
                "Observaciones generales de habitabilidad",
                general["observaciones_generales_habitabilidad"],
            )


def add_apartado_evaluacion_estancias_habitabilidad(
    doc: Document, numero_apartado: int, visitas_habitabilidad
) -> None:
    doc.add_heading(f"{numero_apartado}. Evaluación por estancias", level=1)

    if not visitas_habitabilidad:
        add_parrafo(doc, "No constan estancias evaluadas.")
        return

    for indice_visita, bloque in enumerate(visitas_habitabilidad, start=1):
        visita = bloque["visita"]
        estancias = bloque["habitabilidad"]["estancias"]
        doc.add_heading(
            f"{numero_apartado}.{indice_visita} Visita de fecha {valor_o_guion(visita['fecha'])}",
            level=2,
        )
        if not estancias:
            add_parrafo(doc, "No constan estancias registradas en esta visita.")
            continue

        for indice_estancia, estancia in enumerate(estancias, start=1):
            doc.add_heading(
                f"{numero_apartado}.{indice_visita}.{indice_estancia} {valor_o_guion(estancia['estancia_nombre'])}",
                level=3,
            )
            add_etiqueta_valor(doc, "Tipo de estancia", estancia["tipo_estancia"])
            add_etiqueta_valor(doc, "Planta", estancia["planta"])
            add_etiqueta_valor(doc, "Acabado de pavimento", estancia["acabado_pavimento"])
            add_etiqueta_valor(doc, "Acabado de paramento", estancia["acabado_paramento"])
            add_etiqueta_valor(doc, "Acabado de techo", estancia["acabado_techo"])
            add_estado_checklist_habitabilidad(doc, HABITABILIDAD_ESTANCIA_ITEMS, estancia)
            if limpiar_texto(estancia["observaciones_estancia_habitabilidad"]):
                add_etiqueta_valor(
                    doc,
                    "Observaciones de habitabilidad",
                    estancia["observaciones_estancia_habitabilidad"],
                )


def add_apartado_exterior_habitabilidad(
    doc: Document, numero_apartado: int, visitas_habitabilidad
) -> None:
    doc.add_heading(
        f"{numero_apartado}. Factores exteriores que afectan a la habitabilidad",
        level=1,
    )

    hay_datos = False
    for indice, bloque in enumerate(visitas_habitabilidad, start=1):
        exterior = bloque["habitabilidad"]["exterior"]
        if not any(limpiar_texto(exterior[campo]) for campo, _ in HABITABILIDAD_EXTERIOR_ITEMS) and not limpiar_texto(
            exterior["observaciones_exterior_habitabilidad"]
        ):
            continue

        hay_datos = True
        visita = bloque["visita"]
        doc.add_heading(
            f"{numero_apartado}.{indice} Visita de fecha {valor_o_guion(visita['fecha'])}",
            level=2,
        )
        add_estado_checklist_habitabilidad(doc, HABITABILIDAD_EXTERIOR_ITEMS, exterior)
        if limpiar_texto(exterior["observaciones_exterior_habitabilidad"]):
            add_etiqueta_valor(
                doc,
                "Observaciones exteriores",
                exterior["observaciones_exterior_habitabilidad"],
            )

    if not hay_datos:
        add_parrafo(doc, "No constan factores exteriores registrados que afecten a la habitabilidad.")


def add_apartado_conclusion_habitabilidad(
    doc: Document, numero_apartado: int, visitas_habitabilidad
) -> None:
    doc.add_heading(f"{numero_apartado}. Conclusión global de habitabilidad", level=1)

    if not visitas_habitabilidad:
        add_parrafo(doc, "No consta conclusión global de habitabilidad registrada.")
        return

    for indice, bloque in enumerate(visitas_habitabilidad, start=1):
        visita = bloque["visita"]
        general = bloque["habitabilidad"]["general"]
        doc.add_heading(
            f"{numero_apartado}.{indice} Visita de fecha {valor_o_guion(visita['fecha'])}",
            level=2,
        )
        add_etiqueta_valor(
            doc,
            "Conclusión global",
            conclusion_habitabilidad_legible(general["conclusion_habitabilidad"]),
        )
        if limpiar_texto(general["observaciones_generales_habitabilidad"]):
            add_etiqueta_valor(
                doc,
                "Observaciones complementarias",
                general["observaciones_generales_habitabilidad"],
            )


def cargar_valoracion_visita(cur, visita_id: int) -> dict:
    return fila_o_dict_vacio(
        cur.execute(
            """
            SELECT *
            FROM valoracion_visita
            WHERE visita_id = ?
            """,
            (visita_id,),
        ).fetchone(),
        VALORACION_CAMPOS_CONTEXTO,
    )


VALORACION_GRUPOS_CONTEXTO = [
    ("encargo", "Encargo / solicitante", VALORACION_ENCARGO_ITEMS),
    ("base_valor", "Base de valor", VALORACION_BASE_VALOR_ITEMS),
    ("documentacion", "Documentación utilizada", VALORACION_DOCUMENTACION_ITEMS),
    ("identificacion", "Identificación y superficies", VALORACION_IDENTIFICACION_ITEMS),
    ("situacion_legal", "Situación legal y urbanística", VALORACION_SITUACION_LEGAL_ITEMS),
    ("entorno", "Entorno", VALORACION_ENTORNO_ITEMS),
    ("edificio_inmueble", "Edificio e inmueble", VALORACION_EDIFICIO_INMUEBLE_ITEMS),
    ("constructivo", "Características constructivas", VALORACION_CONSTRUCTIVO_ITEMS),
    ("estado", "Estado actual y ocupación", VALORACION_ESTADO_ITEMS),
    ("metodo", "Método / mercado", VALORACION_METODO_ITEMS),
    ("metodos_eco", "Métodos aplicados y descartados", VALORACION_METODOS_ECO_ITEMS),
    ("resultado", "Resultado de la valoración", VALORACION_RESULTADO_ITEMS),
    ("limitaciones", "Condicionantes y limitaciones", VALORACION_LIMITACIONES_ITEMS),
    ("incidencias", "Incidencias", VALORACION_INCIDENCIAS_ITEMS),
]


class ValoracionContext(list):
    def __init__(self, bloques=None, eco=None):
        super().__init__(bloques or [])
        self.eco = eco or {}

    def __getitem__(self, key):
        if isinstance(key, str):
            return self.eco.get(key)
        return super().__getitem__(key)

    def get(self, key, default=None):
        if isinstance(key, str):
            return self.eco.get(key, default)
        return default


def construir_grupo_valoracion(clave: str, titulo: str, campos, datos: dict) -> dict:
    return {
        "clave": clave,
        "titulo": titulo,
        "campos": construir_campos_informe(
            [(etiqueta, datos.get(campo, "")) for campo, etiqueta in campos]
        ),
        "hay_datos": any(limpiar_texto(datos.get(campo, "")) for campo, _ in campos),
    }


def _tiene_datos_valoracion(datos: dict) -> bool:
    return any(limpiar_texto(datos.get(campo, "")) for campo in VALORACION_CAMPOS_CONTEXTO)


def _valoracion_dict_vacio() -> dict:
    return {campo: "" for campo in VALORACION_CAMPOS_CONTEXTO}


def _valoracion_bool(valor) -> bool:
    return limpiar_texto(valor).lower() in {"1", "true", "si", "sí", "yes", "on"}


def _valoracion_visible(valor) -> bool:
    texto = limpiar_texto(valor)
    return texto == "" or texto.lower() not in {"0", "false", "no", "off"}


def _etiqueta_base_valor(valor: str) -> str:
    etiquetas = {
        "valor_mercado": "Valor de mercado",
        "valor_razonable_estimado": "Valor razonable estimado",
        "valor_reposicion": "Valor de reposición",
        "valor_actualizacion_rentas": "Valor por actualización de rentas",
        "otro": "Otro",
    }
    texto = limpiar_texto(valor)
    return etiquetas.get(texto, texto)


def _campos_eco(campos, datos: dict) -> list[dict]:
    return construir_campos_informe(
        [
            (
                etiqueta,
                "Sí"
                if campo.endswith(("_aplicado", "_descartado"))
                and _valoracion_bool(datos.get(campo))
                else "No"
                if campo.endswith(("_aplicado", "_descartado"))
                else _etiqueta_base_valor(datos.get(campo))
                if campo == "base_valor"
                else datos.get(campo, ""),
            )
            for campo, etiqueta in campos
        ]
    )


def _grupo_eco(clave: str, titulo: str, campos, datos: dict) -> dict:
    return {
        "clave": clave,
        "titulo": titulo,
        "campos": _campos_eco(campos, datos),
        "hay_datos": any(
            _valoracion_bool(datos.get(campo))
            if campo.endswith(("_aplicado", "_descartado"))
            else limpiar_texto(datos.get(campo))
            for campo, _ in campos
        ),
    }


def _lineas_manuales(texto: str) -> list[str]:
    return [
        limpiar_texto(linea).lstrip("-*• ").strip()
        for linea in limpiar_texto(texto).splitlines()
        if limpiar_texto(linea).lstrip("-*• ").strip()
    ]


def _visita_por_id(visitas, visita_id):
    for visita in visitas or []:
        if visita["id"] == visita_id:
            return visita
    return None


def _visita_sintetica_valoracion(texto: str = "Datos del expediente") -> dict:
    return {
        "id": None,
        "fecha": "",
        "tecnico": "",
        "objeto_valoracion": texto,
    }


def _aplicar_observaciones_valoracion(datos: dict, observaciones) -> dict:
    if observaciones is None:
        return datos
    estado_observado = get_row_value(observaciones, "estado_observado")
    ocupacion_observada = get_row_value(observaciones, "ocupacion_observada")
    reforma_observada = get_row_value(observaciones, "reforma_observada")
    observaciones_portal = get_row_value(observaciones, "observaciones_portal")
    observaciones_cuadro_contadores = get_row_value(
        observaciones,
        "observaciones_cuadro_contadores",
    )
    if limpiar_texto(estado_observado):
        datos["estado_inmueble"] = estado_observado
    if limpiar_texto(ocupacion_observada):
        datos["regimen_ocupacion"] = ocupacion_observada
        if not limpiar_texto(datos.get("situacion_ocupacion")):
            datos["situacion_ocupacion"] = ocupacion_observada
    if limpiar_texto(reforma_observada) and not limpiar_texto(datos.get("observaciones_valoracion")):
        datos["observaciones_valoracion"] = f"Reforma observada en visita: {reforma_observada}"
    if limpiar_texto(observaciones_portal):
        datos["observaciones_portal"] = observaciones_portal
    if limpiar_texto(observaciones_cuadro_contadores):
        datos["observaciones_cuadro_contadores"] = observaciones_cuadro_contadores
    return datos


def _aplicar_resultado_valoracion(datos: dict, resultado) -> dict:
    if resultado is None:
        return datos
    if limpiar_texto(get_row_value(resultado, "valor_unitario")):
        datos["valor_unitario"] = get_row_value(resultado, "valor_unitario")
    if limpiar_texto(get_row_value(resultado, "valor_resultante")):
        datos["valor_resultante"] = get_row_value(resultado, "valor_resultante")
    if limpiar_texto(get_row_value(resultado, "valor_tasacion_final")):
        datos["valor_tasacion_final"] = get_row_value(resultado, "valor_tasacion_final")
    resumen = get_row_value(resultado, "resumen_calculo")
    if limpiar_texto(resumen) and not limpiar_texto(datos.get("observaciones_valoracion")):
        datos["observaciones_valoracion"] = resumen
    return datos


def construir_valoracion_visita_contexto(cur, visita, datos: dict | None = None) -> dict:
    visita_dict = row_to_dict(visita)
    visita_id = visita_dict.get("id")
    if datos is None:
        datos = cargar_valoracion_visita(cur, visita_id)
    grupos = [
        construir_grupo_valoracion(clave, titulo, campos, datos)
        for clave, titulo, campos in VALORACION_GRUPOS_CONTEXTO
    ]
    objeto = visita_dict.get("objeto_valoracion")
    if not objeto and visita_id is not None:
        objeto = describir_objeto_visita_informe(cur, visita)
    return {
        "visita": {
            "id": visita_id,
            "fecha": visita_dict.get("fecha", ""),
            "tecnico": visita_dict.get("tecnico", ""),
            "objeto": objeto or "Datos del expediente",
        },
        "grupos": grupos,
        "hay_datos": any(grupo["hay_datos"] for grupo in grupos),
    }


def cargar_valoracion_expediente_con_fallback(cur, expediente_id: int, visitas) -> list[dict]:
    valoracion_expediente = cur.execute(
        """
        SELECT *
        FROM valoracion_expediente
        WHERE expediente_id = ?
        """,
        (expediente_id,),
    ).fetchone()
    if valoracion_expediente is not None:
        datos = {**_valoracion_dict_vacio(), **row_to_dict(valoracion_expediente)}
        observaciones = cur.execute(
            """
            SELECT *
            FROM valoracion_visita_observaciones
            WHERE expediente_id = ?
            ORDER BY id DESC
            LIMIT 1
            """,
            (expediente_id,),
        ).fetchone()
        resultado = cur.execute(
            """
            SELECT *
            FROM valoracion_resultados
            WHERE expediente_id = ?
              AND COALESCE(activo, 1) = 1
            ORDER BY version DESC, id DESC
            LIMIT 1
            """,
            (expediente_id,),
        ).fetchone()
        datos = _aplicar_observaciones_valoracion(datos, observaciones)
        datos = _aplicar_resultado_valoracion(datos, resultado)
        visita = _visita_por_id(visitas, get_row_value(observaciones, "visita_id")) if observaciones else None
        if visita is None:
            visita = visitas[-1] if visitas else _visita_sintetica_valoracion()
        return [construir_valoracion_visita_contexto(cur, visita, datos)]

    valoraciones_legacy = []
    for visita in visitas or []:
        datos = cargar_valoracion_visita(cur, visita["id"])
        if _tiene_datos_valoracion(datos):
            valoraciones_legacy.append((visita, datos))

    if valoraciones_legacy:
        visita, datos = valoraciones_legacy[-1]
        observaciones = cur.execute(
            """
            SELECT *
            FROM valoracion_visita_observaciones
            WHERE visita_id = ?
            ORDER BY id DESC
            LIMIT 1
            """,
            (visita["id"],),
        ).fetchone()
        datos = _aplicar_observaciones_valoracion(datos, observaciones)
        return [construir_valoracion_visita_contexto(cur, visita, datos)]

    if visitas:
        observaciones = cur.execute(
            """
            SELECT *
            FROM valoracion_visita_observaciones
            WHERE expediente_id = ?
            ORDER BY id DESC
            LIMIT 1
            """,
            (expediente_id,),
        ).fetchone()
        if observaciones is not None:
            datos = _aplicar_observaciones_valoracion(
                _valoracion_dict_vacio(),
                observaciones,
            )
            visita = _visita_por_id(visitas, get_row_value(observaciones, "visita_id"))
            return [
                construir_valoracion_visita_contexto(
                    cur,
                    visita or visitas[-1],
                    datos,
                )
            ]
        return [construir_valoracion_visita_contexto(cur, visitas[0])]
    return []


def cargar_valoracion_eco_con_fallback(cur, expediente_id: int, visitas) -> dict:
    valoracion_expediente = cur.execute(
        """
        SELECT *
        FROM valoracion_expediente
        WHERE expediente_id = ?
        """,
        (expediente_id,),
    ).fetchone()
    if valoracion_expediente is not None:
        return {**_valoracion_dict_vacio(), **row_to_dict(valoracion_expediente)}

    for visita in reversed(visitas or []):
        datos = cargar_valoracion_visita(cur, visita["id"])
        if _tiene_datos_valoracion(datos):
            return {**_valoracion_dict_vacio(), **datos}
    return _valoracion_dict_vacio()


def _metodo_eco(datos: dict, clave: str, titulo: str) -> dict:
    aplicado = _valoracion_bool(datos.get(f"metodo_{clave}_aplicado"))
    descartado = _valoracion_bool(datos.get(f"metodo_{clave}_descartado"))
    if clave == "comparacion" and not aplicado:
        aplicado = _valoracion_bool(datos.get("metodo_comparacion_activo"))
    if clave == "coste" and not aplicado:
        aplicado = _valoracion_bool(datos.get("metodo_coste_activo"))
    return {
        "clave": clave,
        "titulo": titulo,
        "aplicado": aplicado,
        "descartado": descartado,
        "justificacion": limpiar_texto(datos.get(f"metodo_{clave}_justificacion")),
        "observaciones": limpiar_texto(datos.get(f"metodo_{clave}_observaciones")),
    }


def _crear_incidencia(tipo: str, origen: str, descripcion: str, visible: bool = True) -> dict:
    return {
        "tipo": tipo,
        "origen": origen,
        "descripcion": descripcion,
        "visible": visible,
    }


def construir_valoracion_eco(
    expediente: dict,
    datos: dict,
    comparables: list[dict],
) -> dict:
    metodos = [
        _metodo_eco(datos, "comparacion", "Comparación"),
        _metodo_eco(datos, "coste", "Coste"),
        _metodo_eco(datos, "actualizacion_rentas", "Actualización de rentas"),
        _metodo_eco(datos, "residual", "Residual"),
    ]
    metodo_comparacion = next(
        metodo for metodo in metodos if metodo["clave"] == "comparacion"
    )
    visibles_auto = datos.get("incidencias_automaticas_visibles", 1)
    visibles_manual = datos.get("incidencias_manuales_visibles", 1)
    incidencias = []
    if not limpiar_texto(expediente.get("referencia_catastral")):
        incidencias.append(
            _crear_incidencia(
                "advertencia",
                "automatica",
                "No consta referencia catastral en el expediente.",
                _valoracion_visible(visibles_auto),
            )
        )
    if not limpiar_texto(datos.get("superficie_adoptada_calculo")):
        incidencias.append(
            _crear_incidencia(
                "advertencia",
                "automatica",
                "No consta superficie adoptada para cálculo.",
                _valoracion_visible(visibles_auto),
            )
        )
    if not limpiar_texto(datos.get("fecha_valoracion")):
        incidencias.append(
            _crear_incidencia(
                "advertencia",
                "automatica",
                "No consta fecha de valoración.",
                _valoracion_visible(visibles_auto),
            )
        )
    if not limpiar_texto(datos.get("finalidad_valoracion")):
        incidencias.append(
            _crear_incidencia(
                "advertencia",
                "automatica",
                "No consta finalidad de valoración.",
                _valoracion_visible(visibles_auto),
            )
        )
    if metodo_comparacion["aplicado"] and len(comparables or []) < 3:
        incidencias.append(
            _crear_incidencia(
                "advertencia",
                "automatica",
                "El método de comparación figura aplicado con menos de 3 testigos.",
                _valoracion_visible(visibles_auto),
            )
        )
    if any(not limpiar_texto(comparable.get("fuente_testigo")) for comparable in comparables or []):
        incidencias.append(
            _crear_incidencia(
                "advertencia",
                "automatica",
                "Existen testigos sin fuente informada.",
                _valoracion_visible(visibles_auto),
            )
        )
    documentacion = limpiar_texto(datos.get("documentacion_utilizada")).lower()
    if not limpiar_texto(datos.get("datos_registrales")) and "nota simple" not in documentacion:
        incidencias.append(
            _crear_incidencia(
                "limitacion",
                "automatica",
                "No consta documentación registral suficiente.",
                _valoracion_visible(visibles_auto),
            )
        )
    if not limpiar_texto(expediente.get("referencia_catastral")) and "catastro" not in documentacion:
        incidencias.append(
            _crear_incidencia(
                "limitacion",
                "automatica",
                "No consta información catastral suficiente.",
                _valoracion_visible(visibles_auto),
            )
        )

    for tipo, campo in [
        ("condicionante", "incidencias_condicionantes_manuales"),
        ("advertencia", "incidencias_advertencias_manuales"),
        ("limitacion", "incidencias_limitaciones_manuales"),
    ]:
        for descripcion in _lineas_manuales(datos.get(campo)):
            incidencias.append(
                _crear_incidencia(
                    tipo,
                    "manual",
                    descripcion,
                    _valoracion_visible(visibles_manual),
                )
            )

    superficies = _grupo_eco(
        "superficies",
        "Superficies consideradas y superficie adoptada",
        [
            ("superficie_util", "Superficie útil"),
            ("superficie_construida", "Superficie construida"),
            ("superficie_registral", "Superficie registral"),
            ("superficie_catastral", "Superficie catastral"),
            ("superficie_comprobada", "Superficie comprobada"),
            ("superficie_computable", "Superficie computable"),
            ("superficie_adoptada_calculo", "Superficie adoptada para cálculo"),
            ("criterio_superficie_adoptada", "Criterio de adopción"),
        ],
        datos,
    )
    return {
        "nota_metodologica": VALORACION_NOTA_ECO,
        "finalidad": _grupo_eco(
            "finalidad",
            "Finalidad y alcance",
            [
                ("finalidad_valoracion", "Finalidad"),
                ("finalidad_otro", "Finalidad: otro / matiz"),
                ("alcance_valoracion", "Alcance"),
                ("fecha_valoracion", "Fecha de valoración"),
            ],
            datos,
        ),
        "base_valor": _grupo_eco(
            "base_valor",
            "Base de valor",
            VALORACION_BASE_VALOR_ITEMS,
            datos,
        ),
        "superficies": superficies,
        "metodos": {
            "titulo": "Métodos aplicados y descartados",
            "items": metodos,
            "hay_datos": any(
                metodo["aplicado"]
                or metodo["descartado"]
                or metodo["justificacion"]
                or metodo["observaciones"]
                for metodo in metodos
            ),
        },
        "incidencias": {
            "titulo": "Condicionantes, advertencias y limitaciones",
            "items": incidencias,
            "visibles": [incidencia for incidencia in incidencias if incidencia["visible"]],
            "hay_datos": bool(incidencias),
        },
    }


def _snapshot_json_valoracion(texto: str) -> dict:
    if not limpiar_texto(texto):
        return {}
    try:
        datos = json.loads(texto)
    except (TypeError, json.JSONDecodeError):
        return {}
    return datos if isinstance(datos, dict) else {}


def _hay_valor_comparable(valor) -> bool:
    if valor is None:
        return False
    if isinstance(valor, (int, float)):
        return True
    return bool(limpiar_texto(valor))


def _primer_valor_comparable(row, snapshot: dict, *campos):
    for campo in campos:
        valor = get_row_value(row, campo, None)
        if _hay_valor_comparable(valor):
            return valor
        valor_snapshot = snapshot.get(campo)
        if _hay_valor_comparable(valor_snapshot):
            return valor_snapshot
    return ""


def _valor_booleano_comparable(valor):
    if isinstance(valor, (int, float)):
        return bool(valor)
    texto = limpiar_texto(valor).lower()
    if texto in {"", "-"}:
        return None
    if texto in {"1", "true", "si", "sí", "yes", "on"}:
        return True
    if texto in {"0", "false", "no", "off"}:
        return False
    return None


def formatear_booleano_comparable(valor) -> str:
    booleano = _valor_booleano_comparable(valor)
    if booleano is None:
        return "-"
    return "Sí" if booleano else "No"


def _extraer_numero_planta(valor):
    texto = limpiar_texto(valor).lower()
    if not texto:
        return None
    if "bajo" in texto:
        return 0
    match = re.search(r"-?\d+", texto)
    if not match:
        return None
    try:
        return int(match.group(0))
    except ValueError:
        return None


def construir_advertencias_tecnicas_testigo(comparable: dict) -> list[str]:
    advertencias = []
    planta = _extraer_numero_planta(get_row_value(comparable, "planta"))
    tiene_ascensor = _valor_booleano_comparable(
        get_row_value(comparable, "ascensor")
    )
    if planta is not None and planta >= 4 and tiene_ascensor is False:
        advertencias.append("4ª planta o superior sin ascensor.")

    superficie_construida = parsear_float_valoracion(
        get_row_value(comparable, "superficie_construida")
    )
    superficie_util = parsear_float_valoracion(
        get_row_value(comparable, "superficie_util")
    )
    if superficie_construida is None and superficie_util is None:
        advertencias.append("Falta superficie útil y construida.")
    elif superficie_construida and superficie_util:
        divergencia = abs(superficie_construida - superficie_util) / max(
            superficie_construida,
            superficie_util,
        )
        if divergencia >= 0.35:
            advertencias.append("Superficie útil/construida muy divergente.")

    if not limpiar_texto(get_row_value(comparable, "estado_conservacion")):
        advertencias.append("Estado de conservación desconocido.")
    if not limpiar_texto(get_row_value(comparable, "ano_construccion")):
        advertencias.append("Año de construcción ausente.")
    return advertencias


def _construir_comparable_nuevo(row) -> dict:
    snapshot = _snapshot_json_valoracion(get_row_value(row, "snapshot_json"))
    comparable = {
        "id": get_row_value(row, "expediente_testigo_id"),
        "visita_id": None,
        "visita_fecha": "",
        "origen": "modelo_nuevo",
        "expediente_testigo_id": get_row_value(row, "expediente_testigo_id"),
        "testigo_id": get_row_value(row, "testigo_id"),
        "orden": get_row_value(row, "orden"),
        "incluido": get_row_value(row, "incluido"),
        "incluido_calculo": get_row_value(row, "incluido_calculo"),
        "peso_porcentaje": get_row_value(row, "peso_porcentaje"),
        "motivo_ponderacion": get_row_value(row, "motivo_ponderacion"),
        "representatividad": get_row_value(row, "representatividad"),
        "motivo_exclusion": get_row_value(row, "motivo_exclusion"),
        "observaciones_ponderacion": get_row_value(row, "observaciones_ponderacion"),
        "notas_seleccion": get_row_value(row, "notas_seleccion"),
        "valor_unitario_base": get_row_value(row, "valor_unitario_base"),
        "valor_unitario_ajustado": get_row_value(row, "valor_unitario_ajustado"),
        "coeficiente_total": get_row_value(row, "coeficiente_total"),
        "justificacion_ajustes": get_row_value(row, "justificacion"),
        "snapshot": snapshot,
        "ajustes": {
            "superficie_construida": get_row_value(row, "ajuste_superficie_construida"),
            "ubicacion": get_row_value(row, "ajuste_ubicacion"),
            "antiguedad": get_row_value(row, "ajuste_antiguedad"),
            "calidades": get_row_value(row, "ajuste_calidades"),
            "caracteristicas_constructivas": get_row_value(
                row, "ajuste_caracteristicas_constructivas"
            ),
            "coeficiente_total": get_row_value(row, "coeficiente_total"),
            "justificacion": get_row_value(row, "justificacion"),
        },
        **{
            "direccion_testigo": _primer_valor_comparable(row, snapshot, "direccion_testigo"),
            "fuente_testigo": _primer_valor_comparable(row, snapshot, "fuente_testigo"),
            "fuente_tipo": _primer_valor_comparable(row, snapshot, "fuente_tipo"),
            "fuente_detalle": _primer_valor_comparable(row, snapshot, "fuente_detalle"),
            "url_fuente": _primer_valor_comparable(row, snapshot, "url_fuente"),
            "fecha_testigo": _primer_valor_comparable(row, snapshot, "fecha_testigo"),
            "fecha_captura": _primer_valor_comparable(row, snapshot, "fecha_captura"),
            "precio_oferta": _primer_valor_comparable(row, snapshot, "precio_oferta"),
            "precio_depurado": _primer_valor_comparable(row, snapshot, "precio_depurado"),
            "superficie_tomada": _primer_valor_comparable(row, snapshot, "superficie_tomada"),
            "tipo_superficie_tomada": _primer_valor_comparable(
                row, snapshot, "tipo_superficie_tomada"
            ),
            "precio_unitario_inicial": _primer_valor_comparable(
                row, snapshot, "precio_unitario_inicial"
            ),
            "valor_unitario": _primer_valor_comparable(
                row,
                snapshot,
                "valor_unitario_ajustado",
                "valor_unitario_base",
                "valor_unitario",
            ),
            "superficie_construida": _primer_valor_comparable(
                row, snapshot, "superficie_construida"
            ),
            "superficie_util": _primer_valor_comparable(row, snapshot, "superficie_util"),
            "tipologia": _primer_valor_comparable(row, snapshot, "tipologia"),
            "planta": _primer_valor_comparable(row, snapshot, "planta"),
            "dormitorios": _primer_valor_comparable(row, snapshot, "dormitorios"),
            "banos": _primer_valor_comparable(row, snapshot, "banos"),
            "aseos": _primer_valor_comparable(row, snapshot, "aseos"),
            "ascensor": _primer_valor_comparable(row, snapshot, "ascensor"),
            "garaje": _primer_valor_comparable(row, snapshot, "garaje"),
            "trastero": _primer_valor_comparable(row, snapshot, "trastero"),
            "terraza": _primer_valor_comparable(row, snapshot, "terraza"),
            "es_exterior": _primer_valor_comparable(row, snapshot, "es_exterior"),
            "balcon": _primer_valor_comparable(row, snapshot, "balcon"),
            "patio": _primer_valor_comparable(row, snapshot, "patio"),
            "estado_conservacion": _primer_valor_comparable(
                row, snapshot, "estado_conservacion"
            ),
            "antiguedad": _primer_valor_comparable(row, snapshot, "antiguedad"),
            "ano_construccion": _primer_valor_comparable(
                row, snapshot, "ano_construccion"
            ),
            "ano_reforma": _primer_valor_comparable(row, snapshot, "ano_reforma"),
            "calidad_constructiva": _primer_valor_comparable(
                row, snapshot, "calidad_constructiva"
            ),
            "aire_acondicionado": _primer_valor_comparable(
                row, snapshot, "aire_acondicionado"
            ),
            "tipo_calefaccion": _primer_valor_comparable(
                row, snapshot, "tipo_calefaccion"
            ),
            "certificacion_energetica": _primer_valor_comparable(
                row, snapshot, "certificacion_energetica"
            ),
            "dato_verificado": _primer_valor_comparable(row, snapshot, "dato_verificado"),
            "testigo_visitado": _primer_valor_comparable(row, snapshot, "testigo_visitado"),
            "fiabilidad_dato": _primer_valor_comparable(row, snapshot, "fiabilidad_dato"),
            "similitud_inmueble": _primer_valor_comparable(row, snapshot, "similitud_inmueble"),
            "estado_mercado": _primer_valor_comparable(row, snapshot, "estado_mercado"),
            "observaciones_economicas": _primer_valor_comparable(
                row, snapshot, "observaciones_economicas"
            ),
            "visitado": _primer_valor_comparable(row, snapshot, "visitado"),
            "observaciones": _primer_valor_comparable(
                row, snapshot, "notas_seleccion", "observaciones"
            ),
        },
    }
    preparado = preparar_testigo_comparacion(comparable)
    if not limpiar_texto(comparable.get("precio_unitario_inicial")):
        comparable["precio_unitario_inicial"] = preparado["precio_unitario_inicial"]
    comparable["advertencias_calculo"] = preparado["advertencias_calculo"]
    comparable["advertencias_calculo_texto"] = "\n".join(preparado["advertencias_calculo"])
    return comparable


def construir_comparable_valoracion_contexto(comparable, visita=None) -> dict:
    visita_fecha = get_row_value(comparable, "visita_fecha")
    if not limpiar_texto(visita_fecha) and visita is not None:
        visita_fecha = get_row_value(visita, "fecha")
    preparado = preparar_testigo_comparacion(
        {
            "precio_oferta": get_row_value(comparable, "precio_oferta"),
            "precio_depurado": get_row_value(comparable, "precio_depurado"),
            "superficie_tomada": get_row_value(comparable, "superficie_tomada")
            or get_row_value(comparable, "superficie_construida"),
            "fuente_testigo": get_row_value(comparable, "fuente_testigo"),
            "fuente_detalle": get_row_value(comparable, "fuente_detalle"),
            "fecha_testigo": get_row_value(comparable, "fecha_testigo"),
            "fiabilidad_dato": get_row_value(comparable, "fiabilidad_dato"),
        }
    )
    advertencias_calculo = get_row_value(
        comparable,
        "advertencias_calculo",
        preparado["advertencias_calculo"],
    )
    precio_unitario_inicial = get_row_value(
        comparable,
        "precio_unitario_inicial",
        preparado["precio_unitario_inicial"],
    )
    homogeneizacion = get_row_value(comparable, "homogeneizacion", {})
    unitario_homogeneizado = homogeneizacion.get("unitario_homogeneizado")
    incluido_calculo = get_row_value(comparable, "incluido_calculo")
    if incluido_calculo in ("", None):
        incluido_calculo = get_row_value(comparable, "incluido", 1)
    unitario_para_resumen = unitario_homogeneizado or precio_unitario_inicial
    advertencias_tecnicas = construir_advertencias_tecnicas_testigo(comparable)
    campos = []
    for campo, etiqueta in COMPARABLES_COLUMNAS:
        valor = (
            precio_unitario_inicial
            if campo == "precio_unitario_inicial"
            else "\n".join(advertencias_calculo)
            if campo == "advertencias_calculo_texto"
            else unitario_homogeneizado
            if campo == "unitario_homogeneizado"
            else unitario_para_resumen
            if campo == "unitario_para_resumen"
            else incluido_calculo
            if campo == "incluido_calculo"
            else get_row_value(comparable, campo)
        )
        formatter = COMPARABLES_FORMATTERS.get(campo)
        campos.append((etiqueta, formatter(valor) if formatter else valor))
    return {
        "id": get_row_value(comparable, "id"),
        "visita_id": get_row_value(comparable, "visita_id"),
        "visita_fecha": visita_fecha,
        "origen": get_row_value(comparable, "origen", "legacy"),
        "direccion_testigo": get_row_value(comparable, "direccion_testigo"),
        "fuente_testigo": get_row_value(comparable, "fuente_testigo"),
        "fuente_tipo": get_row_value(comparable, "fuente_tipo"),
        "fuente_detalle": get_row_value(comparable, "fuente_detalle"),
        "url_fuente": get_row_value(comparable, "url_fuente"),
        "fecha_testigo": get_row_value(comparable, "fecha_testigo"),
        "fecha_captura": get_row_value(comparable, "fecha_captura"),
        "precio_oferta": get_row_value(comparable, "precio_oferta"),
        "precio_depurado": get_row_value(comparable, "precio_depurado"),
        "superficie_tomada": get_row_value(comparable, "superficie_tomada"),
        "tipo_superficie_tomada": get_row_value(comparable, "tipo_superficie_tomada"),
        "precio_unitario_inicial": precio_unitario_inicial,
        "valor_unitario": get_row_value(comparable, "valor_unitario"),
        "superficie_construida": get_row_value(comparable, "superficie_construida"),
        "superficie_util": get_row_value(comparable, "superficie_util"),
        "planta": get_row_value(comparable, "planta"),
        "banos": get_row_value(comparable, "banos"),
        "aseos": get_row_value(comparable, "aseos"),
        "ascensor": get_row_value(comparable, "ascensor"),
        "es_exterior": get_row_value(comparable, "es_exterior"),
        "balcon": get_row_value(comparable, "balcon"),
        "terraza": get_row_value(comparable, "terraza"),
        "patio": get_row_value(comparable, "patio"),
        "ano_construccion": get_row_value(comparable, "ano_construccion"),
        "ano_reforma": get_row_value(comparable, "ano_reforma"),
        "aire_acondicionado": get_row_value(comparable, "aire_acondicionado"),
        "tipo_calefaccion": get_row_value(comparable, "tipo_calefaccion"),
        "garaje": get_row_value(comparable, "garaje"),
        "trastero": get_row_value(comparable, "trastero"),
        "certificacion_energetica": get_row_value(
            comparable,
            "certificacion_energetica",
        ),
        "estado_conservacion": get_row_value(comparable, "estado_conservacion"),
        "dato_verificado": get_row_value(comparable, "dato_verificado"),
        "testigo_visitado": get_row_value(comparable, "testigo_visitado"),
        "fiabilidad_dato": get_row_value(comparable, "fiabilidad_dato"),
        "similitud_inmueble": get_row_value(comparable, "similitud_inmueble"),
        "estado_mercado": get_row_value(comparable, "estado_mercado"),
        "observaciones_economicas": get_row_value(comparable, "observaciones_economicas"),
        "advertencias_calculo": advertencias_calculo,
        "advertencias_calculo_texto": "\n".join(advertencias_calculo),
        "advertencias_tecnicas": advertencias_tecnicas,
        "advertencias_tecnicas_texto": "\n".join(advertencias_tecnicas),
        "ajustes_homogeneizacion": homogeneizacion.get("ajustes", []),
        "unitario_inicial": homogeneizacion.get("unitario_inicial")
        or precio_unitario_inicial,
        "unitario_homogeneizado": unitario_homogeneizado,
        "unitario_para_resumen": unitario_para_resumen,
        "incluido_calculo": incluido_calculo,
        "peso_porcentaje": get_row_value(comparable, "peso_porcentaje"),
        "representatividad": get_row_value(comparable, "representatividad"),
        "motivo_ponderacion": get_row_value(comparable, "motivo_ponderacion"),
        "motivo_exclusion": get_row_value(comparable, "motivo_exclusion"),
        "observaciones_ponderacion": get_row_value(
            comparable,
            "observaciones_ponderacion",
        ),
        "ajuste_total_importe_m2": homogeneizacion.get("ajuste_total_importe_m2"),
        "ajuste_total_porcentaje_equivalente": homogeneizacion.get(
            "ajuste_total_porcentaje_equivalente"
        ),
        "pasos_homogeneizacion": homogeneizacion.get("pasos", []),
        "advertencias_homogeneizacion": homogeneizacion.get("advertencias", []),
        "unitario_homogeneizado_fmt": formatear_precio_unitario_es(
            unitario_homogeneizado
        ),
        "unitario_para_resumen_fmt": formatear_precio_unitario_es(
            unitario_para_resumen
        ),
        "peso_porcentaje_fmt": formatear_porcentaje_es(
            get_row_value(comparable, "peso_porcentaje")
        ),
        "ajuste_total_importe_m2_fmt": formatear_precio_unitario_es(
            homogeneizacion.get("ajuste_total_importe_m2")
        ),
        "expediente_testigo_id": get_row_value(comparable, "expediente_testigo_id"),
        "testigo_id": get_row_value(comparable, "testigo_id"),
        "orden": get_row_value(comparable, "orden"),
        "incluido": get_row_value(comparable, "incluido"),
        "notas_seleccion": get_row_value(comparable, "notas_seleccion"),
        "valor_unitario_base": get_row_value(comparable, "valor_unitario_base"),
        "valor_unitario_ajustado": get_row_value(
            comparable,
            "valor_unitario_ajustado",
        ),
        "coeficiente_total": get_row_value(comparable, "coeficiente_total"),
        "valor_unitario_base_fmt": formatear_precio_unitario_es(
            get_row_value(comparable, "valor_unitario_base")
        ),
        "valor_unitario_ajustado_fmt": formatear_precio_unitario_es(
            get_row_value(comparable, "valor_unitario_ajustado")
        ),
        "coeficiente_total_fmt": formatear_coeficiente_es(
            get_row_value(comparable, "coeficiente_total")
        ),
        "precio_oferta_fmt": formatear_moneda_es(
            get_row_value(comparable, "precio_oferta")
        ),
        "precio_depurado_fmt": formatear_moneda_es(
            get_row_value(comparable, "precio_depurado")
        ),
        "precio_unitario_inicial_fmt": formatear_precio_unitario_es(
            get_row_value(comparable, "precio_unitario_inicial")
            or precio_unitario_inicial
        ),
        "superficie_tomada_fmt": formatear_superficie_es(
            get_row_value(comparable, "superficie_tomada")
        ),
        "valor_unitario_fmt": formatear_precio_unitario_es(
            get_row_value(comparable, "valor_unitario")
        ),
        "superficie_construida_fmt": formatear_superficie_es(
            get_row_value(comparable, "superficie_construida")
        ),
        "superficie_util_fmt": formatear_superficie_es(
            get_row_value(comparable, "superficie_util")
        ),
        "ascensor_fmt": formatear_booleano_comparable(
            get_row_value(comparable, "ascensor")
        ),
        "es_exterior_fmt": formatear_booleano_comparable(
            get_row_value(comparable, "es_exterior")
        ),
        "balcon_fmt": formatear_booleano_comparable(get_row_value(comparable, "balcon")),
        "terraza_fmt": formatear_booleano_comparable(get_row_value(comparable, "terraza")),
        "patio_fmt": formatear_booleano_comparable(get_row_value(comparable, "patio")),
        "aire_acondicionado_fmt": formatear_booleano_comparable(
            get_row_value(comparable, "aire_acondicionado")
        ),
        "garaje_fmt": formatear_booleano_comparable(get_row_value(comparable, "garaje")),
        "trastero_fmt": formatear_booleano_comparable(
            get_row_value(comparable, "trastero")
        ),
        "justificacion_ajustes": get_row_value(comparable, "justificacion_ajustes"),
        "snapshot": get_row_value(comparable, "snapshot", {}),
        "ajustes": get_row_value(comparable, "ajustes", {}),
        "campos": construir_campos_informe(campos),
    }


def cargar_ajustes_homogeneizacion_informe(cur, expediente_testigo_id: int) -> list[dict]:
    return [
        row_to_dict(row)
        for row in cur.execute(
            """
            SELECT *
            FROM valoracion_testigo_ajustes
            WHERE expediente_testigo_id = ?
              AND COALESCE(variable, '') != ''
              AND COALESCE(activo, 1) = 1
            ORDER BY COALESCE(orden, 9999) ASC, id ASC
            """,
            (expediente_testigo_id,),
        ).fetchall()
    ]


def cargar_comparables_valoracion_con_fallback(cur, expediente_id: int, visitas) -> list[dict]:
    comparables_nuevos = cur.execute(
        """
        SELECT vet.id AS expediente_testigo_id,
               vet.expediente_id,
               vet.testigo_id,
               vet.orden,
               vet.incluido,
               vet.incluido_calculo,
               vet.peso_porcentaje,
               vet.motivo_ponderacion,
               vet.representatividad,
               vet.motivo_exclusion,
               vet.observaciones_ponderacion,
               vet.snapshot_json,
               vet.notas_seleccion,
               vet.valor_unitario_base,
               vet.valor_unitario_ajustado,
               vet.valor_resultante,
               tv.direccion_testigo,
               tv.referencia_testigo,
               tv.fuente_testigo,
               tv.url_fuente,
               tv.fecha_testigo,
               tv.codigo_postal,
               tv.municipio,
               tv.provincia,
               tv.precio_oferta,
               tv.precio_depurado,
               tv.precio_unitario_inicial,
               tv.superficie_tomada,
               tv.tipo_superficie_tomada,
               tv.precio_cierre,
               tv.superficie_construida,
               tv.superficie_util,
               tv.valor_unitario,
               tv.tipologia,
               tv.planta,
               tv.dormitorios,
               tv.banos,
               tv.aseos,
               tv.ascensor,
               tv.garaje,
               tv.trastero,
               tv.terraza,
               tv.es_exterior,
               tv.balcon,
               tv.patio,
               tv.estado_conservacion,
               tv.antiguedad,
               tv.ano_construccion,
               tv.ano_reforma,
               tv.calidad_constructiva,
               tv.caracteristicas_constructivas,
               tv.ubicacion,
               tv.aire_acondicionado,
               tv.tipo_calefaccion,
               tv.certificacion_energetica,
               tv.fuente_tipo,
               tv.fuente_detalle,
               tv.fecha_captura,
               tv.dato_verificado,
               tv.testigo_visitado,
               tv.fiabilidad_dato,
               tv.similitud_inmueble,
               tv.estado_mercado,
               tv.observaciones_economicas,
               tv.visitado,
               tv.validacion_estado,
               tv.reutilizable,
               tv.observaciones,
               vta.ajuste_superficie_construida,
               vta.ajuste_ubicacion,
               vta.ajuste_antiguedad,
               vta.ajuste_calidades,
               vta.ajuste_caracteristicas_constructivas,
               vta.coeficiente_total,
               vta.justificacion
        FROM valoracion_expediente_testigos vet
        LEFT JOIN testigos_valoracion tv ON tv.id = vet.testigo_id
        LEFT JOIN valoracion_testigo_ajustes vta
          ON vta.expediente_testigo_id = vet.id
         AND COALESCE(vta.variable, '') = ''
        WHERE vet.expediente_id = ?
          AND COALESCE(vet.incluido, 1) = 1
        ORDER BY COALESCE(vet.orden, 9999) ASC, vet.id ASC
        """,
        (expediente_id,),
    ).fetchall()
    if comparables_nuevos:
        comparables_contexto = []
        for row in comparables_nuevos:
            comparable = _construir_comparable_nuevo(row)
            ajustes_homogeneizacion = cargar_ajustes_homogeneizacion_informe(
                cur,
                comparable["expediente_testigo_id"],
            )
            comparable["homogeneizacion"] = preparar_matriz_homogeneizacion(
                comparable,
                ajustes_homogeneizacion,
            )
            comparables_contexto.append(
                construir_comparable_valoracion_contexto(comparable)
            )
        return comparables_contexto

    comparables_contexto = []
    for visita in visitas or []:
        comparables = cur.execute(
            """
            SELECT *
            FROM comparables_valoracion
            WHERE visita_id = ?
            ORDER BY id ASC
            """,
            (visita["id"],),
        ).fetchall()
        comparables_contexto.extend(
            construir_comparable_valoracion_contexto(comparable, visita)
            for comparable in comparables
        )
    return comparables_contexto


def construir_resumen_comparacion_contexto(comparables: list[dict]) -> dict:
    resumen = preparar_resumen_comparacion(comparables)
    resumen.update(
        {
            "unitario_minimo_fmt": formatear_precio_unitario_es(
                resumen.get("unitario_minimo")
            ),
            "unitario_maximo_fmt": formatear_precio_unitario_es(
                resumen.get("unitario_maximo")
            ),
            "unitario_medio_fmt": formatear_precio_unitario_es(
                resumen.get("unitario_medio")
            ),
            "unitario_mediana_fmt": formatear_precio_unitario_es(
                resumen.get("unitario_mediana")
            ),
            "unitario_ponderado_fmt": formatear_precio_unitario_es(
                resumen.get("unitario_ponderado")
            ),
            "propuesta_unitaria_orientativa_fmt": formatear_precio_unitario_es(
                resumen.get("propuesta_unitaria_orientativa")
            ),
            "suma_pesos_fmt": formatear_porcentaje_es(resumen.get("suma_pesos")),
        }
    )
    return resumen


def grupo_valoracion_por_clave(valoraciones: list[dict], clave: str) -> list[dict]:
    grupos = []
    for bloque in valoraciones or []:
        grupos.extend(
            grupo
            for grupo in bloque.get("grupos", [])
            if grupo.get("clave") == clave
        )
    return grupos


def grupo_valoracion_tiene_datos(valoraciones: list[dict], clave: str) -> bool:
    return any(grupo.get("hay_datos") for grupo in grupo_valoracion_por_clave(valoraciones, clave))


def grupo_valoracion_tiene_campo(valoraciones: list[dict], clave: str, etiquetas: set[str]) -> bool:
    for grupo in grupo_valoracion_por_clave(valoraciones, clave):
        for campo in grupo.get("campos", []):
            if campo.get("label") in etiquetas and limpiar_texto(campo.get("value")) not in {"", "-"}:
                return True
    return False


def grupo_valoracion_tiene_prefijo(valoraciones: list[dict], clave: str, prefijo: str) -> bool:
    for grupo in grupo_valoracion_por_clave(valoraciones, clave):
        for campo in grupo.get("campos", []):
            if str(campo.get("label", "")).startswith(prefijo) and limpiar_texto(campo.get("value")) not in {"", "-"}:
                return True
    return False


def construir_completitud_valoracion(valoraciones: list[dict], comparables: list[dict]) -> dict:
    comprobaciones = [
        (
            "documentacion",
            "Falta documentacion utilizada.",
            grupo_valoracion_tiene_datos(valoraciones, "documentacion"),
        ),
        (
            "identificacion",
            "Falta identificacion del bien.",
            grupo_valoracion_tiene_campo(valoraciones, "identificacion", {"Identificación del bien"}),
        ),
        (
            "superficies",
            "Faltan superficies de referencia.",
            grupo_valoracion_tiene_prefijo(valoraciones, "identificacion", "Superficie"),
        ),
        (
            "situacion_legal",
            "Falta situacion legal u ocupacional.",
            grupo_valoracion_tiene_datos(valoraciones, "situacion_legal"),
        ),
        (
            "entorno",
            "Falta descripcion del entorno.",
            grupo_valoracion_tiene_datos(valoraciones, "entorno"),
        ),
        (
            "metodo",
            "Falta metodo o criterios de valoracion.",
            grupo_valoracion_tiene_datos(valoraciones, "metodo"),
        ),
        (
            "comparables",
            "Faltan comparables de mercado.",
            bool(comparables),
        ),
        (
            "resultado",
            "Falta resultado de valoracion.",
            grupo_valoracion_tiene_datos(valoraciones, "resultado"),
        ),
        (
            "limitaciones",
            "Faltan limitaciones o condicionantes.",
            grupo_valoracion_tiene_datos(valoraciones, "limitaciones"),
        ),
    ]
    advertencias = [
        {"clave": clave, "mensaje": mensaje}
        for clave, mensaje, ok in comprobaciones
        if not ok
    ]
    return {
        "advertencias": advertencias,
        "completo": not advertencias,
        "total_advertencias": len(advertencias),
    }


def add_apartado_valoracion_por_visitas(
    doc: Document, numero_apartado: int, titulo: str, visitas_valoracion, campos
) -> None:
    doc.add_heading(f"{numero_apartado}. {titulo}", level=1)

    hay_datos = False
    for indice, bloque in enumerate(visitas_valoracion, start=1):
        valoracion = bloque["valoracion"]
        if not any(limpiar_texto(valoracion[campo]) for campo, _ in campos):
            continue
        hay_datos = True
        visita = bloque["visita"]
        doc.add_heading(
            f"{numero_apartado}.{indice} Visita de fecha {valor_o_guion(visita['fecha'])}",
            level=2,
        )
        add_bloque_campos_si_hay(doc, campos, valoracion)

    if not hay_datos:
        add_parrafo(doc, "No constan datos registrados en este apartado.")


def add_apartado_comparables_valoracion(
    doc: Document, numero_apartado: int, visitas_valoracion
) -> None:
    doc.add_heading(f"{numero_apartado}. Testigos o comparables", level=1)

    hay_datos = False
    for indice, bloque in enumerate(visitas_valoracion, start=1):
        comparables = bloque["comparables"]
        if not comparables:
            continue
        hay_datos = True
        visita = bloque["visita"]
        doc.add_heading(
            f"{numero_apartado}.{indice} Visita de fecha {valor_o_guion(visita['fecha'])}",
            level=2,
        )
        add_tabla_comparables(doc, comparables)

    if not hay_datos:
        add_parrafo(doc, "No constan comparables registrados.")


def add_apartado_resultado_valoracion(
    doc: Document, numero_apartado: int, visitas_valoracion
) -> None:
    doc.add_heading(f"{numero_apartado}. Resultado de la valoración", level=1)

    if not visitas_valoracion:
        add_parrafo(doc, "No constan resultados de valoración registrados.")
        return

    hay_datos = False
    for indice, bloque in enumerate(visitas_valoracion, start=1):
        valoracion = bloque["valoracion"]
        if not any(
            limpiar_texto(valoracion[campo]) for campo, _ in VALORACION_RESULTADO_ITEMS
        ):
            continue
        hay_datos = True
        visita = bloque["visita"]
        doc.add_heading(
            f"{numero_apartado}.{indice} Visita de fecha {valor_o_guion(visita['fecha'])}",
            level=2,
        )
        add_valor_destacado(doc, "Valor unitario", valoracion["valor_unitario"])
        add_etiqueta_valor(doc, "Valor resultante", valoracion["valor_resultante"])
        add_valor_destacado(
            doc, "Valor de tasación final", valoracion["valor_tasacion_final"]
        )

    if not hay_datos:
        add_parrafo(doc, "No constan resultados de valoración registrados.")


def add_apartado_limitaciones_valoracion(
    doc: Document, numero_apartado: int, visitas_valoracion
) -> None:
    doc.add_heading(f"{numero_apartado}. Condicionantes y limitaciones", level=1)

    hay_datos = False
    for indice, bloque in enumerate(visitas_valoracion, start=1):
        valoracion = bloque["valoracion"]
        if not any(
            limpiar_texto(valoracion[campo]) for campo, _ in VALORACION_LIMITACIONES_ITEMS
        ):
            continue
        hay_datos = True
        visita = bloque["visita"]
        doc.add_heading(
            f"{numero_apartado}.{indice} Visita de fecha {valor_o_guion(visita['fecha'])}",
            level=2,
        )
        add_bloque_campos_si_hay(doc, VALORACION_LIMITACIONES_ITEMS, valoracion)

    if not hay_datos:
        add_parrafo(doc, "No constan condicionantes, limitaciones u observaciones adicionales.")


def cargar_inspeccion_visita(cur, visita_id: int) -> dict:
    general_campos = [
        campo for _, grupo in INSPECCION_GENERAL_GROUPS for campo, _ in grupo
    ] + ["observaciones_generales_inspeccion"]
    exterior_campos = [campo for campo, _ in INSPECCION_EXTERIOR_ITEMS] + [
        "observaciones_exteriores"
    ]
    comunes_campos = [campo for campo, _ in INSPECCION_ELEMENTOS_COMUNES_ITEMS] + [
        "observaciones_elementos_comunes"
    ]

    general = fila_o_dict_vacio(
        cur.execute(
        """
        SELECT *
        FROM inspeccion_general_visita
        WHERE visita_id = ?
        """,
        (visita_id,),
        ).fetchone(),
        general_campos,
    )
    exterior = fila_o_dict_vacio(
        cur.execute(
        """
        SELECT *
        FROM inspeccion_exterior
        WHERE visita_id = ?
        """,
        (visita_id,),
        ).fetchone(),
        exterior_campos,
    )
    comunes = fila_o_dict_vacio(
        cur.execute(
        """
        SELECT *
        FROM inspeccion_elementos_comunes
        WHERE visita_id = ?
        """,
        (visita_id,),
        ).fetchone(),
        comunes_campos,
    )
    estancias = cur.execute(
        """
        SELECT e.id AS estancia_id,
               e.nombre AS estancia_nombre,
               e.tipo_estancia,
               e.planta,
               e.acabado_pavimento,
               e.acabado_paramento,
               e.acabado_techo,
               ie.*
        FROM estancias e
        LEFT JOIN inspeccion_estancias ie
               ON ie.estancia_id = e.id
              AND ie.visita_id = e.visita_id
        WHERE e.visita_id = ?
        ORDER BY e.id ASC
        """,
        (visita_id,),
    ).fetchall()

    return {
        "general": general,
        "exterior": exterior,
        "comunes": comunes,
        "estancias": estancias,
    }


def add_apartado_condiciones_visita_inspeccion(
    doc: Document, numero_apartado: int, visitas_inspeccion
) -> None:
    doc.add_heading(f"{numero_apartado}. Condiciones de la visita / inspección", level=1)

    if not visitas_inspeccion:
        add_parrafo(doc, "No constan visitas registradas en el expediente.")
        return

    for indice, bloque in enumerate(visitas_inspeccion, start=1):
        visita = bloque["visita"]
        climatologia = bloque["climatologia"]
        doc.add_heading(
            f"{numero_apartado}.{indice} Visita de fecha {valor_o_guion(visita['fecha'])}",
            level=2,
        )
        add_etiqueta_valor(doc, "Técnico", visita["tecnico"])
        add_etiqueta_valor(doc, "Observaciones de visita", visita["observaciones_visita"])
        if climatologia:
            add_etiqueta_valor(doc, "Climatología", climatologia["resumen"])
        else:
            add_etiqueta_valor(doc, "Climatología", "No consta climatología registrada")


def add_apartado_comprobaciones_generales(
    doc: Document, numero_apartado: int, visitas_inspeccion
) -> None:
    doc.add_heading(f"{numero_apartado}. Comprobaciones generales", level=1)

    if not visitas_inspeccion:
        add_parrafo(doc, "No constan comprobaciones generales registradas.")
        return

    for indice, bloque in enumerate(visitas_inspeccion, start=1):
        general = bloque["inspeccion"]["general"]
        visita = bloque["visita"]
        doc.add_heading(
            f"{numero_apartado}.{indice} Visita de fecha {valor_o_guion(visita['fecha'])}",
            level=2,
        )
        for subindice, (titulo, items) in enumerate(INSPECCION_GENERAL_GROUPS, start=1):
            doc.add_heading(f"{numero_apartado}.{indice}.{subindice} {titulo}", level=3)
            add_estado_checklist(doc, items, general)
        if limpiar_texto(general["observaciones_generales_inspeccion"]):
            add_etiqueta_valor(
                doc,
                "Observaciones generales de inspección",
                general["observaciones_generales_inspeccion"],
            )


def add_apartado_inspeccion_estancias(
    doc: Document, numero_apartado: int, visitas_inspeccion
) -> None:
    doc.add_heading(f"{numero_apartado}. Inspección por estancias", level=1)

    if not visitas_inspeccion:
        add_parrafo(doc, "No constan estancias inspeccionadas.")
        return

    for indice_visita, bloque in enumerate(visitas_inspeccion, start=1):
        visita = bloque["visita"]
        estancias = bloque["inspeccion"]["estancias"]
        doc.add_heading(
            f"{numero_apartado}.{indice_visita} Visita de fecha {valor_o_guion(visita['fecha'])}",
            level=2,
        )
        if not estancias:
            add_parrafo(doc, "No constan estancias registradas en esta visita.")
            continue

        for indice_estancia, estancia in enumerate(estancias, start=1):
            doc.add_heading(
                f"{numero_apartado}.{indice_visita}.{indice_estancia} {valor_o_guion(estancia['estancia_nombre'])}",
                level=3,
            )
            add_etiqueta_valor(doc, "Tipo de estancia", estancia["tipo_estancia"])
            add_etiqueta_valor(doc, "Planta", estancia["planta"])
            add_etiqueta_valor(doc, "Acabado de pavimento", estancia["acabado_pavimento"])
            add_etiqueta_valor(doc, "Acabado de paramento", estancia["acabado_paramento"])
            add_etiqueta_valor(doc, "Acabado de techo", estancia["acabado_techo"])
            add_estado_checklist(
                doc,
                obtener_items_inspeccion_estancia(estancia["tipo_estancia"]),
                estancia,
            )
            if limpiar_texto(estancia["observaciones_estancia_inspeccion"]):
                add_etiqueta_valor(
                    doc,
                    "Observaciones de inspección",
                    estancia["observaciones_estancia_inspeccion"],
                )


def add_apartado_inspeccion_exterior(
    doc: Document, numero_apartado: int, visitas_inspeccion
) -> None:
    doc.add_heading(f"{numero_apartado}. Inspección exterior", level=1)

    if not visitas_inspeccion:
        add_parrafo(doc, "No constan comprobaciones exteriores registradas.")
        return

    for indice, bloque in enumerate(visitas_inspeccion, start=1):
        exterior = bloque["inspeccion"]["exterior"]
        visita = bloque["visita"]
        doc.add_heading(
            f"{numero_apartado}.{indice} Visita de fecha {valor_o_guion(visita['fecha'])}",
            level=2,
        )
        add_estado_checklist(doc, INSPECCION_EXTERIOR_ITEMS, exterior)
        if limpiar_texto(exterior["observaciones_exteriores"]):
            add_etiqueta_valor(
                doc,
                "Observaciones exteriores",
                exterior["observaciones_exteriores"],
            )


def add_apartado_elementos_comunes(
    doc: Document, numero_apartado: int, visitas_inspeccion
) -> None:
    doc.add_heading(f"{numero_apartado}. Elementos comunes", level=1)

    if not visitas_inspeccion:
        add_parrafo(doc, "No constan elementos comunes inspeccionados.")
        return

    for indice, bloque in enumerate(visitas_inspeccion, start=1):
        comunes = bloque["inspeccion"]["comunes"]
        visita = bloque["visita"]
        doc.add_heading(
            f"{numero_apartado}.{indice} Visita de fecha {valor_o_guion(visita['fecha'])}",
            level=2,
        )
        add_estado_checklist(doc, INSPECCION_ELEMENTOS_COMUNES_ITEMS, comunes)
        if limpiar_texto(comunes["observaciones_elementos_comunes"]):
            add_etiqueta_valor(
                doc,
                "Observaciones de elementos comunes",
                comunes["observaciones_elementos_comunes"],
            )


def add_apartado_observaciones_tecnicas(
    doc: Document, numero_apartado: int, visitas_inspeccion
) -> None:
    doc.add_heading(f"{numero_apartado}. Observaciones técnicas", level=1)

    hay_observaciones = False
    for indice, bloque in enumerate(visitas_inspeccion, start=1):
        visita = bloque["visita"]
        general = bloque["inspeccion"]["general"]
        exterior = bloque["inspeccion"]["exterior"]
        comunes = bloque["inspeccion"]["comunes"]
        estancias = bloque["inspeccion"]["estancias"]
        lineas = []

        if limpiar_texto(visita["observaciones_visita"]):
            lineas.append(("Observaciones de visita", visita["observaciones_visita"]))
        if limpiar_texto(general["observaciones_generales_inspeccion"]):
            lineas.append(
                (
                    "Observaciones generales de inspección",
                    general["observaciones_generales_inspeccion"],
                )
            )
        if limpiar_texto(exterior["observaciones_exteriores"]):
            lineas.append(("Observaciones exteriores", exterior["observaciones_exteriores"]))
        if limpiar_texto(comunes["observaciones_elementos_comunes"]):
            lineas.append(
                (
                    "Observaciones de elementos comunes",
                    comunes["observaciones_elementos_comunes"],
                )
            )

        for estancia in estancias:
            if limpiar_texto(estancia["observaciones_estancia_inspeccion"]):
                lineas.append(
                    (
                        f"Observaciones de {valor_o_guion(estancia['estancia_nombre'])}",
                        estancia["observaciones_estancia_inspeccion"],
                    )
                )

        if not lineas:
            continue

        hay_observaciones = True
        doc.add_heading(
            f"{numero_apartado}.{indice} Visita de fecha {valor_o_guion(visita['fecha'])}",
            level=2,
        )
        for etiqueta, valor in lineas:
            add_etiqueta_valor(doc, etiqueta, valor)

    if not hay_observaciones:
        add_parrafo(doc, "No constan observaciones técnicas adicionales registradas.")


def add_apartado_recomendaciones_inspeccion(
    doc: Document, numero_apartado: int, visitas_inspeccion
) -> None:
    doc.add_heading(f"{numero_apartado}. Recomendaciones", level=1)

    hay_recomendaciones = False
    for indice, bloque in enumerate(visitas_inspeccion, start=1):
        visita = bloque["visita"]
        general = bloque["inspeccion"]["general"]
        exterior = bloque["inspeccion"]["exterior"]
        comunes = bloque["inspeccion"]["comunes"]
        estancias = bloque["inspeccion"]["estancias"]
        incidencias = []

        for titulo, items in INSPECCION_GENERAL_GROUPS:
            for etiqueta, estado in recoger_incidencias_checklist(items, general):
                incidencias.append((f"General - {etiqueta}", estado))
        for etiqueta, estado in recoger_incidencias_checklist(
            INSPECCION_EXTERIOR_ITEMS, exterior
        ):
            incidencias.append((f"Exterior - {etiqueta}", estado))
        for etiqueta, estado in recoger_incidencias_checklist(
            INSPECCION_ELEMENTOS_COMUNES_ITEMS, comunes
        ):
            incidencias.append((f"Elementos comunes - {etiqueta}", estado))
        for estancia in estancias:
            for etiqueta, estado in recoger_incidencias_checklist(
                obtener_items_inspeccion_estancia(estancia["tipo_estancia"]),
                estancia,
            ):
                incidencias.append(
                    (
                        f"{valor_o_guion(estancia['estancia_nombre'])} - {etiqueta}",
                        estado,
                    )
                )

        if not incidencias:
            continue

        hay_recomendaciones = True
        doc.add_heading(
            f"{numero_apartado}.{indice} Visita de fecha {valor_o_guion(visita['fecha'])}",
            level=2,
        )
        add_parrafo(
            doc,
            "Se recomienda revisar y, en su caso, actuar sobre los siguientes elementos marcados en la inspección:",
        )
        for elemento, estado in incidencias:
            add_parrafo(doc, f"- {elemento}: {estado}")

    if not hay_recomendaciones:
        add_parrafo(
            doc,
            "No constan elementos marcados como necesitados de reparación o defecto grave en los checklists registrados.",
        )


def add_apartado_conclusion_inspeccion(
    doc: Document, numero_apartado: int, visitas_inspeccion
) -> None:
    doc.add_heading(f"{numero_apartado}. Conclusión general", level=1)

    total_reparacion = 0
    total_graves = 0
    total_no_inspeccionado = 0

    for bloque in visitas_inspeccion:
        general = bloque["inspeccion"]["general"]
        exterior = bloque["inspeccion"]["exterior"]
        comunes = bloque["inspeccion"]["comunes"]
        estancias = bloque["inspeccion"]["estancias"]
        conjuntos = [
            [campo for _, grupo in INSPECCION_GENERAL_GROUPS for campo, _ in grupo],
            [campo for campo, _ in INSPECCION_EXTERIOR_ITEMS],
            [campo for campo, _ in INSPECCION_ELEMENTOS_COMUNES_ITEMS],
        ]
        fuentes = [general, exterior, comunes]

        for fuente, campos in zip(fuentes, conjuntos):
            for campo in campos:
                estado = limpiar_texto(fuente[campo])
                if estado == "necesita_reparacion":
                    total_reparacion += 1
                elif estado == "defecto_grave":
                    total_graves += 1
                elif estado in {"", "no_inspeccionado"}:
                    total_no_inspeccionado += 1

        for estancia in estancias:
            for campo, _ in obtener_items_inspeccion_estancia(estancia["tipo_estancia"]):
                estado = limpiar_texto(estancia[campo])
                if estado == "necesita_reparacion":
                    total_reparacion += 1
                elif estado == "defecto_grave":
                    total_graves += 1
                elif estado in {"", "no_inspeccionado"}:
                    total_no_inspeccionado += 1

    if total_graves or total_reparacion:
        add_parrafo(
            doc,
            "La inspección incorpora elementos consignados con necesidad de reparación y/o defecto grave. "
            f"Se han registrado {total_reparacion} ítems como 'Necesita reparación' y {total_graves} como 'Defecto grave'.",
        )
    else:
        add_parrafo(
            doc,
            "De acuerdo con los checklists registrados, no constan elementos marcados como necesitados de reparación o defecto grave.",
        )

    if total_no_inspeccionado:
        add_parrafo(
            doc,
            f"Asimismo, constan {total_no_inspeccionado} ítems marcados como 'No inspeccionado', por lo que la valoración debe entenderse limitada a los elementos efectivamente revisados.",
        )


def add_apartado_reforma(doc: Document, expediente) -> None:
    doc.add_heading("3. Antecedentes de reforma", level=1)
    add_etiqueta_valor(doc, "Reformado", expediente["reformado"] or "No")
    add_etiqueta_valor(doc, "Fecha de reforma", expediente["fecha_reforma"])
    add_etiqueta_valor(
        doc, "Observaciones de la reforma", expediente["observaciones_reforma"]
    )


def add_apartado_visita(
    doc: Document, numero_apartado: int, visita, climatologia, estancias, patologias
) -> None:
    doc.add_heading(
        f"{numero_apartado}. Visita de inspección - {valor_o_guion(visita['fecha'])}",
        level=1,
    )

    add_etiqueta_valor(doc, "Técnico", visita["tecnico"])
    add_etiqueta_valor(doc, "Observaciones de visita", visita["observaciones_visita"])

    if climatologia:
        doc.add_heading(f"{numero_apartado}.1 Condiciones climatológicas", level=2)
        add_parrafo(doc, valor_o_guion(climatologia["resumen"]))
    else:
        doc.add_heading(f"{numero_apartado}.1 Condiciones climatológicas", level=2)
        add_parrafo(doc, "No consta climatología registrada para esta visita.")

    doc.add_heading(f"{numero_apartado}.2 Estancias inspeccionadas", level=2)

    if estancias:
        for estancia in estancias:
            add_parrafo(
                doc,
                f"- {valor_o_guion(estancia['nombre'])} | Tipo: {valor_o_guion(estancia['tipo_estancia'])} | "
                f"Planta: {valor_o_guion(estancia['planta'])}",
            )
    else:
        add_parrafo(doc, "No constan estancias registradas en esta visita.")

    doc.add_heading(f"{numero_apartado}.3 Patologías observadas", level=2)

    if not patologias:
        add_parrafo(doc, "No constan patologías registradas en esta visita.")
        return

    indice_patologia = 1
    for patologia in patologias:
        doc.add_heading(
            f"{numero_apartado}.3.{indice_patologia} {valor_o_guion(patologia['estancia_nombre'])}",
            level=3,
        )
        add_etiqueta_valor(doc, "Elemento afectado", patologia["elemento"])
        add_etiqueta_valor(doc, "Patología", patologia["patologia"])
        add_etiqueta_valor(doc, "Observaciones", patologia["observaciones"])

        if patologia["foto"]:
            add_parrafo(doc, "Fotografía asociada:", bold=True)
            add_imagen_si_existe(doc, patologia["foto"])

        indice_patologia += 1


def add_apartado_conclusion(doc: Document) -> None:
    doc.add_heading("Conclusión", level=1)
    add_parrafo(
        doc,
        "El presente documento recopila de forma ordenada la información disponible en el expediente, "
        "las visitas realizadas, las estancias inspeccionadas y las patologías registradas, sirviendo "
        "como base documental para su posterior análisis técnico pericial.",
    )


def generar_informe_patologias(cur, expediente, visitas, doc: Document) -> None:
    add_portada(doc, expediente)

    doc.add_heading("1. Identificación del expediente", level=1)
    add_tabla_datos_expediente(doc, expediente)

    doc.add_heading("2. Objeto del informe", level=1)
    add_parrafo(
        doc,
        "El presente informe pericial documenta las patologías observadas en el inmueble objeto del expediente, "
        "incorporando los antecedentes disponibles, las visitas realizadas y los registros técnicos interiores "
        "y/o exteriores existentes en el sistema.",
    )

    doc.add_heading("3. Descripción del inmueble", level=1)
    add_etiqueta_valor(
        doc, "Referencia catastral", expediente["referencia_catastral"]
    )
    add_etiqueta_valor(
        doc, "Observaciones generales", expediente["observaciones_generales"]
    )
    add_etiqueta_valor(doc, "Planta de la unidad", expediente["planta_unidad"])
    add_etiqueta_valor(doc, "Puerta / unidad", expediente["puerta_unidad"])
    add_etiqueta_valor(
        doc, "Superficie construida", expediente["superficie_construida"]
    )
    add_etiqueta_valor(doc, "Superficie útil", expediente["superficie_util"])
    add_etiqueta_valor(doc, "Dormitorios", expediente["dormitorios_unidad"])
    add_etiqueta_valor(doc, "Baños", expediente["banos_unidad"])
    add_etiqueta_valor(
        doc, "Observaciones de la unidad", expediente["observaciones_unidad"]
    )

    numero_apartado = 4
    add_apartado_datos_patologias(doc, numero_apartado, expediente)
    numero_apartado += 1

    add_apartado_bloque_judicial(doc, numero_apartado, expediente)
    if limpiar_texto(expediente["destinatario"]) == "judicial":
        numero_apartado += 1

    climatologias = {}
    patologias_interiores = []
    patologias_exteriores = []
    bloques_mapas = []

    for visita in visitas:
        climatologias[visita["id"]] = cur.execute(
            """
            SELECT *
            FROM climatologia_visitas
            WHERE visita_id = ?
            ORDER BY id DESC
            LIMIT 1
            """,
            (visita["id"],),
        ).fetchone()

        patologias_interiores.extend(
            cur.execute(
                """
                SELECT rp.*,
                       e.nombre AS estancia_nombre,
                       ue.identificador AS unidad_identificador,
                       ne.nombre_nivel AS nivel_nombre,
                       bp.rol_patologia AS rol_patologia_biblioteca
                FROM registros_patologias rp
                INNER JOIN estancias e ON rp.estancia_id = e.id
                LEFT JOIN unidades_expediente ue ON e.unidad_id = ue.id
                LEFT JOIN niveles_edificio ne ON ue.nivel_id = ne.id
                LEFT JOIN biblioteca_patologias bp
                       ON lower(trim(bp.nombre)) = lower(trim(rp.patologia))
                WHERE rp.visita_id = ?
                ORDER BY ne.nombre_nivel ASC, ue.identificador ASC, e.nombre ASC, rp.id ASC
                """,
                (visita["id"],),
            ).fetchall()
        )
        patologias_exteriores.extend(
            cur.execute(
                """
                SELECT rpe.*,
                       '' AS rol_patologia_observado,
                       bp.rol_patologia AS rol_patologia_biblioteca
                FROM registros_patologias_exteriores rpe
                LEFT JOIN biblioteca_patologias bp
                       ON lower(trim(bp.nombre)) = lower(trim(rpe.patologia))
                WHERE rpe.visita_id = ?
                ORDER BY rpe.zona_exterior ASC, rpe.id ASC
                """,
                (visita["id"],),
            ).fetchall()
        )
        bloques_mapas.append(
            {
                "visita": visita,
                "mapas": cargar_mapas_patologia_utiles_visita(cur, visita),
            }
        )

    add_apartado_inspeccion_visita(doc, numero_apartado, visitas, climatologias)
    numero_apartado += 1

    ambito = limpiar_texto(expediente["ambito_patologias"])
    if ambito in ("", "interior", "interior_exterior"):
        add_apartado_patologias_interiores(
            doc, numero_apartado, patologias_interiores, cur.connection
        )
        numero_apartado += 1

    if ambito in ("exterior", "interior_exterior"):
        add_apartado_patologias_exteriores(
            doc, numero_apartado, patologias_exteriores, cur.connection
        )
        numero_apartado += 1

    if add_apartado_mapas_patologia(doc, numero_apartado, bloques_mapas):
        numero_apartado += 1

    add_apartado_analisis_tecnico(doc, numero_apartado, expediente)
    numero_apartado += 1
    add_apartado_propuesta_reparacion(doc, numero_apartado, expediente)
    numero_apartado += 1
    add_apartado_conclusion_patologias(
        doc,
        numero_apartado,
        patologias_interiores,
        patologias_exteriores,
    )
    numero_apartado += 1
    add_apartado_conclusion_pericial(
        doc,
        numero_apartado,
        patologias_interiores,
        patologias_exteriores,
    )


def obtener_fotos_registro_informe(cur, tabla: str, fk_columna: str, parent_id: int, foto_principal: str | None = None) -> list[dict]:
    fotos = []
    vistos = set()
    foto_principal = limpiar_texto(foto_principal)
    if foto_principal:
        fotos.append({"archivo": foto_principal})
        vistos.add(foto_principal)

    for fila in cur.execute(
        f"""
        SELECT archivo
        FROM {tabla}
        WHERE {fk_columna}=?
        ORDER BY id ASC
        """,
        (parent_id,),
    ).fetchall():
        archivo = limpiar_texto(fila["archivo"])
        if archivo and archivo not in vistos:
            fotos.append({"archivo": archivo})
            vistos.add(archivo)
    return fotos


def obtener_fotos_estancia_informe(cur, estancia_id: int) -> list[dict]:
    return [
        {"archivo": fila["archivo"]}
        for fila in cur.execute(
            """
            SELECT archivo
            FROM estancia_fotos
            WHERE estancia_id=?
            ORDER BY id ASC
            """,
            (estancia_id,),
        ).fetchall()
        if limpiar_texto(fila["archivo"])
    ]


def construir_figura(contador: dict | None, archivo: str, pie_base: str, base_url: str) -> dict:
    return {
        "archivo": archivo,
        "url": imagen_url_pdf(archivo, base_url),
        "caption_base": pie_base,
        "caption": pie_base,
    }


def construir_campos_informe(filas) -> list[dict]:
    return [{"label": etiqueta, "value": valor_o_guion(valor)} for etiqueta, valor in filas]


def construir_toc_items_informe(contexto: dict, paginas: dict | None = None) -> list[dict]:
    paginas = paginas or {}
    offset = 1 if contexto.get("bloque_judicial") else 0
    post_mapas = 1 if contexto.get("mapas") else 0
    if contexto.get("es_valoracion"):
        items = [
            ("identificacion", 1, "Identificación del expediente"),
            ("objeto", 2, "Objeto del informe"),
            ("valoracion-eco", 3, "Criterios periciales de valoración"),
            ("valoracion-encargo", 4, "Encargo"),
            ("valoracion-documentacion", 5, "Documentación utilizada"),
            ("valoracion-identificacion", 6, "Identificación del bien"),
            ("valoracion-situacion_legal", 7, "Situación legal"),
            ("valoracion-entorno", 8, "Entorno"),
            ("valoracion-edificio_inmueble", 9, "Edificio/inmueble"),
            ("valoracion-constructivo", 10, "Características constructivas"),
            ("valoracion-estado", 11, "Estado"),
            ("valoracion-metodo", 12, "Método de valoración"),
            ("valoracion-comparables", 13, "Comparables"),
            ("valoracion-resultado", 14, "Resultado"),
            ("valoracion-limitaciones", 15, "Limitaciones"),
        ]
        return [
            {
                "id": item_id,
                "number": number,
                "title": title,
                "page": paginas.get(item_id),
            }
            for item_id, number, title in items
        ]
    items = [
        ("identificacion", 1, "Identificación del expediente"),
        ("objeto", 2, "Objeto del informe"),
        ("descripcion", 3, "Descripción del inmueble"),
        ("datos-patologias", 4, "Datos específicos del informe de patologías"),
    ]
    if contexto.get("bloque_judicial"):
        items.append(("judicial", 5, "Encargo judicial"))
    items.extend(
        [
            ("inspeccion", 5 + offset, "Inspección y visita"),
            ("patologias-interiores", 6 + offset, "Patologías interiores"),
            ("patologias-exteriores", 7 + offset, "Patologías exteriores"),
        ]
    )
    if contexto.get("mapas"):
        items.append(("mapas", 8 + offset, "Mapas de localización de daños"))
    items.extend(
        [
            ("analisis", 8 + offset + post_mapas, "Análisis técnico"),
            ("propuesta", 9 + offset + post_mapas, "Propuesta de reparación"),
        ]
    )
    post_anexo = 0
    if (contexto.get("anexo_economico_reparacion") or {}).get("incluido"):
        items.append(
            (
                "anexo-economico-reparacion",
                10 + offset + post_mapas,
                "Anexo económico de reparación",
            )
        )
        post_anexo = 1
    items.extend(
        [
            ("conclusiones-tecnicas", 10 + offset + post_mapas + post_anexo, "Conclusiones técnicas"),
            ("conclusiones-periciales", 11 + offset + post_mapas + post_anexo, "Conclusiones periciales"),
        ]
    )
    return [
        {
            "id": item_id,
            "number": number,
            "title": title,
            "page": paginas.get(item_id),
        }
        for item_id, number, title in items
    ]


def asignar_numero_figura(figura: dict | None, contador: dict) -> None:
    if not figura:
        return
    numero = contador["valor"]
    contador["valor"] += 1
    figura["numero"] = numero
    figura["caption"] = f"Figura {numero}. {figura.get('caption_base') or ''}".strip()


def numerar_figuras_en_orden_visual(contexto: dict) -> int:
    contador = {"valor": 1}

    for visita in contexto.get("visitas", []):
        for foto in visita.get("fotos_exteriores", []):
            asignar_numero_figura(foto.get("figura"), contador)

    for estancia in contexto.get("estancias", []):
        for figura in estancia.get("fotos", []):
            asignar_numero_figura(figura, contador)
        for patologia in estancia.get("patologias", []):
            for figura in patologia.get("fotos", []):
                asignar_numero_figura(figura, contador)

    for zona in contexto.get("patologias_exteriores", []):
        for patologia in zona.get("patologias", []):
            for figura in patologia.get("fotos", []):
                asignar_numero_figura(figura, contador)

    for bloque in contexto.get("mapas", []):
        for mapa in bloque.get("mapas", []):
            for cuadrante in mapa.get("cuadrantes", []):
                asignar_numero_figura(cuadrante.get("foto_figura"), contador)

    return contador["valor"] - 1


def cargar_anexo_economico_reparacion_informe(cur, expediente_id: int, incluir: bool) -> dict:
    anexo = {
        "incluido": False,
        "tiene_costes": False,
        "modo": "",
        "actuaciones": [],
        "patologias": [],
        "total_pem": 0.0,
        "fuentes": [],
        "nota": (
            "La valoración económica siguiente se realiza a efectos orientativos/periciales, "
            "basada en mediciones y precios unitarios incorporados al expediente."
        ),
    }
    filas_actuaciones = cur.execute(
        """
        SELECT
            ar.id AS actuacion_id,
            ar.titulo,
            ar.descripcion AS actuacion_descripcion,
            ar.observaciones AS actuacion_observaciones,
            ar.orden,
            ap.descripcion_snapshot,
            ap.unidad_snapshot,
            ap.precio_unitario_snapshot,
            ap.cantidad,
            ap.importe,
            cc.codigo AS concepto_codigo,
            cc.resumen AS concepto_resumen,
            cb.nombre AS base_nombre,
            cb.origen AS base_origen,
            cb.version AS base_version
        FROM actuaciones_reparacion ar
        JOIN actuacion_partidas ap ON ap.actuacion_id = ar.id
        JOIN costes_conceptos cc ON cc.id = ap.concepto_id
        JOIN costes_bases cb ON cb.id = cc.base_id
        WHERE ar.expediente_id = ?
        ORDER BY ar.orden ASC, ar.id ASC, ap.id ASC
        """,
        (expediente_id,),
    ).fetchall()

    if filas_actuaciones:
        anexo["tiene_costes"] = True
        anexo["modo"] = "actuaciones"
        anexo["nota"] = (
            "La valoración económica se estructura por actuaciones de reparación "
            "necesarias para la subsanación de los daños observados."
        )
        if not incluir:
            return anexo

        actuaciones_por_id: dict[int, dict] = {}
        fuentes: dict[str, dict] = {}
        total = 0.0
        for fila in filas_actuaciones:
            actuacion = actuaciones_por_id.setdefault(
                fila["actuacion_id"],
                {
                    "id": fila["actuacion_id"],
                    "titulo": valor_o_guion(fila["titulo"]),
                    "descripcion": limpiar_texto(fila["actuacion_descripcion"]),
                    "observaciones": limpiar_texto(fila["actuacion_observaciones"]),
                    "partidas": [],
                    "subtotal": 0.0,
                },
            )
            importe = round(float(fila["importe"] or 0), 2)
            actuacion["partidas"].append(
                {
                    "codigo": limpiar_texto(fila["concepto_codigo"]),
                    "partida": limpiar_texto(fila["descripcion_snapshot"])
                    or limpiar_texto(fila["concepto_resumen"]),
                    "cantidad": float(fila["cantidad"] or 0),
                    "unidad": limpiar_texto(fila["unidad_snapshot"]),
                    "precio_unitario": float(fila["precio_unitario_snapshot"] or 0),
                    "importe": importe,
                    "base_nombre": limpiar_texto(fila["base_nombre"]),
                    "base_origen": limpiar_texto(fila["base_origen"]),
                }
            )
            actuacion["subtotal"] = round(actuacion["subtotal"] + importe, 2)
            total = round(total + importe, 2)
            clave_fuente = f"{fila['base_nombre']}|{fila['base_origen']}|{fila['base_version']}"
            fuentes[clave_fuente] = {
                "nombre": limpiar_texto(fila["base_nombre"]),
                "origen": limpiar_texto(fila["base_origen"]),
                "version": limpiar_texto(fila["base_version"]),
            }

        anexo["incluido"] = True
        anexo["actuaciones"] = list(actuaciones_por_id.values())
        anexo["total_pem"] = total
        anexo["fuentes"] = list(fuentes.values())
        return anexo

    filas_patologias = cur.execute(
        """
        SELECT
            pc.patologia_id,
            pc.descripcion_actuacion,
            pc.cantidad,
            pc.unidad,
            pc.precio_unitario,
            pc.importe,
            pc.estado AS vinculo_estado,
            pc.observaciones AS vinculo_observaciones,
            rp.patologia,
            rp.elemento,
            rp.localizacion_dano,
            rp.detalle_localizacion,
            es.nombre AS estancia_nombre,
            cc.codigo AS concepto_codigo,
            cc.resumen AS concepto_resumen,
            cc.estado AS concepto_estado,
            cb.nombre AS base_nombre,
            cb.origen AS base_origen,
            cb.version AS base_version
        FROM patologia_costes pc
        JOIN registros_patologias rp ON rp.id = pc.patologia_id
        JOIN visitas v ON v.id = rp.visita_id
        LEFT JOIN estancias es ON es.id = rp.estancia_id
        JOIN costes_conceptos cc ON cc.id = pc.concepto_id
        JOIN costes_bases cb ON cb.id = cc.base_id
        WHERE v.expediente_id = ?
        ORDER BY es.nombre COLLATE NOCASE, rp.id, pc.id
        """,
        (expediente_id,),
    ).fetchall()
    if not filas_patologias:
        return anexo

    anexo["tiene_costes"] = True
    anexo["modo"] = "patologias"
    if not incluir:
        return anexo

    patologias_por_id: dict[int, dict] = {}
    fuentes: dict[str, dict] = {}
    total = 0.0
    for fila in filas_patologias:
        patologia = patologias_por_id.setdefault(
            fila["patologia_id"],
            {
                "id": fila["patologia_id"],
                "patologia": valor_o_guion(fila["patologia"]),
                "zona": " · ".join(
                    parte
                    for parte in [
                        limpiar_texto(fila["estancia_nombre"]),
                        limpiar_texto(fila["elemento"]),
                        limpiar_texto(fila["localizacion_dano"]),
                    ]
                    if parte
                ),
                "detalle_localizacion": limpiar_texto(fila["detalle_localizacion"]),
                "partidas": [],
                "subtotal": 0.0,
            },
        )
        importe = round(float(fila["importe"] or 0), 2)
        patologia["partidas"].append(
            {
                "codigo": limpiar_texto(fila["concepto_codigo"]),
                "partida": limpiar_texto(fila["descripcion_actuacion"]) or limpiar_texto(fila["concepto_resumen"]),
                "cantidad": float(fila["cantidad"] or 0),
                "unidad": limpiar_texto(fila["unidad"]),
                "precio_unitario": float(fila["precio_unitario"] or 0),
                "importe": importe,
                "estado": limpiar_texto(fila["vinculo_estado"]),
                "observaciones": limpiar_texto(fila["vinculo_observaciones"]),
                "base_nombre": limpiar_texto(fila["base_nombre"]),
                "base_origen": limpiar_texto(fila["base_origen"]),
            }
        )
        patologia["subtotal"] = round(patologia["subtotal"] + importe, 2)
        total = round(total + importe, 2)
        clave_fuente = f"{fila['base_nombre']}|{fila['base_origen']}|{fila['base_version']}"
        fuentes[clave_fuente] = {
            "nombre": limpiar_texto(fila["base_nombre"]),
            "origen": limpiar_texto(fila["base_origen"]),
            "version": limpiar_texto(fila["base_version"]),
        }

    anexo["incluido"] = True
    anexo["patologias"] = list(patologias_por_id.values())
    anexo["total_pem"] = total
    anexo["fuentes"] = list(fuentes.values())
    return anexo


def build_informe_context(
    expediente_id: int,
    base_url: str = "",
    incluir_anexo_economico_reparacion: bool = False,
) -> dict:
    """Prepara los datos del informe HTML sin alterar la generación DOCX existente."""
    conn = get_connection()
    cur = conn.cursor()
    try:
        expediente = cur.execute(
            "SELECT * FROM expedientes WHERE id=?",
            (expediente_id,),
        ).fetchone()
        if not expediente:
            raise ValueError("Expediente no encontrado")

        tipo_informe = limpiar_texto(expediente["tipo_informe"])

        visitas = cur.execute(
            """
            SELECT *
            FROM visitas
            WHERE expediente_id=?
            ORDER BY id ASC
            """,
            (expediente_id,),
        ).fetchall()

        ambitos = {
            "interior": "Interior",
            "exterior": "Exterior",
            "interior_exterior": "Interior y exterior",
        }
        zonas_exteriores = {
            "fachada": "Fachada",
            "cubierta": "Cubierta",
            "medianera": "Medianera",
            "patio": "Patio",
            "terraza": "Terraza",
            "exterior_general": "Exterior general",
        }
        elementos_exteriores = {
            "revestimiento": "Revestimiento",
            "cerramiento": "Cerramiento",
            "cornisa": "Cornisa",
            "alero": "Alero",
            "peto": "Peto",
            "impermeabilizacion": "Impermeabilización",
            "carpinteria_exterior": "Carpintería exterior",
            "barandilla": "Barandilla",
            "bajante": "Bajante",
            "canalon": "Canalón",
            "forjado": "Forjado",
            "otro": "Otro",
        }
        localizaciones_exteriores = {
            "horizontal": "Horizontal",
            "vertical": "Vertical",
            "encuentro": "Encuentro",
            "puntual": "Puntual",
        }

        patologias_interiores_rows = []
        patologias_exteriores_rows = []
        visitas_contexto = []
        estancias_por_id: dict[int, dict] = {}
        bloques_mapas = []
        valoracion_contexto = []
        comparables_valoracion_contexto = []
        valoracion_eco_contexto = {}

        for visita in visitas:
            climatologia = cur.execute(
                """
                SELECT *
                FROM climatologia_visitas
                WHERE visita_id=?
                ORDER BY id DESC
                LIMIT 1
                """,
                (visita["id"],),
            ).fetchone()
            fotos_exteriores_visita = [
                {
                    "descripcion": foto["descripcion"],
                    "figura": construir_figura(
                        None,
                        foto["ruta"],
                        f"Fotografía exterior descriptiva ({valor_o_guion(visita['fecha'])})",
                        base_url,
                    ),
                }
                for foto in cur.execute(
                    """
                    SELECT ruta, descripcion
                    FROM visita_fotos
                    WHERE visita_id=? AND categoria='exterior'
                    ORDER BY id ASC
                    """,
                    (visita["id"],),
                ).fetchall()
            ]
            visitas_contexto.append(
                {
                    "id": visita["id"],
                    "fecha": visita["fecha"],
                    "tecnico": visita["tecnico"],
                    "observaciones": visita["observaciones_visita"],
                    "objeto": describir_objeto_visita_informe(cur, visita),
                    "climatologia": get_row_value(climatologia, "resumen", "No consta climatología registrada"),
                    "fotos_exteriores": fotos_exteriores_visita,
                }
            )
            for estancia in cur.execute(
                """
                SELECT e.*,
                       ue.identificador AS unidad_identificador,
                       ne.nombre_nivel AS nivel_nombre
                FROM estancias e
                LEFT JOIN unidades_expediente ue ON e.unidad_id = ue.id
                LEFT JOIN niveles_edificio ne ON ue.nivel_id = ne.id
                WHERE e.visita_id=?
                ORDER BY ne.nombre_nivel ASC, ue.identificador ASC, e.id ASC
                """,
                (visita["id"],),
            ).fetchall():
                fotos_estancia = []
                for foto in obtener_fotos_estancia_informe(cur, estancia["id"]):
                    fotos_estancia.append(
                        construir_figura(
                            None,
                            foto["archivo"],
                            f"Vista general de {valor_o_guion(estancia['nombre'])}",
                            base_url,
                        )
                    )
                estancias_por_id[estancia["id"]] = {
                    "id": estancia["id"],
                    "visita_id": visita["id"],
                    "nombre": limpiar_texto(estancia["nombre"]) or "Sin estancia",
                    "nivel": limpiar_texto(estancia["nivel_nombre"]),
                    "unidad": limpiar_texto(estancia["unidad_identificador"]),
                    "campos": construir_campos_informe(
                        [
                            ("Tipo de estancia", estancia["tipo_estancia"]),
                            ("Planta", estancia["planta"]),
                            ("Ventilación", estancia["ventilacion"]),
                            ("Acabado de pavimento", estancia["acabado_pavimento"]),
                            ("Acabado de paramento", estancia["acabado_paramento"]),
                            ("Acabado de techo", estancia["acabado_techo"]),
                            ("Observaciones técnicas", estancia["observaciones"]),
                        ]
                    ),
                    "fotos": fotos_estancia,
                    "patologias": [],
                }

            patologias_interiores_rows.extend(
                cur.execute(
                    """
                    SELECT rp.*,
                           e.nombre AS estancia_nombre,
                           ue.identificador AS unidad_identificador,
                           ne.nombre_nivel AS nivel_nombre,
                           bp.rol_patologia AS rol_patologia_biblioteca
                    FROM registros_patologias rp
                    INNER JOIN estancias e ON rp.estancia_id = e.id
                    LEFT JOIN unidades_expediente ue ON e.unidad_id = ue.id
                    LEFT JOIN niveles_edificio ne ON ue.nivel_id = ne.id
                    LEFT JOIN biblioteca_patologias bp
                           ON lower(trim(bp.nombre)) = lower(trim(rp.patologia))
                    WHERE rp.visita_id=?
                    ORDER BY ne.nombre_nivel ASC, ue.identificador ASC, e.nombre ASC, rp.id ASC
                    """,
                    (visita["id"],),
                ).fetchall()
            )
            patologias_exteriores_rows.extend(
                cur.execute(
                    """
                    SELECT rpe.*,
                           '' AS rol_patologia_observado,
                           bp.rol_patologia AS rol_patologia_biblioteca
                    FROM registros_patologias_exteriores rpe
                    LEFT JOIN biblioteca_patologias bp
                           ON lower(trim(bp.nombre)) = lower(trim(rpe.patologia))
                    WHERE rpe.visita_id=?
                    ORDER BY rpe.zona_exterior ASC, rpe.id ASC
                    """,
                    (visita["id"],),
                ).fetchall()
            )
            bloques_mapas.append(
                {
                    "visita": visita,
                    "mapas": cargar_mapas_patologia_utiles_visita(cur, visita),
                }
            )

        if tipo_informe == "valoracion":
            valoracion_contexto = cargar_valoracion_expediente_con_fallback(
                cur,
                expediente_id,
                visitas,
            )
            comparables_valoracion_contexto = cargar_comparables_valoracion_con_fallback(
                cur,
                expediente_id,
                visitas,
            )
            datos_valoracion_eco = cargar_valoracion_eco_con_fallback(
                cur,
                expediente_id,
                visitas,
            )
            valoracion_eco_contexto = construir_valoracion_eco(
                row_to_dict(expediente),
                datos_valoracion_eco,
                comparables_valoracion_contexto,
            )
            resumen_comparacion_valoracion = construir_resumen_comparacion_contexto(
                comparables_valoracion_contexto
            )
            valoracion_eco_contexto["resumen_comparacion"] = resumen_comparacion_valoracion
            valoracion_eco_contexto["comparacion"] = {
                "resumen": resumen_comparacion_valoracion
            }
            valoracion_contexto = ValoracionContext(
                valoracion_contexto,
                valoracion_eco_contexto,
            )
        else:
            resumen_comparacion_valoracion = construir_resumen_comparacion_contexto([])

        for registro in patologias_interiores_rows:
            referencias_cuadrantes = obtener_referencias_cuadrantes_patologia(conn, registro["id"])
            fotos = []
            for foto in obtener_fotos_registro_informe(
                cur,
                "registro_patologia_fotos",
                "registro_id",
                registro["id"],
                registro["foto"],
            ):
                pie_base = (
                    f"{valor_o_guion(registro['patologia'])} en "
                    f"{valor_o_guion(registro['localizacion_dano']).lower()} "
                    f"({valor_o_guion(registro['estancia_nombre'])})"
                )
                fotos.append(construir_figura(None, foto["archivo"], pie_base, base_url))
            patologia = {
                "id": registro["id"],
                "titulo": valor_o_guion(registro["patologia"]),
                "campos": construir_campos_informe(
                    [
                        ("Cuadrantes", formatear_referencias_cuadrantes_patologia(referencias_cuadrantes) if referencias_cuadrantes else ""),
                        ("Localización del daño", registro["localizacion_dano"]),
                        ("Detalle de localización", registro["detalle_localizacion"]),
                        ("Elemento", registro["elemento"]),
                        ("Patología", registro["patologia"]),
                        ("Rol técnico", registro["rol_patologia_observado"] or registro["rol_patologia_biblioteca"]),
                        ("Observaciones", registro["observaciones"]),
                    ]
                ),
                "fotos": fotos,
            }
            if registro["estancia_id"] not in estancias_por_id:
                estancias_por_id[registro["estancia_id"]] = {
                    "id": registro["estancia_id"],
                    "visita_id": registro["visita_id"],
                    "nombre": limpiar_texto(registro["estancia_nombre"]) or "Sin estancia",
                    "nivel": limpiar_texto(registro["nivel_nombre"]),
                    "unidad": limpiar_texto(registro["unidad_identificador"]),
                    "campos": [],
                    "fotos": [],
                    "patologias": [],
                }
            estancias_por_id[registro["estancia_id"]]["patologias"].append(patologia)

        estancias = list(estancias_por_id.values())
        for estancia in estancias:
            prefijo = f"En la estancia {estancia['nombre']}"
            registros_estancia = [
                registro
                for registro in patologias_interiores_rows
                if registro["estancia_id"] == estancia["id"]
            ]
            estancia["resumen"] = construir_resumen_patologias(prefijo, registros_estancia)
            estancia["incoherencias"] = detectar_incoherencias(registros_estancia)

        patologias_exteriores = []
        for zona, items in agrupar_patologias_exteriores(patologias_exteriores_rows).items():
            zona_legible = zonas_exteriores.get(zona, zona)
            patologias_zona = []
            for registro in items:
                referencias_cuadrantes = obtener_referencias_cuadrantes_patologia(conn, registro["id"])
                elemento_exterior = elementos_exteriores.get(
                    limpiar_texto(registro["elemento_exterior"]),
                    registro["elemento_exterior"],
                )
                localizacion_exterior = localizaciones_exteriores.get(
                    limpiar_texto(registro["localizacion_dano_exterior"]),
                    registro["localizacion_dano_exterior"],
                )
                fotos = []
                for foto in obtener_fotos_registro_informe(
                    cur,
                    "registro_patologia_exterior_fotos",
                    "registro_id",
                    registro["id"],
                    registro["foto"],
                ):
                    pie_base = (
                        f"{valor_o_guion(registro['patologia'])} en "
                        f"{valor_o_guion(localizacion_exterior).lower()} "
                        f"({valor_o_guion(zona_legible)})"
                    )
                    fotos.append(construir_figura(None, foto["archivo"], pie_base, base_url))
                patologias_zona.append(
                    {
                        "id": registro["id"],
                        "titulo": valor_o_guion(registro["patologia"]),
                        "campos": construir_campos_informe(
                            [
                                ("Cuadrantes", formatear_referencias_cuadrantes_patologia(referencias_cuadrantes) if referencias_cuadrantes else ""),
                                ("Zona exterior", zona_legible),
                                ("Elemento exterior", elemento_exterior),
                                ("Localización del daño exterior", localizacion_exterior),
                                ("Patología", registro["patologia"]),
                                ("Rol técnico", registro["rol_patologia_observado"] or registro["rol_patologia_biblioteca"]),
                                ("Observaciones", registro["observaciones"]),
                            ]
                        ),
                        "fotos": fotos,
                    }
                )
            patologias_exteriores.append(
                {
                    "zona": valor_o_guion(zona_legible),
                    "resumen": construir_resumen_patologias("En el exterior del inmueble", items),
                    "incoherencias": detectar_incoherencias(items),
                    "patologias": patologias_zona,
                }
            )

        mapas_contexto = []
        for bloque in bloques_mapas:
            mapas = []
            for mapa in bloque["mapas"]:
                cuadrantes = []
                for cuadrante in mapa["cuadrantes"]:
                    foto_detalle = ""
                    if limpiar_texto(cuadrante["foto_detalle"]):
                        foto_detalle = construir_figura(
                            None,
                            cuadrante["foto_detalle"],
                            f"Detalle del cuadrante {valor_o_guion(cuadrante['codigo_cuadrante'])}",
                            base_url,
                        )
                    cuadrantes.append({**cuadrante, "foto_figura": foto_detalle})
                mapas.append(
                    {
                        **mapa,
                        "imagen_base_url": imagen_url_pdf(mapa["imagen_base"], base_url),
                        "imagen_overlay_url": mapa_overlay_data_uri({**mapa, "cuadrantes": cuadrantes}),
                        "cuadrantes": cuadrantes,
                    }
                )
            if mapas:
                mapas_contexto.append(
                    {
                        "visita": row_to_dict(bloque["visita"]),
                        "mapas": mapas,
                    }
                )

        expediente_dict = row_to_dict(expediente)
        tipo_trabajo = {
            "patologias": "Informe de patologías",
            "inspeccion": "Informe de inspección técnica",
            "habitabilidad": "Informe de habitabilidad",
            "valoracion": "Informe de valoración",
        }.get(tipo_informe, "Informe pericial")
        anexo_economico_reparacion = cargar_anexo_economico_reparacion_informe(
            cur,
            expediente_id,
            incluir_anexo_economico_reparacion and tipo_informe == "patologias",
        )

        contexto = {
            "expediente": expediente_dict,
            "tipo_informe": tipo_informe,
            "es_valoracion": tipo_informe == "valoracion",
            "fecha_emision": datetime.now().strftime("%d/%m/%Y"),
            "portada": {
                "tecnico": visitas_contexto[0]["tecnico"] if visitas_contexto else "-",
                "tipo_trabajo": tipo_trabajo,
                "direccion": valor_o_guion(expediente["direccion"]),
                "fecha": datetime.now().strftime("%d/%m/%Y"),
            },
            "identificacion": construir_campos_informe(
                [
                    ("Número de expediente", expediente["numero_expediente"]),
                    ("Cliente", expediente["cliente"]),
                    ("Dirección", expediente["direccion"]),
                    ("Código postal", expediente["codigo_postal"]),
                    ("Ciudad", expediente["ciudad"]),
                    ("Provincia", expediente["provincia"]),
                    ("Tipo de inmueble", expediente["tipo_inmueble"]),
                    ("Orientación", expediente["orientacion_inmueble"]),
                    ("Año de construcción", expediente["anio_construccion"]),
                    ("Uso del inmueble", expediente["uso_inmueble"]),
                ]
            ),
            "descripcion_inmueble": construir_campos_informe(
                [
                    ("Referencia catastral", expediente["referencia_catastral"]),
                    ("Observaciones generales", expediente["observaciones_generales"]),
                    ("Observaciones del edificio", expediente["observaciones_bloque"]),
                    ("Planta de la unidad", expediente["planta_unidad"]),
                    ("Puerta / unidad", expediente["puerta_unidad"]),
                    ("Superficie construida", expediente["superficie_construida"]),
                    ("Superficie útil", expediente["superficie_util"]),
                    ("Dormitorios", expediente["dormitorios_unidad"]),
                    ("Baños", expediente["banos_unidad"]),
                    ("Observaciones de la unidad", expediente["observaciones_unidad"]),
                ]
            ),
            "datos_patologias": construir_campos_informe(
                [
                    ("Ámbito de patologías", ambitos.get(limpiar_texto(expediente["ambito_patologias"]), expediente["ambito_patologias"])),
                    ("Descripción del daño", expediente["descripcion_dano"]),
                    ("Causa probable", expediente["causa_probable"]),
                    ("Pruebas e indicios", expediente["pruebas_indicios"]),
                    ("Evolución / preexistencia", expediente["evolucion_preexistencia"]),
                    ("Propuesta de reparación", expediente["propuesta_reparacion"]),
                    ("Urgencia / gravedad", expediente["urgencia_gravedad"]),
                ]
            ),
            "bloque_judicial": construir_campos_informe(
                [
                    ("Procedimiento judicial", expediente["procedimiento_judicial"]),
                    ("Juzgado", expediente["juzgado"]),
                    ("Auto judicial", expediente["auto_judicial"]),
                    ("Parte solicitante", expediente["parte_solicitante"]),
                    ("Objeto de la pericia", expediente["objeto_pericia"]),
                    ("Alcance y limitaciones", expediente["alcance_limitaciones"]),
                    ("Metodología pericial", expediente["metodologia_pericial"]),
                ]
            )
            if limpiar_texto(expediente["destinatario"]) == "judicial"
            else [],
            "visitas": visitas_contexto,
            "valoracion": valoracion_contexto,
            "valoracion_eco": valoracion_eco_contexto,
            "comparables_valoracion": comparables_valoracion_contexto,
            "resumen_comparacion_valoracion": resumen_comparacion_valoracion,
            "completitud_valoracion": construir_completitud_valoracion(
                valoracion_contexto,
                comparables_valoracion_contexto,
            )
            if tipo_informe == "valoracion"
            else {"advertencias": [], "completo": True, "total_advertencias": 0},
            "estancias": estancias,
            "patologias_exteriores": patologias_exteriores,
            "mapas": mapas_contexto,
            "analisis_tecnico": construir_campos_informe(
                [
                    ("Descripción del daño", expediente["descripcion_dano"]),
                    ("Causa probable", expediente["causa_probable"]),
                    ("Pruebas e indicios", expediente["pruebas_indicios"]),
                    ("Evolución / preexistencia", expediente["evolucion_preexistencia"]),
                    ("Urgencia / gravedad", expediente["urgencia_gravedad"]),
                ]
            ),
            "propuesta_reparacion": valor_o_guion(expediente["propuesta_reparacion"])
            if limpiar_texto(expediente["propuesta_reparacion"])
            else "No consta propuesta de reparación registrada.",
            "conclusion_tecnica": construir_conclusion_tecnica_global(
                patologias_interiores_rows,
                patologias_exteriores_rows,
            ),
            "conclusion_pericial": construir_conclusion_pericial(
                patologias_interiores_rows,
                patologias_exteriores_rows,
            ),
            "anexo_economico_reparacion": anexo_economico_reparacion,
        }
        contexto["total_figuras"] = numerar_figuras_en_orden_visual(contexto)
        contexto["toc_items"] = construir_toc_items_informe(contexto)
        return contexto
    finally:
        conn.close()


def nombre_archivo_pdf_informe(expediente) -> str:
    numero = limpiar_nombre_archivo(get_row_value(expediente, "numero_expediente", "expediente"))
    return f"Informe-{numero}.pdf"


def nombre_archivo_docx_editable_informe(expediente) -> str:
    numero = limpiar_nombre_archivo(get_row_value(expediente, "numero_expediente", "expediente"))
    return f"Informe-editable-{numero}.docx"


def get_or_create_paragraph_style(doc: Document, name: str):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def configurar_estilo_parrafo(
    doc: Document,
    name: str,
    *,
    size: int,
    bold: bool = False,
    italic: bool = False,
    color: str = "1F2933",
    alignment=None,
):
    style = get_or_create_paragraph_style(doc, name)
    style.font.name = "Arial"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = RGBColor.from_string(color)
    if alignment is not None:
        style.paragraph_format.alignment = alignment
    return style


def configurar_documento_editable(doc: Document) -> None:
    configurar_documento(doc)
    seccion = doc.sections[-1]
    seccion.top_margin = Cm(1.5)
    seccion.bottom_margin = Cm(1.5)
    seccion.left_margin = Cm(1.7)
    seccion.right_margin = Cm(1.7)

    estilo_normal = doc.styles["Normal"]
    estilo_normal.font.name = "Arial"
    estilo_normal.font.size = Pt(10)

    configurar_estilo_parrafo(
        doc,
        "Informe Portada Titulo",
        size=26,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    configurar_estilo_parrafo(
        doc,
        "Informe Portada Subtitulo",
        size=12,
        italic=True,
        color="4B5563",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    configurar_estilo_parrafo(doc, "Informe Caption", size=8, italic=True, color="4B5563", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    configurar_estilo_parrafo(doc, "Informe Observacion", size=9, color="1F2933")
    configurar_estilo_parrafo(doc, "Informe Card Titulo", size=11, bold=True)


def ajustar_runs_parrafo(parrafo, font_size: int | None = None, bold: bool | None = None) -> None:
    for run in parrafo.runs:
        run.font.name = "Arial"
        if font_size:
            run.font.size = Pt(font_size)
        if bold is not None:
            run.bold = bold


def add_heading_editable(doc: Document, texto: str, level: int = 1) -> None:
    parrafo = doc.add_heading(valor_o_guion(texto), level=level)
    ajustar_runs_parrafo(parrafo, font_size=16 if level == 1 else 13 if level == 2 else 11)
    parrafo.paragraph_format.keep_with_next = True


def aplicar_bordes_celda(celda, color: str = "D8DEE6", size: str = "6") -> None:
    tc_pr = celda._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def aplicar_fila_no_dividir(fila) -> None:
    tr_pr = fila._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def aplicar_sombreado_celda(celda, fill: str) -> None:
    tc_pr = celda._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def aplicar_margenes_celda(celda, margen: int = 120) -> None:
    tc_pr = celda._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for lado in ("top", "left", "bottom", "right"):
        nodo = tc_mar.find(qn(f"w:{lado}"))
        if nodo is None:
            nodo = OxmlElement(f"w:{lado}")
            tc_mar.append(nodo)
        nodo.set(qn("w:w"), str(margen))
        nodo.set(qn("w:type"), "dxa")


def aplicar_estilo_tabla_suave(tabla, sombrear_primera_columna: bool = False) -> None:
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    for fila in tabla.rows:
        aplicar_fila_no_dividir(fila)
        for indice, celda in enumerate(fila.cells):
            celda.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            aplicar_bordes_celda(celda)
            aplicar_margenes_celda(celda)
            if sombrear_primera_columna and indice == 0:
                aplicar_sombreado_celda(celda, "F3F6F8")
            for parrafo in celda.paragraphs:
                parrafo.paragraph_format.space_after = Pt(0)
                for run in parrafo.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
                    if sombrear_primera_columna and indice == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor.from_string("6B7280")


def add_parrafo_editable(
    doc: Document,
    texto: str,
    *,
    bold: bool = False,
    italic: bool = False,
    centrado: bool = False,
    espacio_despues: int = 6,
) -> None:
    parrafo = doc.add_paragraph()
    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER if centrado else WD_ALIGN_PARAGRAPH.LEFT
    parrafo.paragraph_format.space_after = Pt(espacio_despues)
    run = parrafo.add_run(valor_o_guion(texto))
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.bold = bold
    run.italic = italic


def add_tabla_campos_editable(doc: Document, campos: list[dict]) -> None:
    tabla = doc.add_table(rows=0, cols=2)

    for campo in campos:
        row = tabla.add_row().cells
        row[0].text = valor_o_guion(campo.get("label")).upper()
        row[1].text = valor_o_guion(campo.get("value"))

    aplicar_estilo_tabla_suave(tabla, sombrear_primera_columna=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def obtener_imagen_origen_docx(figura: dict):
    if not figura:
        return None

    archivo = limpiar_texto(figura.get("archivo"))
    if archivo:
        ruta_foto = os.path.join(UPLOAD_DIR, archivo)
        if os.path.exists(ruta_foto):
            return ruta_foto

    data_uri = limpiar_texto(figura.get("data_uri"))
    if data_uri.startswith("data:image/") and "," in data_uri:
        try:
            return BytesIO(base64.b64decode(data_uri.split(",", 1)[1]))
        except Exception:
            return None

    return None


def add_imagen_en_parrafo(parrafo, figura: dict, ancho_cm: float) -> bool:
    imagen_origen = obtener_imagen_origen_docx(figura)
    if imagen_origen is None:
        return False
    run = parrafo.add_run()
    run.add_picture(imagen_origen, width=Cm(ancho_cm))
    return True


def add_imagen_editable(doc: Document, figura: dict, ancho_cm: float = 11.5) -> None:
    if not figura:
        return

    imagen_origen = obtener_imagen_origen_docx(figura)
    archivo = limpiar_texto(figura.get("archivo"))
    if imagen_origen is None:
        add_parrafo_editable(doc, f"Imagen no localizada: {archivo or '-'}")
        return

    try:
        doc.add_picture(imagen_origen, width=Cm(ancho_cm))
        ultimo = doc.paragraphs[-1]
        ultimo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ultimo.paragraph_format.space_after = Pt(2)
        if limpiar_texto(figura.get("caption")):
            parrafo_pie = doc.add_paragraph(style="Informe Caption")
            parrafo_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
            parrafo_pie.paragraph_format.space_after = Pt(10)
            run = parrafo_pie.add_run(figura["caption"])
            run.font.name = "Arial"
            run.font.size = Pt(9)
            run.italic = True
    except Exception:
        add_parrafo_editable(doc, f"No se ha podido insertar la imagen: {archivo or '-'}")


def add_bloque_figuras_editable(doc: Document, figuras: list[dict], ancho_cm: float = 11.5) -> None:
    for figura in figuras or []:
        add_imagen_editable(doc, figura, ancho_cm=ancho_cm)


def add_grid_figuras_editable(doc: Document, figuras: list[dict], ancho_cm: float = 7.5) -> None:
    figuras = [figura for figura in (figuras or []) if figura]
    if not figuras:
        return

    tabla = doc.add_table(rows=0, cols=2)
    for indice in range(0, len(figuras), 2):
        row = tabla.add_row()
        aplicar_fila_no_dividir(row)
        cells = row.cells
        for offset, celda in enumerate(cells):
            posicion = indice + offset
            aplicar_bordes_celda(celda, color="FFFFFF", size="0")
            aplicar_margenes_celda(celda, margen=80)
            if posicion >= len(figuras):
                continue
            figura = figuras[posicion]
            parrafo_imagen = celda.paragraphs[0]
            parrafo_imagen.alignment = WD_ALIGN_PARAGRAPH.CENTER
            parrafo_imagen.paragraph_format.keep_with_next = True
            if not add_imagen_en_parrafo(parrafo_imagen, figura, ancho_cm):
                parrafo_imagen.add_run("Imagen no localizada")
            parrafo_pie = celda.add_paragraph()
            parrafo_pie.style = "Informe Caption"
            parrafo_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
            parrafo_pie.paragraph_format.space_after = Pt(6)
            run = parrafo_pie.add_run(valor_o_guion(figura.get("caption")))
            run.font.name = "Arial"
            run.font.size = Pt(8)
            run.italic = True

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_bloque_observaciones_docx(doc: Document, items: list[str]) -> None:
    if not items:
        return
    tabla = doc.add_table(rows=1, cols=1)
    aplicar_fila_no_dividir(tabla.rows[0])
    celda = tabla.rows[0].cells[0]
    aplicar_bordes_celda(celda, color="E5C46A")
    aplicar_sombreado_celda(celda, "FFF8E6")
    aplicar_margenes_celda(celda, margen=150)
    parrafo = celda.paragraphs[0]
    parrafo.paragraph_format.keep_with_next = True
    run = parrafo.add_run("Observaciones técnicas")
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.bold = True
    for item in items:
        p = celda.add_paragraph()
        p.style = "Informe Observacion"
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"- {item}")
        r.font.name = "Arial"
        r.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_bloque_destacado_docx(doc: Document, titulo: str, texto: str, fill: str = "F8FAFC") -> None:
    tabla = doc.add_table(rows=1, cols=1)
    aplicar_fila_no_dividir(tabla.rows[0])
    celda = tabla.rows[0].cells[0]
    aplicar_bordes_celda(celda, color="D8DEE6")
    aplicar_sombreado_celda(celda, fill)
    aplicar_margenes_celda(celda, margen=170)
    parrafo_titulo = celda.paragraphs[0]
    parrafo_titulo.paragraph_format.keep_with_next = True
    run_titulo = parrafo_titulo.add_run(valor_o_guion(titulo))
    run_titulo.font.name = "Arial"
    run_titulo.font.size = Pt(12)
    run_titulo.bold = True
    parrafo_texto = celda.add_paragraph()
    parrafo_texto.paragraph_format.space_after = Pt(0)
    run_texto = parrafo_texto.add_run(valor_o_guion(texto))
    run_texto.font.name = "Arial"
    run_texto.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_portada_editable_docx(doc: Document, contexto: dict) -> None:
    expediente = contexto["expediente"]
    portada = contexto["portada"]
    es_valoracion = contexto.get("es_valoracion")

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph(style="Informe Portada Titulo")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("INFORME DE VALORACIÓN INMOBILIARIA" if es_valoracion else "INFORME PERICIAL")
    run.font.name = "Arial"
    run.font.size = Pt(26)
    run.bold = True

    add_parrafo_editable(
        doc,
        "Valoración inmobiliaria según datos registrados"
        if es_valoracion
        else "Inspección técnica y registro de patologías",
        italic=True,
        centrado=True,
        espacio_despues=26,
    )

    tabla = doc.add_table(rows=0, cols=2)
    for etiqueta, valor in [
        ("Expediente", expediente.get("numero_expediente")),
        ("Cliente", expediente.get("cliente")),
        ("Dirección", expediente.get("direccion")),
        ("Fecha de emisión", contexto.get("fecha_emision")),
        ("Técnico", portada.get("tecnico")),
        ("Tipo de trabajo", portada.get("tipo_trabajo")),
    ]:
        row = tabla.add_row().cells
        row[0].text = valor_o_guion(etiqueta).upper()
        row[1].text = valor_o_guion(valor)
    aplicar_estilo_tabla_suave(tabla, sombrear_primera_columna=True)
    add_salto_pagina(doc)


def add_patologia_card_docx(doc: Document, patologia: dict, indice: int) -> None:
    tabla_titulo = doc.add_table(rows=1, cols=1)
    aplicar_fila_no_dividir(tabla_titulo.rows[0])
    celda = tabla_titulo.rows[0].cells[0]
    aplicar_bordes_celda(celda, color="D8DEE6")
    aplicar_sombreado_celda(celda, "F8FAFC")
    aplicar_margenes_celda(celda, margen=130)
    p = celda.paragraphs[0]
    p.style = "Informe Card Titulo"
    p.paragraph_format.keep_with_next = True
    run = p.add_run(f"Patología {indice} · {valor_o_guion(patologia.get('titulo'))}")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.bold = True

    add_tabla_campos_editable(doc, patologia.get("campos", []))
    add_bloque_figuras_editable(doc, patologia.get("fotos", []), ancho_cm=11.5)


def add_estancia_card_docx(doc: Document, estancia: dict, indice: int) -> None:
    meta = " · ".join(
        parte
        for parte in [
            f"Nivel: {estancia.get('nivel')}" if limpiar_texto(estancia.get("nivel")) else "",
            f"Unidad: {estancia.get('unidad')}" if limpiar_texto(estancia.get("unidad")) else "",
        ]
        if parte
    )
    tabla_titulo = doc.add_table(rows=1, cols=1)
    aplicar_fila_no_dividir(tabla_titulo.rows[0])
    celda_titulo = tabla_titulo.rows[0].cells[0]
    aplicar_bordes_celda(celda_titulo, color="D8DEE6")
    aplicar_sombreado_celda(celda_titulo, "F8FAFC")
    aplicar_margenes_celda(celda_titulo, margen=150)
    p_titulo = celda_titulo.paragraphs[0]
    p_titulo.style = "Informe Card Titulo"
    p_titulo.paragraph_format.keep_with_next = True
    run_titulo = p_titulo.add_run(f"{indice}. {valor_o_guion(estancia.get('nombre'))}")
    run_titulo.font.name = "Arial"
    run_titulo.font.size = Pt(13)
    run_titulo.bold = True
    if meta:
        p_meta = celda_titulo.add_paragraph()
        p_meta.paragraph_format.space_after = Pt(0)
        r_meta = p_meta.add_run(meta)
        r_meta.font.name = "Arial"
        r_meta.font.size = Pt(9)
        r_meta.italic = True
    if limpiar_texto(estancia.get("resumen")):
        add_parrafo_editable(doc, estancia["resumen"], espacio_despues=4)

    add_bloque_observaciones_docx(doc, estancia.get("incoherencias", []))
    add_tabla_campos_editable(doc, estancia.get("campos", []))
    add_grid_figuras_editable(doc, estancia.get("fotos", []), ancho_cm=7.5)

    if estancia.get("patologias"):
        for indice_patologia, patologia in enumerate(estancia["patologias"], start=1):
            add_patologia_card_docx(doc, patologia, indice_patologia)
    else:
        add_parrafo_editable(
            doc,
            "No constan patologías interiores registradas en esta estancia.",
            italic=True,
            espacio_despues=8,
        )


def add_zona_exterior_card_docx(doc: Document, zona: dict, indice: int) -> None:
    add_heading_editable(doc, f"{indice}. {valor_o_guion(zona.get('zona'))}", level=2)
    if limpiar_texto(zona.get("resumen")):
        add_parrafo_editable(doc, zona["resumen"], espacio_despues=4)
    add_bloque_observaciones_docx(doc, zona.get("incoherencias", []))
    for indice_patologia, patologia in enumerate(zona.get("patologias", []), start=1):
        add_patologia_card_docx(doc, patologia, indice_patologia)


TITULOS_VALORACION_DOCX = {
    "encargo": "4. Encargo",
    "documentacion": "5. Documentación utilizada",
    "identificacion": "6. Identificación del bien",
    "situacion_legal": "7. Situación legal",
    "entorno": "8. Entorno",
    "edificio_inmueble": "9. Edificio/inmueble",
    "constructivo": "10. Características constructivas",
    "estado": "11. Estado",
    "metodo": "12. Método de valoración",
    "resultado": "14. Resultado",
    "limitaciones": "15. Limitaciones",
}


def add_grupo_valoracion_docx(doc: Document, grupo: dict, bloque: dict, titulo: str) -> None:
    add_heading_editable(doc, titulo, level=1)
    visita = bloque.get("visita", {})
    add_parrafo_editable(
        doc,
        f"Visita de fecha {valor_o_guion(visita.get('fecha'))}. {valor_o_guion(visita.get('objeto'))}",
        italic=True,
        espacio_despues=4,
    )
    if grupo.get("hay_datos"):
        add_tabla_campos_editable(doc, grupo.get("campos", []))
    else:
        add_parrafo_editable(doc, "No constan datos registrados para este apartado.")


def add_comparables_valoracion_docx(doc: Document, comparables: list[dict]) -> None:
    add_heading_editable(doc, "13. Comparables", level=1)
    add_bloque_destacado_docx(
        doc,
        "Nota de cálculo inicial",
        "Los valores unitarios iniciales no incorporan todavía homogeneización ni ponderación técnica salvo que se indique expresamente.\n"
        "La homogeneización recoge ajustes técnicos introducidos por el perito. Los ajustes cualitativos no cuantificados se reflejan como observaciones y no modifican el valor unitario.",
        fill="FFF8E6",
    )
    if not comparables:
        add_parrafo_editable(doc, "No constan comparables de valoración registrados.")
        return

    resumen = construir_resumen_comparacion_contexto(comparables)
    add_heading_editable(doc, "Resumen comparativo y ponderación", level=2)
    add_bloque_destacado_docx(
        doc,
        "Nota de ponderación",
        "El resumen comparativo tiene carácter preparatorio. La adopción del valor unitario corresponde al criterio técnico del perito y debe quedar justificada expresamente.",
        fill="FFF8E6",
    )
    add_tabla_campos_editable(
        doc,
        construir_campos_informe(
            [
                ("Testigos incluidos", resumen.get("testigos_incluidos")),
                ("Testigos excluidos", resumen.get("testigos_excluidos")),
                ("€/m² mínimo", resumen.get("unitario_minimo_fmt")),
                ("€/m² máximo", resumen.get("unitario_maximo_fmt")),
                ("€/m² medio", resumen.get("unitario_medio_fmt")),
                ("€/m² mediana", resumen.get("unitario_mediana_fmt")),
                ("€/m² ponderado", resumen.get("unitario_ponderado_fmt")),
                ("Suma de pesos", resumen.get("suma_pesos_fmt")),
                (
                    "Propuesta orientativa",
                    resumen.get("propuesta_unitaria_orientativa_fmt"),
                ),
            ]
        )
    )
    if resumen.get("advertencias"):
        add_bloque_destacado_docx(
            doc,
            "Advertencias de ponderación",
            "\n".join(resumen.get("advertencias", [])),
            fill="FFF8E6",
        )

    for indice, comparable in enumerate(comparables, start=1):
        add_heading_editable(
            doc,
            f"13.{indice} Comparable {indice} · visita {valor_o_guion(comparable.get('visita_fecha'))}",
            level=2,
        )
        add_tabla_campos_editable(doc, comparable.get("campos", []))
        if comparable.get("pasos_homogeneizacion"):
            add_heading_editable(doc, "Matriz de homogeneización", level=3)
            add_tabla_campos_editable(
                doc,
                construir_campos_informe(
                    [
                        (
                            paso.get("variable") or "Ajuste",
                            (
                                f"Inmueble: {valor_o_guion(paso.get('valor_inmueble'))}\n"
                                f"Testigo: {valor_o_guion(paso.get('valor_testigo'))}\n"
                                f"Efecto: {formatear_precio_unitario_es(paso.get('efecto_importe_m2'))}\n"
                                f"Justificación: {valor_o_guion(paso.get('justificacion'))}"
                            ),
                        )
                        for paso in comparable.get("pasos_homogeneizacion", [])
                    ]
                ),
            )
        if comparable.get("advertencias_homogeneizacion"):
            add_bloque_destacado_docx(
                doc,
                "Advertencias de homogeneización",
                "\n".join(comparable.get("advertencias_homogeneizacion", [])),
                fill="FFF8E6",
            )


def add_valoracion_eco_docx(doc: Document, contexto: dict) -> None:
    valoracion_eco = contexto.get("valoracion_eco") or {}
    if not valoracion_eco:
        return
    add_heading_editable(doc, "3. Criterios periciales de valoración", level=1)
    add_bloque_destacado_docx(
        doc,
        "Nota metodológica",
        valoracion_eco.get("nota_metodologica", VALORACION_NOTA_ECO),
        fill="EEF6FF",
    )
    for grupo in [
        valoracion_eco.get("finalidad"),
        valoracion_eco.get("base_valor"),
        valoracion_eco.get("superficies"),
    ]:
        if not grupo:
            continue
        add_heading_editable(doc, grupo.get("titulo", "Apartado de valoración"), level=2)
        if grupo.get("hay_datos"):
            add_tabla_campos_editable(doc, grupo.get("campos", []))
        else:
            add_parrafo_editable(doc, "No constan datos registrados para este apartado.")

    metodos = valoracion_eco.get("metodos") or {}
    add_heading_editable(doc, metodos.get("titulo", "Métodos aplicados y descartados"), level=2)
    if metodos.get("hay_datos"):
        campos_metodos = []
        for metodo in metodos.get("items", []):
            detalle = (
                f"Aplicado: {'Sí' if metodo.get('aplicado') else 'No'} · "
                f"Descartado: {'Sí' if metodo.get('descartado') else 'No'}"
            )
            if metodo.get("justificacion"):
                detalle += f"\nJustificación: {metodo['justificacion']}"
            if metodo.get("observaciones"):
                detalle += f"\nObservaciones: {metodo['observaciones']}"
            campos_metodos.append((metodo.get("titulo", "Método"), detalle))
        add_tabla_campos_editable(doc, construir_campos_informe(campos_metodos))
    else:
        add_parrafo_editable(doc, "No constan métodos aplicados o descartados.")

    incidencias = valoracion_eco.get("incidencias") or {}
    add_heading_editable(
        doc,
        incidencias.get("titulo", "Condicionantes, advertencias y limitaciones"),
        level=2,
    )
    visibles = incidencias.get("visibles", [])
    if visibles:
        add_tabla_campos_editable(
            doc,
            construir_campos_informe(
                [
                    (
                        f"{incidencia.get('tipo', '').capitalize()} · {incidencia.get('origen', '')}",
                        incidencia.get("descripcion", ""),
                    )
                    for incidencia in visibles
                ]
            ),
        )
    else:
        add_parrafo_editable(doc, "No constan incidencias visibles para el informe.")


def add_cuerpo_valoracion_docx(doc: Document, contexto: dict) -> None:
    add_heading_editable(doc, "2. Objeto del informe", level=1)
    add_parrafo_editable(
        doc,
        "El presente informe tiene por objeto documentar la valoración inmobiliaria del bien identificado en "
        "el expediente, incorporando el encargo, la documentación disponible, los datos observados en visita, "
        "los comparables registrados, el método utilizado, el resultado y las limitaciones declaradas.",
    )
    advertencias = contexto.get("completitud_valoracion", {}).get("advertencias", [])
    if advertencias:
        add_bloque_destacado_docx(
            doc,
            "Advertencias de completitud de valoración",
            "Estas advertencias no bloquean la generación manual del informe.\n"
            + "\n".join(advertencia["mensaje"] for advertencia in advertencias),
            fill="FFF8E6",
        )

    add_valoracion_eco_docx(doc, contexto)

    valoraciones = contexto.get("valoracion") or []
    if valoraciones:
        for bloque in valoraciones:
            for grupo in bloque.get("grupos", []):
                clave = grupo.get("clave")
                if clave in {"base_valor", "metodos_eco", "incidencias", "resultado", "limitaciones"}:
                    continue
                add_grupo_valoracion_docx(
                    doc,
                    grupo,
                    bloque,
                    TITULOS_VALORACION_DOCX.get(clave, valor_o_guion(grupo.get("titulo"))),
                )
    else:
        for clave, titulo in TITULOS_VALORACION_DOCX.items():
            if clave in {"base_valor", "metodos_eco", "incidencias", "resultado", "limitaciones"}:
                continue
            add_heading_editable(doc, titulo, level=1)
            add_parrafo_editable(doc, "No constan datos de valoración registrados.")

    add_comparables_valoracion_docx(doc, contexto.get("comparables_valoracion") or [])

    if valoraciones:
        for bloque in valoraciones:
            for grupo in bloque.get("grupos", []):
                clave = grupo.get("clave")
                if clave not in {"resultado", "limitaciones"}:
                    continue
                add_grupo_valoracion_docx(
                    doc,
                    grupo,
                    bloque,
                    TITULOS_VALORACION_DOCX.get(clave, valor_o_guion(grupo.get("titulo"))),
                )
    else:
        add_heading_editable(doc, "14. Resultado", level=1)
        add_parrafo_editable(doc, "No consta resultado de valoración registrado.")
        add_heading_editable(doc, "15. Limitaciones", level=1)
        add_parrafo_editable(doc, "No constan limitaciones de valoración registradas.")


def generar_informe_docx_editable_bytes(
    expediente_id: int,
    incluir_anexo_economico_reparacion: bool = False,
) -> bytes:
    contexto = build_informe_context(
        expediente_id,
        incluir_anexo_economico_reparacion=incluir_anexo_economico_reparacion,
    )
    doc = Document()
    configurar_documento_editable(doc)

    add_portada_editable_docx(doc, contexto)

    add_heading_editable(doc, "1. Identificación", level=1)
    add_tabla_campos_editable(doc, contexto["identificacion"])

    if contexto.get("es_valoracion"):
        add_cuerpo_valoracion_docx(doc, contexto)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    add_heading_editable(doc, "2. Objeto del informe", level=1)
    add_parrafo_editable(
        doc,
        "El presente informe pericial documenta las patologías observadas en el inmueble objeto del expediente, "
        "incorporando los antecedentes disponibles, las visitas realizadas y los registros técnicos interiores "
        "y/o exteriores existentes en el sistema.",
    )

    add_heading_editable(doc, "3. Descripción", level=1)
    add_tabla_campos_editable(doc, contexto["descripcion_inmueble"])

    add_heading_editable(doc, "4. Datos específicos del informe de patologías", level=1)
    add_tabla_campos_editable(doc, contexto["datos_patologias"])

    if contexto.get("bloque_judicial"):
        add_heading_editable(doc, "5. Encargo judicial", level=1)
        add_tabla_campos_editable(doc, contexto["bloque_judicial"])
        offset = 1
    else:
        offset = 0

    add_heading_editable(doc, f"{5 + offset}. Inspección", level=1)
    if contexto["visitas"]:
        for indice, visita in enumerate(contexto["visitas"], start=1):
            add_heading_editable(doc, f"{5 + offset}.{indice} Visita de fecha {valor_o_guion(visita.get('fecha'))}", level=2)
            add_tabla_campos_editable(
                doc,
                construir_campos_informe(
                    [
                        ("Técnico", visita.get("tecnico")),
                        ("Objeto de visita", visita.get("objeto")),
                        ("Climatología", visita.get("climatologia")),
                        ("Observaciones de visita", visita.get("observaciones")),
                    ]
                ),
            )
            if visita.get("fotos_exteriores"):
                add_heading_editable(doc, "Fotografías exteriores descriptivas", level=3)
                add_grid_figuras_editable(
                    doc,
                    [foto["figura"] for foto in visita["fotos_exteriores"]],
                    ancho_cm=7.5,
                )
    else:
        add_parrafo_editable(doc, "No constan visitas registradas en el expediente.")

    add_heading_editable(doc, f"{6 + offset}. Patologías interiores", level=1)
    if contexto["estancias"]:
        for indice_estancia, estancia in enumerate(contexto["estancias"], start=1):
            add_estancia_card_docx(doc, estancia, indice_estancia)
    else:
        add_parrafo_editable(doc, "No constan estancias registradas.")

    add_heading_editable(doc, f"{7 + offset}. Patologías exteriores", level=1)
    if contexto["patologias_exteriores"]:
        for indice_zona, zona in enumerate(contexto["patologias_exteriores"], start=1):
            add_zona_exterior_card_docx(doc, zona, indice_zona)
    else:
        add_parrafo_editable(doc, "No constan patologías exteriores registradas.")

    if contexto.get("mapas"):
        add_heading_editable(doc, f"{8 + offset}. Mapas de localización de daños", level=1)
        for bloque in contexto["mapas"]:
            for mapa in bloque.get("mapas", []):
                add_heading_editable(doc, valor_o_guion(mapa.get("titulo")), level=2)
                add_tabla_campos_editable(
                    doc,
                    construir_campos_informe(
                        [
                            ("Descripción", mapa.get("descripcion")),
                            ("Contexto", mapa.get("objeto_visita_label")),
                        ]
                    ),
                )
                if limpiar_texto(mapa.get("imagen_overlay_url")):
                    add_imagen_editable(
                        doc,
                        {"data_uri": mapa["imagen_overlay_url"], "caption": ""},
                        ancho_cm=12,
                    )
                elif limpiar_texto(mapa.get("imagen_base")):
                    add_imagen_editable(
                        doc,
                        {"archivo": mapa["imagen_base"], "caption": ""},
                        ancho_cm=12,
                    )
                for cuadrante in mapa.get("cuadrantes", []):
                    add_heading_editable(
                        doc,
                        f"Cuadrante {valor_o_guion(cuadrante.get('codigo_cuadrante'))}",
                        level=3,
                    )
                    add_tabla_campos_editable(
                        doc,
                        construir_campos_informe(
                            [
                                ("Descripción", cuadrante.get("descripcion")),
                                ("Gravedad", cuadrante.get("gravedad_label")),
                                ("Patología vinculada", cuadrante.get("patologia_vinculada_resumen")),
                            ]
                        ),
                    )
                    if cuadrante.get("foto_figura"):
                        add_imagen_editable(doc, cuadrante["foto_figura"], ancho_cm=10.5)
        post_mapas = 1
    else:
        post_mapas = 0

    add_heading_editable(doc, f"{8 + offset + post_mapas}. Análisis técnico", level=1)
    add_tabla_campos_editable(doc, contexto["analisis_tecnico"])

    add_heading_editable(doc, f"{9 + offset + post_mapas}. Propuesta de reparación", level=1)
    add_parrafo_editable(doc, contexto["propuesta_reparacion"])

    post_anexo = 0
    if contexto.get("anexo_economico_reparacion", {}).get("incluido"):
        anexo = contexto["anexo_economico_reparacion"]
        add_heading_editable(doc, f"{10 + offset + post_mapas}. Anexo económico de reparación", level=1)
        add_parrafo_editable(doc, anexo["nota"])
        add_parrafo_editable(
            doc,
            "Alcance económico: PEM orientativo. No incluye IVA, beneficio industrial, "
            "gastos generales ni constituye oferta comercial.",
        )
        if anexo.get("modo") == "actuaciones":
            for actuacion in anexo["actuaciones"]:
                add_heading_editable(doc, actuacion["titulo"], level=2)
                add_tabla_campos_editable(
                    doc,
                    construir_campos_informe(
                        [
                            ("Descripción", actuacion.get("descripcion")),
                            ("Observaciones", actuacion.get("observaciones")),
                            ("Subtotal", f"{actuacion.get('subtotal', 0):.2f} €"),
                        ]
                    ),
                )
                for partida in actuacion["partidas"]:
                    add_tabla_campos_editable(
                        doc,
                        construir_campos_informe(
                            [
                                ("Código", partida.get("codigo")),
                                ("Partida", partida.get("partida")),
                                ("Cantidad", partida.get("cantidad")),
                                ("Unidad", partida.get("unidad")),
                                ("Precio unitario", f"{partida.get('precio_unitario', 0):.2f} €"),
                                ("Importe", f"{partida.get('importe', 0):.2f} €"),
                            ]
                        ),
                    )
        else:
            for patologia in anexo["patologias"]:
                add_heading_editable(doc, patologia["patologia"], level=2)
                add_tabla_campos_editable(
                    doc,
                    construir_campos_informe(
                        [
                            ("Zona", patologia.get("zona")),
                            ("Subtotal", f"{patologia.get('subtotal', 0):.2f} €"),
                        ]
                    ),
                )
                for partida in patologia["partidas"]:
                    add_tabla_campos_editable(
                        doc,
                        construir_campos_informe(
                            [
                                ("Código", partida.get("codigo")),
                                ("Partida", partida.get("partida")),
                                ("Cantidad", partida.get("cantidad")),
                                ("Unidad", partida.get("unidad")),
                                ("Precio unitario", f"{partida.get('precio_unitario', 0):.2f} €"),
                                ("Importe", f"{partida.get('importe', 0):.2f} €"),
                            ]
                        ),
                    )
        add_parrafo_editable(doc, f"Total PEM de reparación: {anexo['total_pem']:.2f} €")
        post_anexo = 1

    add_salto_pagina(doc)
    add_heading_editable(doc, f"{10 + offset + post_mapas + post_anexo}. Conclusiones técnicas", level=1)
    add_bloque_destacado_docx(
        doc,
        "Conclusiones técnicas",
        contexto["conclusion_tecnica"],
        fill="F8FAFC",
    )

    add_heading_editable(doc, f"{11 + offset + post_mapas + post_anexo}. Conclusiones periciales", level=1)
    add_bloque_destacado_docx(
        doc,
        "Conclusiones periciales",
        contexto["conclusion_pericial"],
        fill="F8FAFC",
    )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generar_informe_pdf_bytes(
    request: Request,
    expediente_id: int,
    incluir_anexo_economico_reparacion: bool = False,
) -> bytes:
    try:
        from playwright.sync_api import sync_playwright
    except ImportError as exc:
        raise HTTPException(
            status_code=500,
            detail="Playwright no está instalado. Actualiza dependencias para generar PDFs.",
        ) from exc

    contexto = build_informe_context(
        expediente_id,
        base_url=str(request.base_url).rstrip("/"),
        incluir_anexo_economico_reparacion=incluir_anexo_economico_reparacion,
    )
    template = request.app.state.templates.env.get_template("informes/imprimir.html")

    def render_html():
        return template.render({
            "request": request,
            "current_user": getattr(request.state, "current_user", None),
            "modo_pdf": True,
            **contexto,
        })

    html = render_html()

    try:
        with sync_playwright() as p:
            browser = p.chromium.launch()
            page = browser.new_page(viewport={"width": 794, "height": 1123})
            page.emulate_media(media="print")
            page.set_content(html, wait_until="networkidle")
            paginas_toc = page.evaluate(
                """
                () => {
                    const probe = document.createElement('div');
                    probe.style.width = '100mm';
                    probe.style.height = '1px';
                    probe.style.position = 'absolute';
                    probe.style.left = '-10000px';
                    document.body.appendChild(probe);
                    const pxPerMm = probe.getBoundingClientRect().width / 100;
                    probe.remove();
                    const pageHeight = 264 * pxPerMm;
                    const documentTop = document.querySelector('.document')?.getBoundingClientRect().top + window.scrollY || 0;
                    const result = {};
                    document.querySelectorAll('[data-toc-id]').forEach((el) => {
                        const id = el.getAttribute('data-toc-id');
                        const top = el.getBoundingClientRect().top + window.scrollY - documentTop;
                        result[id] = Math.max(1, Math.floor(top / pageHeight) + 1);
                    });
                    return result;
                }
                """
            )
            if paginas_toc:
                contexto["toc_items"] = construir_toc_items_informe(contexto, paginas_toc)
                html = render_html()
                page.set_content(html, wait_until="networkidle")
            pdf_bytes = page.pdf(
                format="A4",
                margin={
                    "top": "15mm",
                    "right": "16mm",
                    "bottom": "18mm",
                    "left": "16mm",
                },
                display_header_footer=True,
                header_template="<span></span>",
                footer_template=(
                    "<div style='width:100%;font-family:Arial,Helvetica,sans-serif;"
                    "font-size:8px;color:#6b7280;text-align:center;'>"
                    "Página <span class='pageNumber'></span> de <span class='totalPages'></span>"
                    "</div>"
                ),
                print_background=True,
                prefer_css_page_size=True,
            )
            browser.close()
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail="No se pudo generar el PDF del informe.",
        ) from exc

    return pdf_bytes


def generar_informe_v2_pdf_bytes(request: Request, contexto: dict) -> bytes:
    try:
        from playwright.sync_api import sync_playwright
    except ImportError as exc:
        raise HTTPException(
            status_code=500,
            detail="Playwright no está instalado. Actualiza dependencias para generar PDFs.",
        ) from exc

    template = request.app.state.templates.env.get_template("informes/v2_pdf.html")
    numero_expediente = str((contexto.get("expediente") or {}).get("numero_expediente") or "").strip()
    pie_profesional = "Carlos Blanco | Arquitecto Técnico | Colegiado nº 5866"
    pie_derecha = f"Expediente {numero_expediente}" if numero_expediente else "Informe pericial"
    header_template = "<span></span>"
    footer_template = (
        "<div style='width:100%;padding:0 16mm;font-family:Aptos,Segoe UI,Arial,Helvetica,sans-serif;"
        "font-size:8px;color:#6b7280;display:flex;justify-content:space-between;gap:12px;"
        "box-sizing:border-box;'>"
        f"<span>{html.escape(pie_profesional)}</span>"
        f"<span>{html.escape(pie_derecha)}</span>"
        "</div>"
    )

    def renderizar_html(contexto_pdf: dict) -> str:
        return template.render(
            {
                "request": request,
                "current_user": getattr(request.state, "current_user", None),
                "modo_pdf": True,
                **contexto_pdf,
            }
        )

    def generar_pdf_desde_html(html: str) -> bytes:
        with sync_playwright() as p:
            browser = p.chromium.launch()
            page = browser.new_page(viewport={"width": 794, "height": 1123})
            page.emulate_media(media="print")
            page.set_content(html, wait_until="networkidle")
            pdf = page.pdf(
                format="A4",
                margin={
                    "top": "15mm",
                    "right": "16mm",
                    "bottom": "18mm",
                    "left": "16mm",
                },
                display_header_footer=True,
                header_template=header_template,
                footer_template=footer_template,
                print_background=True,
                prefer_css_page_size=True,
            )
            browser.close()
            return pdf

    def contar_paginas_pdf(pdf_bytes: bytes) -> int:
        try:
            from pypdf import PdfReader

            return len(PdfReader(BytesIO(pdf_bytes)).pages)
        except Exception:
            return 0

    def completar_paginas_documentos_aportados(
        indice_paginas: dict[str, int],
        contexto_pdf: dict,
        total_paginas_informe: int,
    ) -> dict[str, int]:
        if total_paginas_informe <= 0:
            return indice_paginas
        paginas = dict(indice_paginas)
        desplazamiento_documentacion = int(
            contexto_pdf.get("desplazamiento_paginas_documentacion_aportada") or 0
        )
        pagina_documento = total_paginas_informe + desplazamiento_documentacion + 1
        for item in contexto_pdf.get("indice") or []:
            if item.get("grupo") != "documentacion_documento":
                continue
            clave = str(item.get("clave") or "").strip()
            if not clave:
                continue
            paginas[clave] = pagina_documento
            pagina_documento += 1 + int(item.get("paginas_pdf") or 0)
        return paginas

    try:
        pdf_bytes = generar_pdf_desde_html(renderizar_html(contexto))
        total_paginas_informe = contar_paginas_pdf(pdf_bytes)
        indice_paginas = extraer_paginas_indice_informe_v2(
            pdf_bytes,
            contexto.get("indice") or [],
        )
        desplazamiento_documentacion = int(
            contexto.get("desplazamiento_paginas_documentacion_aportada") or 0
        )
        if desplazamiento_documentacion and "documentacion_aportada" in indice_paginas:
            indice_paginas["documentacion_aportada"] += desplazamiento_documentacion
        indice_paginas = completar_paginas_documentos_aportados(
            indice_paginas,
            contexto,
            total_paginas_informe,
        )
        if indice_paginas:
            contexto_con_indice = {
                **contexto,
                "indice_paginas": indice_paginas,
            }
            pdf_bytes = generar_pdf_desde_html(renderizar_html(contexto_con_indice))
            total_paginas_informe_actualizado = contar_paginas_pdf(pdf_bytes)
            if total_paginas_informe_actualizado and total_paginas_informe_actualizado != total_paginas_informe:
                indice_paginas = completar_paginas_documentos_aportados(
                    indice_paginas,
                    contexto,
                    total_paginas_informe_actualizado,
                )
                contexto_con_indice = {
                    **contexto,
                    "indice_paginas": indice_paginas,
                }
                pdf_bytes = generar_pdf_desde_html(renderizar_html(contexto_con_indice))
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail="No se pudo generar el PDF del informe.",
        ) from exc

    return pdf_bytes


def normalizar_texto_indice_pdf_v2(texto: str) -> str:
    texto = unicodedata.normalize("NFKD", str(texto or "").lower())
    texto = "".join(c for c in texto if not unicodedata.combining(c))
    return re.sub(r"[^a-z0-9]+", " ", texto).strip()


def extraer_paginas_indice_informe_v2(pdf_bytes: bytes, indice: list[dict]) -> dict[str, int]:
    try:
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(pdf_bytes))
    except Exception:
        return {}

    textos_paginas = []
    for pagina in reader.pages:
        try:
            textos_paginas.append(normalizar_texto_indice_pdf_v2(pagina.extract_text() or ""))
        except Exception:
            textos_paginas.append("")

    paginas: dict[str, int] = {}
    ultima_pagina_anexo: int | None = None
    for item in indice:
        clave = str(item.get("clave") or "")
        if not clave:
            continue
        if clave == "portada":
            paginas[clave] = 1
            continue
        patrones = []
        numero = str(item.get("numero") or "").strip()
        titulo = str(item.get("titulo") or "").strip()
        if clave == "indice":
            patrones = ["indice"]
        elif item.get("grupo") == "anexos":
            patrones = [f"ANEXO {numero} {titulo}"]
        elif item.get("grupo") == "documentacion_documento":
            continue
        else:
            patrones = [f"{numero} {titulo}"]
        patrones_normalizados = [
            normalizar_texto_indice_pdf_v2(patron)
            for patron in patrones
            if patron
        ]
        if clave == "indice":
            pagina_inicio_busqueda = 1
        elif item.get("grupo") == "anexos":
            pagina_inicio_busqueda = max(
                3,
                (ultima_pagina_anexo + 1)
                if ultima_pagina_anexo is not None
                else (paginas.get("conclusiones", 3) + 1),
            )
        else:
            pagina_inicio_busqueda = 3
        for indice_pagina, texto_pagina in enumerate(textos_paginas, start=1):
            if indice_pagina < pagina_inicio_busqueda:
                continue
            if all(patron in texto_pagina for patron in patrones_normalizados):
                paginas[clave] = indice_pagina
                if item.get("grupo") == "anexos":
                    ultima_pagina_anexo = indice_pagina
                break
    return paginas


def generar_informe_inspeccion(cur, expediente, visitas, doc: Document) -> None:
    add_portada(doc, expediente)

    doc.add_heading("1. Identificación del expediente", level=1)
    add_tabla_datos_expediente(doc, expediente)

    doc.add_heading("2. Objeto del informe", level=1)
    add_parrafo(
        doc,
        "El presente informe técnico recoge las comprobaciones de inspección registradas en el inmueble objeto del expediente, "
        "incluyendo las condiciones de la visita, los checklists generales, la revisión por estancias, la inspección exterior "
        "y, en su caso, los elementos comunes del edificio.",
    )

    doc.add_heading("3. Descripción del inmueble", level=1)
    add_etiqueta_valor(doc, "Referencia catastral", expediente["referencia_catastral"])
    add_etiqueta_valor(doc, "Observaciones generales", expediente["observaciones_generales"])
    add_etiqueta_valor(doc, "Planta de la unidad", expediente["planta_unidad"])
    add_etiqueta_valor(doc, "Puerta / unidad", expediente["puerta_unidad"])
    add_etiqueta_valor(doc, "Superficie construida", expediente["superficie_construida"])
    add_etiqueta_valor(doc, "Superficie útil", expediente["superficie_util"])
    add_etiqueta_valor(doc, "Dormitorios", expediente["dormitorios_unidad"])
    add_etiqueta_valor(doc, "Baños", expediente["banos_unidad"])
    add_etiqueta_valor(doc, "Observaciones del bloque", expediente["observaciones_bloque"])
    add_etiqueta_valor(doc, "Observaciones de la unidad", expediente["observaciones_unidad"])

    numero_apartado = 4
    add_apartado_bloque_judicial(doc, numero_apartado, expediente)
    if limpiar_texto(expediente["destinatario"]) == "judicial":
        numero_apartado += 1

    visitas_inspeccion = []
    for visita in visitas:
        climatologia = cur.execute(
            """
            SELECT *
            FROM climatologia_visitas
            WHERE visita_id = ?
            ORDER BY id DESC
            LIMIT 1
            """,
            (visita["id"],),
        ).fetchone()
        visitas_inspeccion.append(
            {
                "visita": visita,
                "climatologia": climatologia,
                "inspeccion": cargar_inspeccion_visita(cur, visita["id"]),
            }
        )

    add_apartado_condiciones_visita_inspeccion(doc, numero_apartado, visitas_inspeccion)
    numero_apartado += 1
    add_apartado_comprobaciones_generales(doc, numero_apartado, visitas_inspeccion)
    numero_apartado += 1
    add_apartado_inspeccion_estancias(doc, numero_apartado, visitas_inspeccion)
    numero_apartado += 1
    add_apartado_inspeccion_exterior(doc, numero_apartado, visitas_inspeccion)
    numero_apartado += 1
    add_apartado_elementos_comunes(doc, numero_apartado, visitas_inspeccion)
    numero_apartado += 1
    add_apartado_observaciones_tecnicas(doc, numero_apartado, visitas_inspeccion)
    numero_apartado += 1
    add_apartado_recomendaciones_inspeccion(doc, numero_apartado, visitas_inspeccion)
    numero_apartado += 1
    add_apartado_conclusion_inspeccion(doc, numero_apartado, visitas_inspeccion)


def generar_informe_habitabilidad(cur, expediente, visitas, doc: Document) -> None:
    add_portada(doc, expediente)

    doc.add_heading("1. Identificación del expediente", level=1)
    add_tabla_datos_expediente(doc, expediente)

    doc.add_heading("2. Objeto del informe", level=1)
    add_parrafo(
        doc,
        "El presente informe técnico recoge la evaluación de habitabilidad registrada en el inmueble objeto del expediente, "
        "incluyendo las condiciones generales, la valoración por estancias y los factores exteriores mínimos que pueden afectar al uso residencial.",
    )

    doc.add_heading("3. Descripción del inmueble", level=1)
    add_etiqueta_valor(doc, "Referencia catastral", expediente["referencia_catastral"])
    add_etiqueta_valor(doc, "Observaciones generales", expediente["observaciones_generales"])
    add_etiqueta_valor(doc, "Planta de la unidad", expediente["planta_unidad"])
    add_etiqueta_valor(doc, "Puerta / unidad", expediente["puerta_unidad"])
    add_etiqueta_valor(doc, "Superficie construida", expediente["superficie_construida"])
    add_etiqueta_valor(doc, "Superficie útil", expediente["superficie_util"])
    add_etiqueta_valor(doc, "Dormitorios", expediente["dormitorios_unidad"])
    add_etiqueta_valor(doc, "Baños", expediente["banos_unidad"])
    add_etiqueta_valor(doc, "Observaciones del bloque", expediente["observaciones_bloque"])
    add_etiqueta_valor(doc, "Observaciones de la unidad", expediente["observaciones_unidad"])

    visitas_habitabilidad = []
    for visita in visitas:
        climatologia = cur.execute(
            """
            SELECT *
            FROM climatologia_visitas
            WHERE visita_id = ?
            ORDER BY id DESC
            LIMIT 1
            """,
            (visita["id"],),
        ).fetchone()
        visitas_habitabilidad.append(
            {
                "visita": visita,
                "climatologia": climatologia,
                "habitabilidad": cargar_habitabilidad_visita(cur, visita["id"]),
            }
        )

    numero_apartado = 4
    add_apartado_condiciones_habitabilidad(doc, numero_apartado, visitas_habitabilidad)
    numero_apartado += 1
    add_apartado_evaluacion_estancias_habitabilidad(
        doc, numero_apartado, visitas_habitabilidad
    )
    numero_apartado += 1
    add_apartado_exterior_habitabilidad(doc, numero_apartado, visitas_habitabilidad)
    numero_apartado += 1
    add_apartado_conclusion_habitabilidad(doc, numero_apartado, visitas_habitabilidad)
    numero_apartado += 1
    add_apartado_bloque_judicial(doc, numero_apartado, expediente)


def generar_informe_valoracion(cur, expediente, visitas, doc: Document) -> None:
    add_portada(doc, expediente)

    doc.add_heading("1. Identificación del expediente", level=1)
    add_tabla_datos_expediente(doc, expediente)

    doc.add_heading("2. Objeto del informe", level=1)
    add_parrafo(
        doc,
        "El presente informe técnico recoge la información disponible para la valoración inmobiliaria del bien objeto del expediente, "
        "incluyendo la documentación consultada, las características del inmueble, los criterios de valoración y los comparables registrados.",
    )

    doc.add_heading("3. Identificación y descripción del bien", level=1)
    add_etiqueta_valor(doc, "Referencia catastral", expediente["referencia_catastral"])
    add_etiqueta_valor(doc, "Observaciones generales", expediente["observaciones_generales"])
    add_etiqueta_valor(doc, "Planta de la unidad", expediente["planta_unidad"])
    add_etiqueta_valor(doc, "Puerta / unidad", expediente["puerta_unidad"])
    add_etiqueta_valor(doc, "Superficie construida", expediente["superficie_construida"])
    add_etiqueta_valor(doc, "Superficie útil", expediente["superficie_util"])
    add_etiqueta_valor(doc, "Dormitorios", expediente["dormitorios_unidad"])
    add_etiqueta_valor(doc, "Baños", expediente["banos_unidad"])
    add_etiqueta_valor(doc, "Observaciones del bloque", expediente["observaciones_bloque"])
    add_etiqueta_valor(doc, "Observaciones de la unidad", expediente["observaciones_unidad"])

    visitas_valoracion = []
    for visita in visitas:
        comparables = cur.execute(
            """
            SELECT *
            FROM comparables_valoracion
            WHERE visita_id = ?
            ORDER BY id ASC
            """,
            (visita["id"],),
        ).fetchall()
        visitas_valoracion.append(
            {
                "visita": visita,
                "valoracion": cargar_valoracion_visita(cur, visita["id"]),
                "comparables": comparables,
            }
        )

    numero_apartado = 4
    add_apartado_valoracion_por_visitas(
        doc, numero_apartado, "Encargo / solicitante", visitas_valoracion, VALORACION_ENCARGO_ITEMS
    )
    numero_apartado += 1
    add_apartado_valoracion_por_visitas(
        doc, numero_apartado, "Documentación utilizada", visitas_valoracion, VALORACION_DOCUMENTACION_ITEMS
    )
    numero_apartado += 1
    add_apartado_valoracion_por_visitas(
        doc, numero_apartado, "Situación legal y urbanística", visitas_valoracion, VALORACION_SITUACION_LEGAL_ITEMS
    )
    numero_apartado += 1
    add_apartado_valoracion_por_visitas(
        doc, numero_apartado, "Entorno", visitas_valoracion, VALORACION_ENTORNO_ITEMS
    )
    numero_apartado += 1
    add_apartado_valoracion_por_visitas(
        doc, numero_apartado, "Descripción del edificio e inmueble", visitas_valoracion, VALORACION_IDENTIFICACION_ITEMS + VALORACION_EDIFICIO_INMUEBLE_ITEMS
    )
    numero_apartado += 1
    add_apartado_valoracion_por_visitas(
        doc, numero_apartado, "Características constructivas", visitas_valoracion, VALORACION_CONSTRUCTIVO_ITEMS
    )
    numero_apartado += 1
    add_apartado_valoracion_por_visitas(
        doc, numero_apartado, "Estado actual y ocupación", visitas_valoracion, VALORACION_ESTADO_ITEMS
    )
    numero_apartado += 1
    add_apartado_valoracion_por_visitas(
        doc, numero_apartado, "Método / criterios de valoración", visitas_valoracion, VALORACION_METODO_ITEMS
    )
    numero_apartado += 1
    add_apartado_comparables_valoracion(doc, numero_apartado, visitas_valoracion)
    numero_apartado += 1
    add_apartado_resultado_valoracion(doc, numero_apartado, visitas_valoracion)
    numero_apartado += 1
    add_apartado_limitaciones_valoracion(doc, numero_apartado, visitas_valoracion)
    numero_apartado += 1
    add_apartado_bloque_judicial(doc, numero_apartado, expediente)


def generar_informe(expediente_id: int) -> tuple[str, str]:
    os.makedirs(INFORMES_DIR, exist_ok=True)

    conn = get_connection()
    cur = conn.cursor()

    expediente = cur.execute(
        """
        SELECT *
        FROM expedientes
        WHERE id = ?
        """,
        (expediente_id,),
    ).fetchone()

    if not expediente:
        conn.close()
        raise ValueError("Expediente no encontrado")

    visitas = cur.execute(
        """
        SELECT *
        FROM visitas
        WHERE expediente_id = ?
        ORDER BY id ASC
        """,
        (expediente_id,),
    ).fetchall()

    doc = Document()
    configurar_documento(doc)

    tipo_informe = limpiar_texto(expediente["tipo_informe"])

    if tipo_informe == "patologias":
        generar_informe_patologias(cur, expediente, visitas, doc)
    elif tipo_informe == "inspeccion":
        generar_informe_inspeccion(cur, expediente, visitas, doc)
    elif tipo_informe == "habitabilidad":
        generar_informe_habitabilidad(cur, expediente, visitas, doc)
    elif tipo_informe == "valoracion":
        generar_informe_valoracion(cur, expediente, visitas, doc)
    else:
        add_portada(doc, expediente)
        add_apartado_introduccion(doc)
        add_apartado_datos_generales(doc, expediente)
        add_apartado_reforma(doc, expediente)

        numero_apartado = 4

        for visita in visitas:
            climatologia = cur.execute(
                """
                SELECT *
                FROM climatologia_visitas
                WHERE visita_id = ?
                ORDER BY id DESC
                LIMIT 1
                """,
                (visita["id"],),
            ).fetchone()

            estancias = cur.execute(
                """
                SELECT *
                FROM estancias
                WHERE visita_id = ?
                ORDER BY id ASC
                """,
                (visita["id"],),
            ).fetchall()

            patologias = cur.execute(
                """
                SELECT rp.*, e.nombre AS estancia_nombre
                FROM registros_patologias rp
                INNER JOIN estancias e ON rp.estancia_id = e.id
                WHERE rp.visita_id = ?
                ORDER BY e.nombre ASC, rp.id ASC
                """,
                (visita["id"],),
            ).fetchall()

            doc.add_section(WD_SECTION.NEW_PAGE)
            configurar_documento(doc)

            add_apartado_visita(
                doc,
                numero_apartado,
                visita,
                climatologia,
                estancias,
                patologias,
            )

            numero_apartado += 1

        doc.add_section(WD_SECTION.NEW_PAGE)
        configurar_documento(doc)
        add_apartado_conclusion(doc)

    base_nombre = limpiar_nombre_archivo(
        f"{expediente['numero_expediente']}_{expediente['cliente']}"
    )
    marca_tiempo = datetime.now().strftime("%Y%m%d_%H%M%S")
    nombre_archivo = f"informe_{base_nombre}_{marca_tiempo}.docx"
    ruta_archivo = os.path.join(INFORMES_DIR, nombre_archivo)

    doc.save(ruta_archivo)
    conn.close()

    return ruta_archivo, nombre_archivo
